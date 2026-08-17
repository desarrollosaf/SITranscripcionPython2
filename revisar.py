#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
revisar_v2_modificado.py
==========
Interfaz web LOCAL para revisar y corregir los oradores de las sesiones
guardadas en la base de datos (sesiones.db).

Uso:
    python revisar_v2_modificado.py                  (usa sesiones.db en la carpeta actual)
    python revisar_v2_modificado.py --db otra.db     (otra base de datos)

Se abre solo en tu navegador. Todo queda en tu computadora: no sube nada
a internet. Ctrl+C en la terminal para cerrarlo.

Si existe un archivo diputados.txt (un nombre por línea), esos nombres
aparecen como sugerencias al corregir.
"""

import argparse
import html
import html.parser
import json
import math
import mimetypes
import os
import re
import secrets
import sqlite3
import subprocess
import sys
import threading
import time
import webbrowser
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from urllib.parse import urlparse, parse_qs
from docx import Document
from docx.shared import Pt
from io import BytesIO

# ---------------------------------------------------------------------------
# Autenticación opcional (--requiere-login): reutiliza los mismos usuarios
# de MySQL y tokens JWT que la API (api/security.py, api/db_mysql.py). Solo
# se activa con la bandera --requiere-login (la usa el servicio Docker
# "revisor"), así que el uso local de siempre (python revisar.py, sin login,
# sin depender de MySQL) sigue funcionando igual.
# ---------------------------------------------------------------------------

COOKIE_SESION = "sesion_revisor"

# Todo lo demás (lanzar transcripciones, editar sesiones, trabajos de la
# API, usuarios) es de administrador. Una cuenta no-admin (agente de
# captura, corrector) solo debería poder tocar esto — la lista es un
# candado por defecto: lo que no está aquí, exige admin.
_RUTAS_NO_ADMIN = {
    "/api/yo", "/esteno", "/esteno/", "/esteno.html",
    "/api/sesiones",  # para el selector de sesión del corrector
    "/api/esteno/estado", "/api/esteno/segmentos", "/api/esteno/nombres",
    "/api/esteno/audio",
    "/api/esteno/tomar", "/api/esteno/latido", "/api/esteno/soltar",
    "/api/esteno/guardar",
    # Reutilizada por el editor del corrector para partir un turno en dos
    # oradores; el handler mismo exige que sea dentro de su bloque abierto.
    "/api/dividir",
}

# Se rellenan en main() solo si --requiere-login está activo.
_verificar_password = None
_obtener_usuario = None
_crear_token = None
_jwt = None
_jwt_secret = None
_jwt_algoritmo = None

PAGINA_LOGIN = """<!doctype html>
<html lang="es"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Iniciar sesión — revisión de oradores</title>
<style>
  :root{--papel:#FBFAF6;--tinta:#23241F;--tenue:#6E6C63;--verde:#1E5A38;
        --verde-suave:#E9F0EA;--linea:#E3E1D7;--ambar:#A9691F}
  *{box-sizing:border-box}
  html,body{margin:0;height:100%;background:var(--papel);color:var(--tinta);
    font:15px/1.5 -apple-system,"Segoe UI",sans-serif;
    display:flex;align-items:center;justify-content:center}
  form{background:#fff;border:1px solid var(--linea);border-radius:10px;
    padding:32px;width:320px;box-shadow:0 2px 10px rgba(0,0,0,.04)}
  h1{font-size:18px;margin:0 0 20px}
  label{display:block;font-size:13px;color:var(--tenue);margin:14px 0 4px}
  input{width:100%;padding:9px 10px;border:1px solid var(--linea);
    border-radius:6px;font-size:14px}
  button{margin-top:20px;width:100%;padding:10px;border:0;border-radius:6px;
    background:var(--verde);color:#fff;font-size:14px;cursor:pointer}
  button:hover{opacity:.92}
  .err{color:var(--ambar);font-size:13px;margin-top:12px;min-height:16px}
</style></head>
<body>
  <form id="f">
    <h1>Revisión de oradores</h1>
    <label>Email</label>
    <input type="email" name="email" required autofocus>
    <label>Contraseña</label>
    <input type="password" name="password" required>
    <div class="err" id="err"></div>
    <button type="submit">Entrar</button>
  </form>
<script>
document.getElementById('f').addEventListener('submit', async (e) => {
  e.preventDefault();
  const fd = new FormData(e.target);
  const r = await fetch('/login', {method:'POST', body: JSON.stringify({
    email: fd.get('email'), password: fd.get('password')})});
  if (r.ok) { location.href = '/'; }
  else {
    const d = await r.json().catch(() => ({}));
    document.getElementById('err').textContent =
      d.error || 'Email o contraseña incorrectos';
  }
});
</script>
</body></html>"""


# ---------------------------------------------------------------------------
# Catálogo opcional de diputados
# ---------------------------------------------------------------------------

def cargar_catalogo():
    base = os.path.dirname(os.path.abspath(__file__))
    for ruta in (os.path.join(base, "diputados.txt"), "diputados.txt"):
        if os.path.isfile(ruta):
            with open(ruta, encoding="utf-8") as f:
                return [ln.strip() for ln in f
                        if ln.strip() and not ln.strip().startswith("#")]
    return []


# ---------------------------------------------------------------------------
# Página (HTML + CSS + JS en un solo bloque)
# ---------------------------------------------------------------------------

PAGINA = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Versión estenográfica — revisión de oradores</title>
<style>
:root{
  --papel:#FAF8F2; --panel:#FFFFFF; --tinta:#20221C; --tinta-suave:#494A41;
  --tenue:#63625A; --verde:#1E5A38; --verde-hondo:#154029;
  --verde-suave:#E7EFE9; --linea:#E4E1D6; --linea-fuerte:#D2CFC2;
  --ambar:#9C5E18; --ambar-suave:#FAF0DE; --blanco:#FFFFFF;
  --sombra:0 1px 2px rgba(30,40,25,.05), 0 6px 20px rgba(30,40,25,.05);
  --sombra-barra:0 6px 14px -8px rgba(30,40,25,.28);
  --r:8px; --r-sm:6px;
}
*{box-sizing:border-box}
html,body{margin:0;padding:0;background:var(--papel);color:var(--tinta);
  font:15.5px/1.55 -apple-system,"Segoe UI",Roboto,"Helvetica Neue",Arial,sans-serif}
:focus-visible{outline:2px solid var(--ambar);outline-offset:2px}
@media (prefers-reduced-motion: reduce){*{transition:none!important}}
.oculto{display:none!important}

header{border-bottom:1px solid var(--linea-fuerte);background:var(--panel);
  padding:14px 24px;display:flex;flex-wrap:wrap;gap:8px 18px;align-items:center}
header .marca{display:flex;flex-direction:column;gap:2px;line-height:1}
header h1{font:700 22px/1.05 Georgia,"Times New Roman",serif;margin:0;
  color:var(--verde);letter-spacing:.01em}
header .sub{color:var(--tenue);font-size:11.5px;text-transform:uppercase;
  letter-spacing:.16em;font-weight:600}
header .controles{margin-left:auto;display:flex;gap:10px;flex-wrap:wrap;
  align-items:center}
select,input[type=search],input[type=text]{font:inherit;font-size:14.5px;padding:8px 11px;
  border:1px solid var(--linea-fuerte);border-radius:var(--r-sm);
  background:var(--panel);color:var(--tinta);max-width:340px}
select:hover,input[type=search]:hover,input[type=text]:hover{border-color:var(--verde)}
#buscar{min-width:200px}
.combo-sesion{position:relative;min-width:220px}
.combo-sesion input{width:100%;font-weight:600;min-width:170px}
.combo-sesion input:focus{border-color:var(--verde);outline:2px solid var(--verde-suave)}
.combo-lista{position:absolute;top:calc(100% + 4px);left:0;right:0;z-index:30;
  max-height:320px;overflow-y:auto;background:var(--panel);
  border:1px solid var(--linea-fuerte);border-radius:var(--r-sm);
  box-shadow:var(--sombra-barra)}
.combo-opcion{padding:8px 11px;font-size:14px;cursor:pointer;font-weight:400}
.combo-opcion:hover,.combo-opcion.resaltado{background:var(--verde-suave);color:var(--verde-hondo)}
.combo-opcion.vacio{color:var(--tenue);cursor:default}
.combo-opcion.vacio:hover{background:none;color:var(--tenue)}
.ver{display:inline-block;margin-left:8px;padding:1px 7px;border-radius:999px;
  background:var(--verde-suave);color:var(--verde);font-size:10px;
  letter-spacing:.06em;vertical-align:middle}

#meta{padding:11px 24px;color:var(--tenue);font-size:13.5px;
  border-bottom:1px solid var(--linea);background:var(--papel)}
#meta a{color:var(--verde);font-weight:600;text-decoration:none}
#meta a:hover{text-decoration:underline}
#meta strong{color:var(--tinta);font-weight:700}

.cuerpo{display:grid;grid-template-columns:250px minmax(0,1fr);gap:0;
  max-width:1160px;margin:0 auto}
aside{padding:20px 18px;border-right:1px solid var(--linea)}
aside h2{font-size:11px;text-transform:uppercase;letter-spacing:.14em;
  color:var(--tenue);margin:0 0 12px;font-weight:700}
#listaOradores{list-style:none;margin:0;padding:0}
#listaOradores li{margin:2px 0}
#listaOradores button{width:100%;text-align:left;font:inherit;font-size:13.5px;
  border:0;background:none;padding:7px 10px;border-radius:var(--r-sm);
  cursor:pointer;color:var(--tinta);display:flex;justify-content:space-between;
  gap:8px;transition:background .12s}
#listaOradores button:hover{background:var(--verde-suave)}
#listaOradores button.activo{background:var(--verde);color:#fff}
#listaOradores .min{color:var(--tenue);font-variant-numeric:tabular-nums}
#listaOradores button.activo .min{color:#D7E4DA}
aside .nota{font-size:12.5px;line-height:1.55;color:var(--tenue);
  margin-top:18px;padding-top:14px;border-top:1px solid var(--linea)}

main{padding:26px 30px 90px;max-width:820px}
.turno{display:grid;grid-template-columns:78px minmax(0,1fr);gap:16px;
  margin:0 0 22px;padding-bottom:2px}
.cuerpo-turno{min-width:0}
.cabecera-turno{display:flex;flex-wrap:wrap;align-items:center;gap:6px 10px;
  margin-bottom:6px}
.tiempo{font-size:12px;color:var(--tenue);text-decoration:none;
  font-variant-numeric:tabular-nums;padding-top:5px;white-space:nowrap}
.tiempo:hover{color:var(--ambar)}
.orador{font:700 13px/1.3 -apple-system,"Segoe UI",Roboto,sans-serif;
  text-transform:uppercase;letter-spacing:.05em;color:var(--verde);
  background:none;border:0;border-bottom:1px dashed var(--verde);
  padding:0 0 1px;cursor:pointer;text-align:left}
.orador:hover{color:var(--ambar);border-color:var(--ambar)}
.voz-pista{font:600 12px/1.4 -apple-system,"Segoe UI",Roboto,sans-serif;
  padding:2px 9px;border-radius:999px;cursor:pointer;white-space:nowrap;
  border:1px solid var(--ambar);background:#FBF3E4;color:var(--ambar)}
.voz-pista:hover{background:var(--ambar);color:#fff}
.insignia-ia{font:700 11px/1.4 -apple-system,"Segoe UI",Roboto,sans-serif;
  padding:2px 9px;border-radius:999px;white-space:nowrap;letter-spacing:.02em}
.insignia-ia.validado{background:var(--verde-suave);color:var(--verde);
  border:1px solid var(--verde)}
.insignia-ia.descartado{background:#F1EFE8;color:var(--tenue);
  border:1px solid var(--linea)}
.insignia-ia.media{background:#FBF3E4;color:var(--ambar);
  border:1px solid var(--ambar)}
.motivo-ia{font:italic 12px/1.4 Georgia,"Times New Roman",serif;
  color:var(--tenue);max-width:360px;overflow:hidden;text-overflow:ellipsis;
  white-space:nowrap;cursor:help}
/* Barra de acciones: fija, agrupada por etapa del flujo de trabajo */
.barra{position:sticky;top:0;z-index:20;display:flex;flex-wrap:wrap;
  align-items:center;gap:8px 14px;padding:9px 24px;background:var(--panel);
  border-bottom:1px solid var(--linea-fuerte);box-shadow:var(--sombra-barra)}
.grupo{display:flex;align-items:center;gap:8px;position:relative;
  padding-right:14px}
.grupo:not(:last-of-type)::after{content:"";position:absolute;right:0;top:50%;
  transform:translateY(-50%);width:1px;height:22px;background:var(--linea)}
.grupo-tit{font:700 10.5px/1 -apple-system,"Segoe UI",Roboto,sans-serif;
  text-transform:uppercase;letter-spacing:.12em;color:var(--tenue);
  padding-right:2px}
.grupo-tit .paso{display:inline-flex;align-items:center;justify-content:center;
  width:16px;height:16px;border-radius:999px;background:var(--verde-suave);
  color:var(--verde);font-size:10px;margin-right:5px}

.accion{font:600 14px/1.2 -apple-system,"Segoe UI",Roboto,sans-serif;
  padding:8px 13px;border-radius:var(--r-sm);cursor:pointer;
  border:1px solid var(--linea-fuerte);background:var(--panel);
  color:var(--tinta);transition:background .12s,border-color .12s,color .12s;
  white-space:nowrap}
.accion:hover{border-color:var(--verde);background:var(--verde-suave);
  color:var(--verde-hondo)}
.accion.ia{border-color:#E1D2B4;background:var(--ambar-suave);color:var(--ambar)}
.accion.ia:hover{background:var(--ambar);border-color:var(--ambar);color:#fff}
.accion.primario{border-color:var(--verde);background:var(--verde);color:#fff}
.accion.primario:hover{background:var(--verde-hondo);border-color:var(--verde-hondo)}
.accion:disabled{opacity:.55;cursor:default;filter:saturate(.6)}
#listaOradores button.pendiente{color:var(--ambar);font-weight:700}
#listaOradores button.pendiente.activo{background:var(--ambar);color:#fff}
.mini{font:600 12px/1.4 -apple-system,"Segoe UI",Roboto,sans-serif;
  padding:2px 9px;border-radius:999px;cursor:pointer;white-space:nowrap;
  border:1px solid var(--linea);background:#fff;color:var(--tinta-suave)}
.mini:hover{border-color:var(--verde);color:var(--verde)}
.mini.guardado{border-color:var(--verde);background:var(--verde-suave);
  color:var(--verde);font-weight:700}
.vivo{margin-left:auto;display:flex;align-items:center;gap:7px;font-size:14px;
  font-weight:600;color:var(--tinta-suave);cursor:pointer;white-space:nowrap;
  padding:6px 10px;border:1px solid var(--linea);border-radius:999px}
.vivo:hover{border-color:var(--verde);color:var(--verde)}
.vivo input{accent-color:var(--verde);width:15px;height:15px}
.editor-texto{grid-column:1 / -1;margin-top:10px;padding:12px;
  border:1px solid var(--linea);border-radius:8px;background:#FCFAF5}
.editor-texto label{display:flex;gap:10px;margin-bottom:8px;
  align-items:flex-start}
.seg-tiempo{font:600 12px/2 -apple-system,"Segoe UI",Roboto,sans-serif;
  color:var(--verde);white-space:nowrap;min-width:64px}
.editor-texto textarea{flex:1;font:15px/1.55 Georgia,serif;
  color:var(--tinta);padding:7px 10px;border:1px solid var(--linea);
  border-radius:6px;resize:vertical;background:#fff}
.editor-texto textarea:focus{outline:2px solid var(--verde-suave);
  border-color:var(--verde)}
.divisor{grid-column:1 / -1;margin-top:10px;padding:12px;
  border:1px solid var(--linea);border-radius:8px;background:#FCFAF5}
.divisor .ayuda{margin:0 0 8px;color:var(--tenue);font-size:13px}
.divisor .frases{font:15px/1.8 Georgia,"Times New Roman",serif}
.divisor .frase{cursor:pointer;padding:1px 3px;border-radius:4px}
.divisor .frase:hover{background:var(--verde-suave)}
.divisor .frase.p2{background:#F4E7CE}
.split-oradores{display:flex;flex-wrap:wrap;gap:10px;align-items:center;
  margin-top:12px}
.split-oradores label{display:flex;gap:6px;align-items:center;
  font-size:13px;color:var(--tenue)}
.split-oradores input{font:inherit;padding:6px 9px;
  border:1px solid var(--linea);border-radius:6px}
.split-oradores .b-div-ok{font:inherit;font-size:13.5px;padding:7px 12px;
  border-radius:6px;border:1px solid var(--verde);background:var(--verde);
  color:#fff;cursor:pointer}
.fila-botones{display:flex;gap:8px;align-items:center;margin-top:4px}
.fila-botones strong{margin-right:auto;color:var(--verde)}
.panel-resumen{grid-column:1 / -1;margin-top:12px;padding:18px 22px;
  border-radius:10px;border:1px solid var(--verde);
  background:var(--verde-suave);max-width:100%}
.panel-resumen .cabecera{display:flex;gap:8px;align-items:center;
  margin-bottom:10px}
.panel-resumen .cabecera strong{margin-right:auto;color:var(--verde);
  font:700 15px/1.3 -apple-system,"Segoe UI",Roboto,sans-serif}
.panel-resumen .md{font:15.5px/1.65 Georgia,"Times New Roman",serif;
  color:var(--tinta);overflow-wrap:break-word}
.panel-resumen .md h1,.panel-resumen .md h2,.panel-resumen .md h3{
  font:700 16px/1.35 -apple-system,"Segoe UI",Roboto,sans-serif;
  color:var(--verde);margin:14px 0 6px}
.panel-resumen .md h1{font-size:18px}
.panel-resumen .md p{margin:8px 0}
.panel-resumen .md strong{color:var(--tinta);font-weight:700}
.panel-resumen .md ul{margin:8px 0;padding-left:22px}
.panel-resumen .md li{margin:3px 0}
.panel-resumen .md hr{border:0;border-top:1px solid var(--verde);
  margin:12px 0;opacity:.4}
.texto{font:16.5px/1.7 Georgia,"Times New Roman",serif;margin:6px 0 0;
  text-align:justify;hyphens:auto;color:var(--tinta)}
.seccion-od{font:700 15.5px/1.45 Georgia,"Times New Roman",serif;
  color:var(--verde);text-align:justify;
  margin:26px 0 10px;padding:8px 0 6px;border-top:2px solid var(--verde);
  border-bottom:1px solid var(--linea)}
.resaltado{background:#F4E7CE}

.editor{grid-column:1 / -1;background:var(--blanco);border:1px solid var(--linea);
  border-left:3px solid var(--verde);border-radius:6px;padding:12px;
  margin-top:10px;display:flex;flex-wrap:wrap;gap:8px;align-items:center}
.editor input{flex:1 1 240px;font:inherit;padding:7px 10px;
  border:1px solid var(--linea);border-radius:6px}
.editor button{font:inherit;font-size:13.5px;padding:7px 12px;border-radius:6px;
  border:1px solid var(--verde);background:var(--verde);color:#fff;cursor:pointer}
.editor button:hover{filter:brightness(1.1)}
.editor .secundario{background:var(--blanco);color:var(--verde)}
.editor .cancelar{border-color:var(--linea);background:var(--blanco);
  color:var(--tenue)}

.vacio{color:var(--tenue);padding:40px 0;text-align:center}
#aviso{position:fixed;left:50%;bottom:24px;transform:translateX(-50%) translateY(80px);
  background:var(--tinta);color:#fff;padding:10px 18px;border-radius:8px;
  font-size:14px;transition:transform .25s;pointer-events:none}
#aviso.visible{transform:translateX(-50%) translateY(0)}


.lanzador{border:1px solid var(--linea);border-radius:var(--r);
  background:var(--panel);margin:16px 24px 0;padding:0;box-shadow:var(--sombra)}
.lanzador>summary{cursor:pointer;padding:13px 18px;
  font:700 14px/1.2 -apple-system,"Segoe UI",Roboto,sans-serif;
  color:var(--verde);list-style:none;display:flex;align-items:center;gap:8px}
.lanzador>summary::-webkit-details-marker{display:none}
.lanzador>summary::before{content:"\\002B";display:inline-flex;
  align-items:center;justify-content:center;width:20px;height:20px;
  border-radius:999px;background:var(--verde-suave);color:var(--verde);
  font-weight:700}
.lanzador[open]>summary::before{content:"\\2212"}
.lanzador .cuerpo-lan{padding:6px 18px 18px;display:flex;flex-wrap:wrap;
  gap:14px;align-items:flex-end;border-top:1px solid var(--linea)}
.lanzador .campo{display:flex;flex-direction:column;gap:5px;font-size:12.5px;
  font-weight:600;color:var(--tenue)}
.lanzador .campo input,.lanzador .campo select{max-width:none;font-weight:400}
.lanzador #lanUrl{min-width:320px}
.lanzador #lanComisiones{min-width:260px;min-height:96px}
#camposUrl,#camposEvento{flex-basis:100%;display:flex;flex-wrap:wrap;
  gap:14px 18px;align-items:flex-end;padding:14px 16px;margin-top:2px;
  background:var(--papel);border:1px solid var(--linea);border-radius:var(--r-sm)}
.lanzador .cuerpo-lan>.campo:first-child{flex-basis:100%}
.lanzador .cuerpo-lan>.campo:has(#lanModelo){margin-top:4px}
#btnTranscribir{margin-left:auto}
.lanzador .nota-lan{flex-basis:100%;font-size:12.5px;color:var(--tenue);margin:0}
.lanzador pre.log-lan{flex-basis:100%;margin:8px 0 0;max-height:220px;
  overflow:auto;background:#1d1f1b;color:#e7e9e3;padding:10px 12px;
  border-radius:var(--r-sm);font:12px/1.5 ui-monospace,Consolas,monospace;
  white-space:pre-wrap;word-break:break-word}
.lanzador .estado-lan{flex-basis:100%;font-size:13px;font-weight:600}
.lanzador .estado-lan.corriendo{color:var(--ambar)}
.lanzador .estado-lan.fin{color:var(--verde)}
#listaTrabajosApi,#listaUsuarios{flex-basis:100%}
.tabla-scroll{max-height:280px;overflow-y:auto;
  border:1px solid var(--linea);border-radius:var(--r-sm);margin-top:8px}
table.tabla-lista{width:100%;border-collapse:collapse;font-size:13px}
table.tabla-lista thead th{position:sticky;top:0;background:var(--panel);
  text-align:left;padding:7px 10px;font-size:11px;text-transform:uppercase;
  letter-spacing:.04em;color:var(--tenue);
  border-bottom:1px solid var(--linea-fuerte)}
table.tabla-lista tbody td{padding:7px 10px;border-bottom:1px solid var(--linea);
  vertical-align:top}
table.tabla-lista tbody tr:last-child td{border-bottom:none}
table.tabla-lista tbody tr:hover{background:var(--verde-suave)}
table.tabla-lista .accion{font-size:12.5px;padding:4px 10px}

@media (max-width:860px){
  header{padding:12px 16px}
  header .controles{width:100%;margin-left:0}
  .combo-sesion,#buscar{flex:1 1 45%;min-width:0;max-width:none}
  .barra{padding:9px 16px;gap:8px 10px;top:0}
  .grupo{padding-right:10px}
  .grupo-tit{display:none}
  .vivo{margin-left:0}
  .cuerpo{grid-template-columns:1fr}
  aside{border-right:0;border-bottom:1px solid var(--linea)}
  #listaOradores{display:flex;flex-wrap:wrap;gap:4px}
  #listaOradores li{flex:0 0 auto}
  #listaOradores button{border:1px solid var(--linea);border-radius:999px;
    padding:5px 13px}
  .turno{grid-template-columns:1fr}
  .editor{grid-column:1}
  main{padding:20px 16px 90px}
  .lanzador{margin:14px 16px 0}
  .lanzador #lanUrl{min-width:0;width:100%}
}
@media (max-width:560px){
  .grupo{padding-right:0}
  .grupo:not(:last-of-type)::after{display:none}
  .accion{flex:1 1 auto}
}
</style>
</head>
<body>
<header>
  <div class="marca">
    <h1>Versión estenográfica</h1>
    <span class="sub">Revisión de oradores <span class="ver">diseño v9</span></span>
  </div>
  <div class="controles">
    <div class="combo-sesion" id="comboSesion">
      <input id="selSesionTxt" type="text" placeholder="— Elige una sesión —"
             autocomplete="off" aria-label="Sesión">
      <div class="combo-lista oculto" id="listaSesionesCombo"></div>
    </div>
    <input id="buscar" type="search" placeholder="Buscar en el texto…"
           aria-label="Buscar en el texto">
    <a id="btnSalir" href="/logout" class="accion" style="display:none;text-decoration:none">Cerrar sesión</a>
  </div>
</header>
<nav class="barra" aria-label="Acciones de la sesión">
  <div class="grupo">
    <span class="grupo-tit"><span class="paso">1</span>Preparar</span>
    <button id="btnCompactar" class="accion"
            title="Une en la base de datos los registros consecutivos del mismo orador">Unir iguales</button>
  </div>
  <div class="grupo">
    <span class="grupo-tit"><span class="paso">2</span>Revisar con IA</span>
    <button id="btnCorregirEstilo" class="accion ia"
            title="Pasa todo el texto por IA para corregir ortografía, gramática y separar párrafos">✨ Corregir estilo</button>
    <button id="btnEstructurar" class="accion ia"
            title="La IA detecta dónde empieza cada punto del orden del día y los numera (1., 2., 3.)">🗂 Numerar orden del día</button>
  </div>
  <div class="grupo">
    <span class="grupo-tit"><span class="paso">3</span>Exportar</span>
    <button id="btnExportarWord" class="accion primario"
            title="Descargar el documento Word formateado">📄 Descargar Word</button>
  </div>
  <div class="grupo">
    <span class="grupo-tit">Colaborar</span>
    <a id="lnkEsteno" class="accion" href="/esteno"
       title="Reparte la sesión en tramos entre varios correctores">👥 Módulo estenográfico</a>
  </div>
  <label class="vivo" title="Sigue la transcripción casi en tiempo real (se actualiza cada ~2.5 s y se pausa mientras editas). Desmárcala para pausar.">
    <input type="checkbox" id="chkVivo" checked> Auto-actualizar</label>
</nav>
<details class="lanzador" id="lanzador">
  <summary>Nueva transcripción</summary>
  <div class="cuerpo-lan">
    <label class="campo">Origen
      <select id="lanOrigen">
        <option value="url">URL manual</option>
        <option value="evento">Evento del Congreso</option>
      </select>
    </label>
    <div id="camposUrl">
      <label class="campo">URL del video o transmisión
        <input id="lanUrl" type="url" placeholder="https://www.youtube.com/watch?v=…">
      </label>
      <label class="campo">Tipo de sesión
        <select id="lanTipo">
          <option value="pleno">Pleno (sesión)</option>
          <option value="comision">Comisión</option>
        </select>
      </label>
      <label class="campo oculto" id="campoComisiones">Comisión(es) — Ctrl/Cmd para varias unidas
        <select id="lanComisiones" multiple></select>
      </label>
      <label class="campo oculto" id="campoFecha">Fecha (opcional)
        <input id="lanFecha" type="text" placeholder="AAAA-MM-DD" size="12">
      </label>
    </div>
    <div id="camposEvento" class="oculto">
      <label class="campo">Tipo
        <select id="evTipo">
          <option value="1">Sesión / Diputación permanente</option>
          <option value="0">Comisión</option>
        </select>
      </label>
      <label class="campo">Evento
        <select id="evEvento" style="min-width:360px"></select>
      </label>
      <p class="nota-lan" id="notaEvento">Los participantes y, si el evento
      trae liga de YouTube, la URL, salen automáticos de ahí — no hace
      falta escribir nada más.</p>
    </div>
    <label class="campo">Modelo
      <select id="lanModelo">
        <option value="tiny">tiny (rápido)</option>
        <option value="base">base</option>
        <option value="small" selected>small (recomendado)</option>
        <option value="medium">medium</option>
        <option value="large-v3">large-v3 (lento, preciso)</option>
      </select>
    </label>
    <button id="btnTranscribir" class="accion primario">Iniciar transcripción</button>
    <button id="btnDetener" class="accion oculto">Detener</button>
    <p class="nota-lan" id="notaLan"></p>
    <div class="estado-lan oculto" id="estadoLan"></div>
    <pre class="log-lan oculto" id="logLan"></pre>
  </div>
</details>
<details class="lanzador" id="panelTrabajosApi">
  <summary>Trabajos de la API (Evento del Congreso)</summary>
  <div class="cuerpo-lan">
    <p class="nota-lan">Los que se crearon eligiendo "Evento del Congreso"
    arriba. Puedes detenerlos desde aquí — un trabajo detenido desaparece
    de la lista del operador de audio (el .exe).</p>
    <button id="btnRefrescarTrabajosApi" class="accion">🔄 Actualizar</button>
    <div id="listaTrabajosApi"></div>
  </div>
</details>
<details class="lanzador oculto" id="panelUsuarios">
  <summary>Usuarios (login del agente de captura / correctores)</summary>
  <div class="cuerpo-lan">
    <p class="nota-lan">Crea aquí la cuenta que usará el operador para
    entrar al agente de captura (.exe) o un corrector para entrar al
    módulo estenográfico. La contraseña se genera sola y solo se muestra
    una vez — cópiala antes de cerrar el aviso, después no se puede
    recuperar (solo crear una nueva).</p>
    <label class="campo">Email (será el usuario para iniciar sesión)
      <input id="nuEmail" type="email" placeholder="operador1@congreso.local">
    </label>
    <label style="display:flex;align-items:center;gap:6px;font-size:13px">
      <input type="checkbox" id="nuEsAdmin"> Es administrador (puede crear
      trabajos, sesiones y otros usuarios — no marcar para el agente de
      captura ni correctores)
    </label>
    <button id="btnCrearUsuario" class="accion primario">Crear usuario</button>
    <div id="nuevaCredencial" class="oculto" style="margin-top:10px;padding:10px;
         border:1px solid var(--linea);border-radius:8px;background:var(--fondo-suave,#f6f6f6)">
    </div>
    <button id="btnRefrescarUsuarios" class="accion" style="margin-top:10px">🔄 Actualizar lista</button>
    <div id="listaUsuarios"></div>
  </div>
</details>
<div id="meta"></div>
<div class="cuerpo">
  <aside>
    <h2>Oradores de la sesión</h2>
    <ul id="listaOradores"></ul>
    <p class="nota">Haz clic en un <strong>nombre</strong> dentro del texto
    para corregirlo. La <strong>hora</strong> abre YouTube en ese minuto,
    para verificar quién habla.</p>
  </aside>
  <main id="transcripcion"></main>
</div>
<datalist id="dlOradores"></datalist>
<div id="aviso" role="status" aria-live="polite"></div>

<script>
if (__REQUIERE_LOGIN__) document.getElementById('btnSalir').style.display = '';
const $ = s => document.querySelector(s);
const estado = {sesiones:[], sesion:null, filas:[], turnos:[],
                catalogo:[], filtroOrador:null, q:''};

const esc = s => s.replace(/[&<>"]/g,
  c => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;'}[c]));

// El módulo estenográfico guarda texto con formato (HTML: <div>, <b>, etc,
// de su editor con negritas/cursivas — así se conserva el formato en el
// Word). Esta pantalla siempre trabajó con texto plano; sin esto, esas
// etiquetas salían tal cual como texto pegado en la transcripción.
function htmlAPlano(texto){
  if(!texto || texto.indexOf('<') === -1) return texto || '';
  const d = document.createElement('div');
  d.innerHTML = texto;
  d.querySelectorAll('div,p,br,li').forEach(el => {
    el.insertAdjacentText('beforebegin', ' ');
  });
  return (d.textContent || '').replace(/\\s+/g, ' ').trim();
}

// Convierte el Markdown del resumen (títulos #, negritas **, listas -/•,
// separadores ---) en HTML legible. Escapa todo primero para que sea seguro.
function mdAhtml(texto){
  const enlinea = s => esc(s)
    .replace(/\\*\\*([^*]+)\\*\\*/g, '<strong>$1</strong>')
    .replace(/\\*([^*]+)\\*/g, '<em>$1</em>');
  const lineas = (texto||'').replace(/\\r/g,'').split('\\n');
  let html = '', enLista = false;
  const cerrarLista = () => { if(enLista){ html += '</ul>'; enLista = false; } };
  for(let cruda of lineas){
    const l = cruda.trim();
    if(!l){ cerrarLista(); continue; }
    if(/^(-{3,}|_{3,}|\\*{3,})$/.test(l)){ cerrarLista(); html += '<hr>'; continue; }
    let m;
    if((m = l.match(/^(#{1,6})\\s+(.*)$/))){
      cerrarLista();
      const n = Math.min(m[1].length, 3);
      html += '<h'+n+'>'+enlinea(m[2])+'</h'+n+'>';
    } else if((m = l.match(/^[-*•]\\s+(.*)$/))){
      if(!enLista){ html += '<ul>'; enLista = true; }
      html += '<li>'+enlinea(m[1])+'</li>';
    } else {
      cerrarLista();
      html += '<p>'+enlinea(l)+'</p>';
    }
  }
  cerrarLista();
  return html;
}

const norm = s => s.normalize('NFD').replace(/[\\u0300-\\u036f]/g,'').toLowerCase();

// Selector de sesión con buscador (reemplaza al <select> nativo, que se
// vuelve incómodo con muchas sesiones). onChange(id) se dispara al elegir
// una opción de verdad (id=0 vacía la vista); setValor(id,false) solo
// pinta el texto sin disparar nada (para refrescos/auto-selección).
const ComboSesion = {
  opciones: [], filtro: '', resaltado: -1, valorActual: 0, onChange: null,
  init(){
    this.elInput = $('#selSesionTxt'); this.elLista = $('#listaSesionesCombo');
    this.elInput.addEventListener('input', () => {
      this.filtro = this.elInput.value; this.resaltado = -1; this.pintar();
    });
    this.elInput.addEventListener('focus', () => {
      this.filtro = ''; this.elInput.select(); this.pintar();
    });
    this.elInput.addEventListener('blur', () => setTimeout(() => this.cerrar(), 150));
    this.elInput.addEventListener('keydown', e => this.tecla(e));
  },
  cargar(opciones){ this.opciones = opciones; this.setValor(this.valorActual, false); },
  filtradas(){
    if(!this.filtro) return this.opciones;
    const f = norm(this.filtro);
    return this.opciones.filter(o => norm(o.label).includes(f));
  },
  pintar(){
    const lista = this.filtradas();
    this.elLista.innerHTML = lista.length
      ? lista.map((o,i) => '<div class="combo-opcion'+(i===this.resaltado?' resaltado':'')
          +'" data-id="'+o.id+'">'+esc(o.label)+'</div>').join('')
      : '<div class="combo-opcion vacio">Sin resultados</div>';
    this.elLista.classList.remove('oculto');
    this.elLista.querySelectorAll('.combo-opcion[data-id]').forEach(d => {
      d.onmousedown = e => { e.preventDefault(); this.elegir(+d.dataset.id); };
    });
  },
  cerrar(){
    this.elLista.classList.add('oculto');
    const actual = this.opciones.find(o => o.id === this.valorActual);
    this.elInput.value = actual ? actual.label : '';
    this.filtro = '';
  },
  tecla(e){
    const lista = this.filtradas();
    if(e.key === 'ArrowDown'){ e.preventDefault();
      this.resaltado = Math.min(lista.length-1, this.resaltado+1); this.pintar(); }
    else if(e.key === 'ArrowUp'){ e.preventDefault();
      this.resaltado = Math.max(0, this.resaltado-1); this.pintar(); }
    else if(e.key === 'Enter'){ e.preventDefault();
      const o = lista[this.resaltado]; if(o) this.elegir(o.id); }
    else if(e.key === 'Escape'){ this.elInput.blur(); }
  },
  elegir(id){ this.setValor(id, true); this.elLista.classList.add('oculto'); this.elInput.blur(); },
  setValor(id, disparar){
    this.valorActual = id;
    const o = this.opciones.find(x => x.id === id);
    this.elInput.value = o ? o.label : '';
    if(disparar && this.onChange) this.onChange(id);
  },
};
const hmsDe = seg => { seg=Math.floor(seg);
  const h=String(Math.floor(seg/3600)).padStart(2,'0'),
        m=String(Math.floor(seg%3600/60)).padStart(2,'0'),
        s=String(seg%60).padStart(2,'0');
  return h+':'+m+':'+s; };

async function api(ruta, datos){
  const r = datos
    ? await fetch(ruta,{method:'POST',
        headers:{'Content-Type':'application/json'},
        body:JSON.stringify(datos)})
    : await fetch(ruta);
  return r.json();
}

function avisar(msg){
  const a = $('#aviso'); a.textContent = msg; a.classList.add('visible');
  clearTimeout(a._t); a._t = setTimeout(()=>a.classList.remove('visible'), 2600);
}

function enlaceYT(seg){
  const url = estado.sesion && estado.sesion.url;
  if(!url || !url.startsWith('http')) return null;
  return url + (url.includes('?') ? '&' : '?') + 't=' + Math.floor(seg) + 's';
}

function agrupar(filas){
  const turnos = []; let t = null;
  for(const f of filas){
    if(!t || t.orador !== f.orador){
      t = {orador:f.orador, inicio:f.inicio_seg, ids:[], textos:[],
           voz:null, vozSim:0, revIa:null, motivoIa:''};
      turnos.push(t);
    }
    t.ids.push(f.id); t.textos.push(htmlAPlano(f.texto)); t.fin = f.fin_seg;
    if(f.voz_orador && (f.voz_similitud||0) > t.vozSim){
      t.voz = f.voz_orador; t.vozSim = f.voz_similitud||0;
    }
    // Veredicto del corrector: prioridad validado > media > descartado
    const _rango = {validado:3, media:2, descartado:1};
    if(f.revisado_ia && (_rango[f.revisado_ia]||0) > (_rango[t.revIa]||0)){
      t.revIa = f.revisado_ia; t.motivoIa = f.motivo_ia || '';
    }
  }
  return turnos;
}

function pintarLateral(){
  const acc = {};
  for(const f of estado.filas){
    const a = acc[f.orador] || (acc[f.orador] = {seg:0});
    a.seg += Math.max(0, f.fin_seg - f.inicio_seg);
  }
  const orden = Object.entries(acc).sort((x,y)=>{
    const dx = x[0]==='Desconocido' ? 1 : 0, dy = y[0]==='Desconocido' ? 1 : 0;
    if(dx !== dy) return dy - dx;          // Desconocido siempre primero
    return y[1].seg - x[1].seg;
  });
  $('#listaOradores').innerHTML = orden.map(([nom,a]) =>
    '<li><button data-orador="'+esc(nom)+'"'
    + ' class="'+(estado.filtroOrador===nom ? 'activo ' : '')
    + (nom==='Desconocido' ? 'pendiente' : '')+'">'
    + '<span>'+(nom==='Desconocido' ? '⚠ ' : '')+esc(nom)+'</span>'
    + '<span class="min">'+Math.round(a.seg/60)+' min</span></button></li>'
  ).join('');
  document.querySelectorAll('#listaOradores button').forEach(b =>
    b.onclick = () => {
      const n = b.dataset.orador;
      estado.filtroOrador = (estado.filtroOrador===n) ? null : n;
      pintarLateral(); pintarTranscripcion();
    });
}

function pintarTranscripcion(){
  const q = norm(estado.q||'');
  const visibles = estado.turnos.filter(t =>
    (!estado.filtroOrador || t.orador===estado.filtroOrador) &&
    (!q || norm(t.textos.join(' ')).includes(q) || norm(t.orador).includes(q)));
  if(!visibles.length){
    $('#transcripcion').innerHTML =
      '<p class="vacio">' + (estado.turnos.length
        ? 'Sin resultados con este filtro.'
        : 'Esta sesión no tiene participaciones todavía.') + '</p>';
    return;
  }
  let htmlTurnos = visibles.map((t,i) => {
    const yt = enlaceYT(t.inicio);
    const tiempo = yt
      ? '<a class="tiempo" target="_blank" rel="noopener" href="'+yt
        +'" title="Ver este momento en YouTube">'+hmsDe(t.inicio)+' ▸</a>'
      : '<span class="tiempo">'+hmsDe(t.inicio)+'</span>';
    let texto = esc(t.textos.join(' '));
    if(q){
      const rex = new RegExp('('+estado.q.replace(/[.*+?^${}()|[\\]\\\\]/g,'\\\\$&')+')','gi');
      texto = texto.replace(rex,'<mark class="resaltado">$1</mark>');
    }
    const encabezado = (estado.seccionesPorAncla
                        && estado.seccionesPorAncla[t.ids[0]])
      ? '<h3 class="seccion-od">'
        + esc(estado.seccionesPorAncla[t.ids[0]])+'</h3>'
      : '';
    return encabezado
      + '<article class="turno" data-i="'+estado.turnos.indexOf(t)+'">'
      + tiempo
      + '<div class="cuerpo-turno"><div class="cabecera-turno">'
      + '<button class="orador" title="Corregir orador">'
      + esc(t.orador)+':</button>'
      + (t.revIa
         ? '<span class="insignia-ia '+t.revIa+'" title="'
           + esc(t.motivoIa || '')+'">'
           + (t.revIa==='validado' ? '✓ validado por IA'
              : t.revIa==='media' ? '◐ validado (media) por IA'
              : '✗ descartado por IA')+'</span>'
           + (t.motivoIa
              ? '<span class="motivo-ia" title="'+esc(t.motivoIa)+'">'
                + esc(t.motivoIa)+'</span>'
              : '')
         : '')
      + ((t.voz && t.voz !== t.orador)
         ? '<button class="voz-pista" title="Aplicar esta sugerencia de voz">'
           + '🎙 voz: '+esc(t.voz)+' ('+t.vozSim.toFixed(2)+') · aplicar</button>'
         : '')
      + '<button class="mini b-editar" title="Corregir el texto de este bloque">✏ texto</button>'
      + '<button class="mini b-dividir" title="Separar dos oradores dentro de este bloque">✂ dividir</button>'
      + '<button class="mini b-resumen'
        + ((estado.resumenes && estado.resumenes[t.ids[0]]) ? ' guardado' : '')
        + '" title="Resumen ejecutivo de esta intervención">📋 resumen'
        + ((estado.resumenes && estado.resumenes[t.ids[0]]) ? ' ✓' : '')
        + '</button>'
      + '</div>'
      + '<p class="texto">'+texto+'</p></div></article>';
  }).join('');
  $('#transcripcion').innerHTML = htmlTurnos;
  document.querySelectorAll('.orador').forEach(b =>
    b.onclick = e => abrirEditor(e.target.closest('.turno')));
  document.querySelectorAll('.voz-pista').forEach(b =>
    b.onclick = async e => {
      const t = estado.turnos[+e.target.closest('.turno').dataset.i];
      await asignar(t, t.voz, 'Sugerencia de voz aplicada');
    });
  document.querySelectorAll('.b-editar').forEach(b =>
    b.onclick = e => editarTexto(e.target.closest('.turno')));
  document.querySelectorAll('.b-dividir').forEach(b =>
    b.onclick = e => dividirBloque(e.target.closest('.turno')));
  document.querySelectorAll('.b-resumen').forEach(b =>
    b.onclick = e => pedirResumen(e.target.closest('.turno'), e.target));
  // Mostrar automáticamente los resúmenes ya guardados
  if(estado.resumenes){
    document.querySelectorAll('.turno').forEach(art => {
      const t = estado.turnos[+art.dataset.i];
      if(t && estado.resumenes[t.ids[0]])
        mostrarResumen(art, t, estado.resumenes[t.ids[0]]);
    });
  }
}

function editarTexto(art){
  document.querySelectorAll('.editor, .editor-texto, .panel-resumen')
    .forEach(e => e.remove());
  const t = estado.turnos[+art.dataset.i];
  const cont = document.createElement('div');
  cont.className = 'editor-texto';
  cont.innerHTML = t.ids.map((id, k) => {
    const f = estado.porId[id] || {};
    return '<label><span class="seg-tiempo">'+hmsDe(f.inicio_seg||0)+'</span>'
      + '<textarea data-id="'+id+'" rows="2">'+esc(t.textos[k]||'')
      + '</textarea></label>';
  }).join('')
  + '<div class="fila-botones">'
  + '<button class="b-guardar">Guardar texto</button>'
  + '<button class="b-cancelar cancelar">Cancelar</button></div>';
  art.appendChild(cont);
  cont.querySelector('textarea').focus();
  cont.querySelector('.b-cancelar').onclick = () => cont.remove();
  cont.querySelector('.b-guardar').onclick = async () => {
    const cambios = [];
    cont.querySelectorAll('textarea').forEach((ta, k) => {
      if(ta.value !== (t.textos[k]||''))
        cambios.push({id:+ta.dataset.id, texto:ta.value});
    });
    if(!cambios.length){ cont.remove(); return; }
    const r = await api('/api/textos', {cambios});
    avisar('Texto corregido: '+(r.cambios||0)+' segmento(s).');
    cargarSesion(estado.sesion.id);
  };
}

function dividirBloque(art){
  document.querySelectorAll('.editor, .editor-texto, .panel-resumen, .divisor')
    .forEach(e => e.remove());
  const t = estado.turnos[+art.dataset.i];
  const full = t.textos.join(' ');
  // Partir el texto del bloque en frases, conservando el offset de cada una
  const frases = []; let m; const rex = /[^.?!]+[.?!]*\\s*/g;
  while((m = rex.exec(full)) !== null){
    if(m[0].trim()) frases.push({txt:m[0].trim(), off:m.index});
    if(rex.lastIndex === m.index) rex.lastIndex++;
  }
  if(frases.length < 2){
    avisar('Bloque muy corto para dividir por frases; usa "✏ texto".');
    return;
  }
  const cont = document.createElement('div');
  cont.className = 'divisor';
  cont.innerHTML =
    '<p class="ayuda">Haz clic en la frase donde empieza el '
    + '<strong>segundo orador</strong>:</p>'
    + '<div class="frases">'
    + frases.map((f,k) => '<span class="frase" data-k="'+k+'" data-off="'
        + f.off+'">'+esc(f.txt)+'</span>').join(' ')
    + '</div>'
    + '<div class="split-oradores" hidden>'
    + '<label>1ª parte <input class="o1" list="dlOradores"></label>'
    + '<label>2ª parte <input class="o2" list="dlOradores"></label>'
    + '<button class="b-div-ok">Dividir aquí</button>'
    + '<button class="b-cancelar cancelar">Cancelar</button></div>';
  art.appendChild(cont);
  let corte = null;
  const panelO = cont.querySelector('.split-oradores');
  cont.querySelectorAll('.frase').forEach(sp => sp.onclick = () => {
    const k = +sp.dataset.k;
    if(k === 0){ avisar('Elige una frase posterior a la primera.'); return; }
    corte = +sp.dataset.off;
    cont.querySelectorAll('.frase').forEach((s,j) =>
      s.classList.toggle('p2', j >= k));
    panelO.hidden = false;
    cont.querySelector('.o1').value = t.orador;
    cont.querySelector('.o2').focus();
  });
  cont.querySelector('.b-cancelar').onclick = () => cont.remove();
  cont.querySelector('.b-div-ok').onclick = async () => {
    const o1 = cont.querySelector('.o1').value.trim();
    const o2 = cont.querySelector('.o2').value.trim();
    if(corte === null || !o1 || !o2){
      avisar('Marca la frase y escribe los dos oradores.'); return;
    }
    const r = await api('/api/dividir', {sesion_id:estado.sesion.id,
      ids:t.ids, corte:corte, orador1:o1, orador2:o2});
    if(r && r.error){ avisar('No se pudo dividir: '+r.error); return; }
    avisar('Bloque dividido en dos oradores.');
    cargarSesion(estado.sesion.id);
  };
}

function mostrarResumen(art, t, resumen){
  art.querySelectorAll('.panel-resumen').forEach(e => e.remove());
  const p = document.createElement('div');
  p.className = 'panel-resumen';
  p.innerHTML =
    '<div class="cabecera"><strong>Resumen ejecutivo</strong>'
    + '<button class="mini b-copiar">Copiar</button>'
    + '<button class="mini b-borrar">Borrar</button>'
    + '<button class="mini b-cerrar">Cerrar</button></div>'
    + '<div class="md">'+mdAhtml(resumen)+'</div>';
  art.appendChild(p);
  p.querySelector('.b-cerrar').onclick = () => p.remove();
  p.querySelector('.b-copiar').onclick = async () => {
    await navigator.clipboard.writeText(resumen);
    avisar('Resumen copiado al portapapeles.');
  };
  p.querySelector('.b-borrar').onclick = async () => {
    await api('/api/borrar_resumen', {ancla_id:t.ids[0]});
    delete estado.resumenes[t.ids[0]];
    p.remove();
    avisar('Resumen borrado.');
  };
}

async function pedirResumen(art, btn){
  const t = estado.turnos[+art.dataset.i];
  // Si ya hay uno guardado, mostrarlo sin volver a gastar API
  if(estado.resumenes && estado.resumenes[t.ids[0]]){
    mostrarResumen(art, t, estado.resumenes[t.ids[0]]);
    return;
  }
  const original = btn.textContent;
  btn.disabled = true; btn.textContent = 'Generando…';
  try{
    const r = await api('/api/resumen',
      {orador:t.orador, texto:t.textos.join(' '),
       sesion_id:estado.sesion.id, ancla_id:t.ids[0]});
    if(r.modo === 'auto' && r.resumen){
      estado.resumenes[t.ids[0]] = r.resumen;   // recordar en memoria
      mostrarResumen(art, t, r.resumen);
    } else if(r.prompt){
      await navigator.clipboard.writeText(r.prompt);
      avisar(r.error ? r.error
        : 'Prompt copiado: pégalo en Claude (claude.ai) para generar el '
          +'resumen. Para hacerlo automático configura tu clave API '
          +'(ver README).');
    }
  } finally {
    btn.disabled = false; btn.textContent = original;
  }
}

async function asignar(t, orador, mensaje){
  const r = await api('/api/actualizar',
    {ids:t.ids, orador:orador, sesion_id:estado.sesion.id});
  avisar(mensaje+': '+(r.cambios||0)+' segmento(s)'
    + (r.unidos ? ' · '+r.unidos+' unidos en la base' : '')+'.');
  cargarSesion(estado.sesion.id);
}

function opcionesOradores(){
  const set = new Set(['Presidencia / Mesa Directiva']);
  estado.catalogo.forEach(n => set.add('Dip. ' + n));
  estado.filas.forEach(f => set.add(f.orador));
  $('#dlOradores').innerHTML = [...set].sort().map(n =>
    '<option value="'+esc(n)+'">').join('');
}

function abrirEditor(art){
  document.querySelectorAll('.editor').forEach(e => e.remove());
  const i = +art.dataset.i;
  const t = estado.turnos[i];
  const prev = estado.turnos[i-1], next = estado.turnos[i+1];
  const ed = document.createElement('div');
  ed.className = 'editor';
  ed.innerHTML =
    '<input list="dlOradores" value="'+esc(t.orador)+'" aria-label="Nuevo orador">'
    +'<button class="b-bloque">Cambiar solo este bloque</button>'
    +(prev ? '<button class="b-prev secundario" title="Asignar: '
        +esc(prev.orador)+'">⬆ Unir con el anterior</button>' : '')
    +(next ? '<button class="b-next secundario" title="Asignar: '
        +esc(next.orador)+'">⬇ Unir con el siguiente</button>' : '')
    +'<button class="b-sesion secundario">Cambiar en toda la sesión</button>'
    +'<button class="b-cancelar cancelar">Cancelar</button>';
  art.appendChild(ed);
  const input = ed.querySelector('input');
  input.focus(); input.select();
  ed.querySelector('.b-cancelar').onclick = () => ed.remove();
  ed.querySelector('.b-bloque').onclick = async () => {
    const v = input.value.trim(); if(!v) return;
    await asignar(t, v, 'Listo');
  };
  if(prev) ed.querySelector('.b-prev').onclick = () =>
    asignar(t, prev.orador, 'Unido con '+prev.orador);
  if(next) ed.querySelector('.b-next').onclick = () =>
    asignar(t, next.orador, 'Unido con '+next.orador);
  ed.querySelector('.b-sesion').onclick = async () => {
    const v = input.value.trim(); if(!v) return;
    const r = await api('/api/renombrar',
      {sesion_id:estado.sesion.id, de:t.orador, a:v});
    avisar('Listo: '+(r.cambios||0)+' renombrados'
      + (r.unidos ? ' · '+r.unidos+' unidos en la base' : '')+'.');
    cargarSesion(estado.sesion.id);
  };
}

async function cargarSesion(id, opts){
  opts = opts || {};
  const d = await api('/api/participaciones?sesion='+id);
  const filas = d.filas || [];
  // Firma barata para detectar bloques nuevos (nº de filas + último id +
  // largo del último texto). Sirve para no redibujar si nada cambió.
  const ult = filas.length ? filas[filas.length-1] : null;
  const firma = filas.length + ':' + (ult ? ult.id+':'+((ult.texto||'').length) : '0');
  const mismaSesion = estado.sesion && estado.sesion.id === (d.sesion && d.sesion.id);
  if(opts.soloSiCambia && mismaSesion && firma === estado.firma){
    return false;   // sin novedades: no se redibuja
  }
  // Seguro extra: si abriste algo para editar mientras este refresco ya iba
  // en curso, NO redibujamos (perderías lo escrito). Dejamos la firma vieja
  // para que el próximo tic —cuando cierres el editor— muestre lo nuevo.
  if(opts.soloSiCambia &&
     document.querySelector('.editor, .editor-texto, .panel-resumen, .divisor')){
    return false;
  }
  estado.firma = firma;
  estado.sesion = d.sesion; estado.filas = filas;
  estado.resumenes = d.resumenes || {};
  estado.estructura = d.estructura || null;
  estado.seccionesPorAncla = {};
  if(estado.estructura && estado.estructura.secciones){
    estado.estructura.secciones.forEach(s => {
      estado.seccionesPorAncla[s.ancla_id] = s.titulo;
    });
  }
  estado.porId = {};
  filas.forEach(f => estado.porId[f.id] = f);
  estado.turnos = agrupar(filas);
  const s = d.sesion || {};
  $('#meta').innerHTML = s.id
    ? '<strong>'+esc(s.titulo||'Sesión '+s.id)+'</strong> · inicio '
      + esc(s.inicio||'?')
      + (s.url && s.url.startsWith('http')
         ? ' · <a href="'+esc(s.url)+'" target="_blank" rel="noopener">ver video</a>'
         : '')
    : '';
  const lnkEsteno = $('#lnkEsteno');
  if(lnkEsteno) lnkEsteno.href = s.id ? ('/esteno?sesion='+s.id) : '/esteno';
  opcionesOradores(); pintarLateral(); pintarTranscripcion();
  return true;    // se redibujó
}

// Deja la vista vacía (sin sesión cargada): al abrir la app o al elegir
// la opción "— Elige una sesión —" del selector.
function limpiarVista(){
  estado.sesion = null; estado.filas = []; estado.turnos = [];
  estado.filtroOrador = null; estado.estructura = null;
  estado.resumenes = {}; estado.seccionesPorAncla = {};
  $('#meta').innerHTML = '';
  $('#listaOradores').innerHTML = '';
  const dl = $('#dlOradores'); if(dl) dl.innerHTML = '';
  $('#transcripcion').innerHTML =
    '<p class="vacio">Elige una sesión en el menú de arriba para empezar.</p>';
}

async function iniciar(){
  ComboSesion.init();
  ComboSesion.onChange = id => { if(id) cargarSesion(id); else limpiarVista(); };
  estado.catalogo = await api('/api/catalogo');
  estado.sesiones = await api('/api/sesiones');
  if(!estado.sesiones.length){
    $('#transcripcion').innerHTML =
      '<p class="vacio">Aún no hay sesiones en la base de datos.<br>'
      +'Corre primero <code>transcribir_en_vivo.py</code> y vuelve aquí.</p>';
    return;
  }
  ComboSesion.cargar(estado.sesiones.map(s => ({id:s.id,
    label:'#'+s.id+' — '+(s.titulo||'').slice(0,60)+' ('+s.segmentos+' seg.)'})));
  $('#buscar').oninput = e => {estado.q = e.target.value; pintarTranscripcion();};
  
  $('#btnCompactar').onclick = async () => {
    if(!estado.sesion) return;
    const r = await api('/api/compactar', {sesion_id:estado.sesion.id});
    avisar(r.unidos ? r.unidos+' registro(s) unidos en la base.'
                    : 'No había registros consecutivos por unir.');
    cargarSesion(estado.sesion.id);
  };

  $('#btnExportarWord').onclick = () => {
    if(!estado.sesion) return avisar('No hay sesión activa.');
    window.open('/api/exportar_word?sesion=' + estado.sesion.id, '_blank');
  };

  $('#btnCorregirEstilo').onclick = async () => {
    if(!estado.sesion) return avisar('No hay sesión activa.');
    if(!confirm('¿Estás seguro? Esto enviará todos los segmentos a la IA para corrección ortográfica y gramatical. Tomará unos momentos.')) return;
    
    const btn = $('#btnCorregirEstilo');
    const txtOriginal = btn.textContent;
    btn.textContent = 'Procesando... ⏳';
    btn.disabled = true;
    
    try {
      const r = await api('/api/corregir_estilo', {sesion_id: estado.sesion.id});
      if(r.error) {
        avisar('Error: ' + r.error);
      } else {
        avisar('Revisión completada. ' + r.cambios + ' segmentos corregidos.'
               + (r.aviso ? ' ⚠ ' + r.aviso : ''));
        cargarSesion(estado.sesion.id);
      }
    } finally {
      btn.textContent = txtOriginal;
      btn.disabled = false;
    }
  };

  $('#btnEstructurar').onclick = async () => {
    if(!estado.sesion) return avisar('No hay sesión activa.');
    if(!confirm('La IA leerá toda la sesión para detectar dónde empieza '
      + 'cada punto del orden del día y numerarlos (1., 2., 3., …). '
      + 'No toca el texto ni identifica acuerdos. ¿Continuar?')) return;
    const btn = $('#btnEstructurar');
    const txt = btn.textContent;
    btn.textContent = 'Analizando... ⏳'; btn.disabled = true;
    try {
      const r = await api('/api/estructurar', {sesion_id: estado.sesion.id});
      if(r.error){ avisar('Error: ' + r.error); }
      else {
        avisar('Orden del día detectado: ' + r.secciones
          + ' punto(s) numerado(s). Ya aparece en la transcripción y en el Word.');
        cargarSesion(estado.sesion.id);
      }
    } finally { btn.textContent = txt; btn.disabled = false; }
  };

  async function autoRefrescoTick(){
    if(document.hidden) return;                 // pestaña en segundo plano
    // no interrumpir si hay algo abierto para editar
    if(document.querySelector('.editor, .editor-texto, .panel-resumen, .divisor'))
      return;
    if(!estado.sesion) return;
    const doc = document.documentElement;
    // ¿el usuario ya está pegado al final? (para "seguir" la transcripción)
    const cercaFinal = (window.innerHeight + window.scrollY) >= (doc.scrollHeight - 140);
    const y = window.scrollY;
    const cambio = await cargarSesion(estado.sesion.id, {soloSiCambia:true});
    if(cambio){
      // Si venías al final, salta a los bloques nuevos; si estabas leyendo
      // más arriba, respeta tu posición.
      window.scrollTo(0, cercaFinal ? document.documentElement.scrollHeight : y);
    }
  }
  function iniciarAutoRefresco(){
    clearInterval(estado.timerVivo);
    estado.timerVivo = setInterval(autoRefrescoTick, 2500);
  }

  $('#chkVivo').onchange = e => {
    if(e.target.checked){
      iniciarAutoRefresco();
      avisar('Auto-actualización activada: sigue la transcripción casi en '
        +'tiempo real (se pausa mientras editas).');
    } else {
      clearInterval(estado.timerVivo);
      avisar('Auto-actualización en pausa.');
    }
  };
  // Arranca sola: el interruptor viene marcado por defecto.
  if($('#chkVivo').checked) iniciarAutoRefresco();
  limpiarVista();   // arrancamos en blanco; el usuario elige la sesión
}

async function refrescarSelectorSesiones(){
  estado.sesiones = await api('/api/sesiones');
  if(!estado.sesiones.length) return;
  // cargar() vuelve a pintar con lo que ya tenía elegido (valorActual no se
  // toca), así que conserva la selección sin necesidad de guardarla aparte.
  ComboSesion.cargar(estado.sesiones.map(s => ({id:s.id,
    label:'#'+s.id+' — '+(s.titulo||'').slice(0,60)+' ('+s.segmentos+' seg.)'})));
}

function pintarEstadoLan(est){
  const elLog = $('#logLan'), elEst = $('#estadoLan');
  if(est.lineas){
    elLog.classList.remove('oculto');
    const abajo = elLog.scrollTop + elLog.clientHeight >= elLog.scrollHeight - 20;
    elLog.textContent = est.lineas;
    if(abajo) elLog.scrollTop = elLog.scrollHeight;
  }
  elEst.classList.remove('oculto');
  if(est.corriendo){
    elEst.textContent = 'Transcribiendo… (PID '+est.pid+'). Puedes revisar '
      +'mientras tanto; la página se actualiza sola y los bloques irán '
      +'apareciendo.';
    elEst.className = 'estado-lan corriendo';
    $('#btnDetener').classList.remove('oculto');
    $('#btnTranscribir').classList.add('oculto');
  } else {
    elEst.textContent = (est.codigo === 0 || est.codigo === null)
      ? 'Transcripción finalizada.'
      : 'La transcripción terminó (código '+est.codigo+'). Revisa el detalle abajo.';
    elEst.className = 'estado-lan fin';
    $('#btnDetener').classList.add('oculto');
    $('#btnTranscribir').classList.remove('oculto');
  }
}

async function vigilarTranscripcion(){
  clearInterval(estado.timerLan);
  const tic = async () => {
    const est = await api('/api/transcripcion_estado');
    pintarEstadoLan(est);
    if(est.corriendo){
      // Mantén la lista de sesiones al día (la sesión en curso puede ser
      // nueva) y, si aún no hay ninguna elegida, engánchate a la más
      // reciente para seguir la transcripción en vivo en el panel.
      await refrescarSelectorSesiones();
      if(!estado.sesion && estado.sesiones.length){
        const viva = estado.sesiones[0].id;   // la más reciente (id DESC)
        ComboSesion.setValor(viva, false);
        await cargarSesion(viva);
      }
    } else {
      clearInterval(estado.timerLan);
      await refrescarSelectorSesiones();
    }
  };
  await tic();
  estado.timerLan = setInterval(tic, 2500);
}

async function cargarLanzador(){
  let info;
  try { info = await api('/api/contextos'); }
  catch(e){ info = {disponible:false, comisiones:[]}; }
  const selCom = $('#lanComisiones'), nota = $('#notaLan');
  selCom.innerHTML = (info.comisiones||[]).map(c =>
    '<option value="'+esc(c)+'">'+esc(c)+'</option>').join('');

  const avisos = [];
  if(!info.tiene_transcriptor)
    avisos.push('No encuentro el script de transcripción en la carpeta; '
      +'el botón no podrá arrancar hasta que esté junto a revisar.py.');
  if(!info.disponible)
    avisos.push('No hay contextos.json: puedes transcribir el Pleno, pero el '
      +'tipo Comisión no tendrá lista hasta que crees ese archivo.');
  else if(!(info.comisiones||[]).length)
    avisos.push('contextos.json no tiene comisiones todavía.');
  nota.textContent = avisos.join(' ');

  const sincronizarTipo = () => {
    const esCom = $('#lanTipo').value === 'comision';
    $('#campoComisiones').classList.toggle('oculto', !esCom);
    $('#campoFecha').classList.toggle('oculto', esCom);
  };
  $('#lanTipo').onchange = sincronizarTipo;
  sincronizarTipo();

  let eventosCargados = false;
  const cargarEventos = async () => {
    const sel = $('#evEvento');
    sel.innerHTML = '<option>Cargando…</option>';
    const eventos = await api('/api/eventos_parlamentarios?tipo='
      + $('#evTipo').value);
    if(eventos.error){
      sel.innerHTML = '<option value="">(error al consultar)</option>';
      return avisar(eventos.error);
    }
    sel.innerHTML = eventos.map(e => {
      const fecha = (e.fecha || '').slice(0, 10);
      const marca = e.liga ? ' 🔗' : '';
      return '<option value="' + esc(e.id) + '">'
        + esc(fecha + ' — ' + (e.descripcion || '').slice(0, 70) + marca)
        + '</option>';
    }).join('') || '<option value="">(sin eventos)</option>';
  };
  $('#evTipo').onchange = cargarEventos;

  const sincronizarOrigen = () => {
    const esEvento = $('#lanOrigen').value === 'evento';
    $('#camposUrl').classList.toggle('oculto', esEvento);
    $('#camposEvento').classList.toggle('oculto', !esEvento);
    if(esEvento && !eventosCargados){ eventosCargados = true; cargarEventos(); }
  };
  $('#lanOrigen').onchange = sincronizarOrigen;
  sincronizarOrigen();

  $('#btnTranscribir').onclick = async () => {
    $('#btnTranscribir').disabled = true;
    if($('#lanOrigen').value === 'evento'){
      const eventoId = $('#evEvento').value;
      if(!eventoId){
        $('#btnTranscribir').disabled = false;
        return avisar('Elige un evento de la lista.');
      }
      const r = await api('/api/transcripciones_evento', {
        evento_id: eventoId, tipo: $('#evTipo').value,
        modelo: $('#lanModelo').value});
      $('#btnTranscribir').disabled = false;
      if(r.error) return avisar(typeof r.error === 'string'
        ? r.error : JSON.stringify(r.error));
      let msg = 'Trabajo creado (' + r.fuente + ').';
      if(r.fuente === 'srt') msg += ' Esperando audio en el puerto ' + r.puerto + '.';
      if((r.participantes_no_encontrados || []).length)
        msg += ' Sin huella de voz: ' + r.participantes_no_encontrados.join(', ') + '.';
      avisar(msg);
      setTimeout(refrescarSelectorSesiones, 4000);
      return;
    }
    const tipo = $('#lanTipo').value;
    const cuerpo = {
      url: $('#lanUrl').value.trim(),
      tipo: tipo,
      modelo: $('#lanModelo').value,
      fecha: (tipo === 'pleno' ? $('#lanFecha').value.trim() : ''),
      comisiones: (tipo === 'comision'
        ? Array.from($('#lanComisiones').selectedOptions).map(o => o.value)
        : [])
    };
    if(!cuerpo.url.startsWith('http')){
      $('#btnTranscribir').disabled = false;
      return avisar('Pega una URL de video válida.');
    }
    if(tipo === 'comision' && !cuerpo.comisiones.length){
      $('#btnTranscribir').disabled = false;
      return avisar('Elige al menos una comisión.');
    }
    const r = await api('/api/transcribir', cuerpo);
    $('#btnTranscribir').disabled = false;
    if(r.error) return avisar(r.error);
    avisar('Transcripción iniciada.');
    vigilarTranscripcion();
  };

  $('#btnDetener').onclick = async () => {
    $('#btnDetener').disabled = true;
    const r = await api('/api/detener', {});
    $('#btnDetener').disabled = false;
    if(r.error) return avisar(r.error);
    avisar(r.detenida ? 'Transcripción detenida.' : 'No había nada corriendo.');
    vigilarTranscripcion();
  };

  // Si al abrir la página ya había una transcripción en curso, engancharse.
  const est = await api('/api/transcripcion_estado');
  if(est.corriendo){ $('#lanzador').open = true; vigilarTranscripcion(); }
}

async function cargarTrabajosApi(){
  const cont = $('#listaTrabajosApi');
  cont.innerHTML = '<p class="nota-lan">Cargando…</p>';
  const trabajos = await api('/api/trabajos_api');
  if(trabajos.error){
    cont.innerHTML = '<p class="nota-lan">' + esc(trabajos.error) + '</p>';
    return;
  }
  if(!trabajos.length){
    cont.innerHTML = '<p class="nota-lan">No hay trabajos creados por la API todavía.</p>';
    return;
  }
  cont.innerHTML = '<div class="tabla-scroll"><table class="tabla-lista">'
    + '<thead><tr><th>URL / etiqueta</th><th>Fuente</th><th>Estado</th>'
    + '<th>Sesión</th><th></th></tr></thead><tbody>'
    + trabajos.map(t => {
      const activo = t.estado === 'ejecutando' || t.estado === 'deteniendo';
      const detalle = t.fuente === 'srt' ? ('SRT, puerto ' + t.puerto) : 'YouTube';
      return '<tr><td>' + esc((t.url || '').slice(0, 70)) + '</td>'
        + '<td>' + esc(detalle) + '</td>'
        + '<td>' + esc(t.estado) + '</td>'
        + '<td>' + (t.sesion_id ? '#' + t.sesion_id : '—') + '</td>'
        + '<td>' + (activo
           ? '<button class="accion" data-id="' + esc(t.id) + '">Detener</button>'
           : '<span class="nota-lan">—</span>')
        + '</td></tr>';
    }).join('')
    + '</tbody></table></div>';
  cont.querySelectorAll('button[data-id]').forEach(b => {
    b.onclick = async () => {
      b.disabled = true;
      const r = await api('/api/trabajos_api/detener', {id: b.dataset.id});
      if(r.error) avisar(typeof r.error === 'string' ? r.error : JSON.stringify(r.error));
      else avisar('Trabajo detenido.');
      cargarTrabajosApi();
    };
  });
}
$('#btnRefrescarTrabajosApi').onclick = cargarTrabajosApi;
cargarTrabajosApi();

async function cargarUsuarios(){
  const r = await api('/api/usuarios');
  if(r.error){
    // No es admin (o la API no respondió): el panel se queda oculto,
    // no tiene caso mostrarlo a quien no lo puede usar.
    $('#panelUsuarios').classList.add('oculto');
    return;
  }
  $('#panelUsuarios').classList.remove('oculto');
  const cont = $('#listaUsuarios');
  cont.innerHTML = !r.length
    ? '<p class="nota-lan">Aún no hay usuarios.</p>'
    : '<div class="tabla-scroll"><table class="tabla-lista">'
      + '<thead><tr><th>Email</th><th>Rol</th></tr></thead><tbody>'
      + r.map(u => '<tr><td>' + esc(u.email) + '</td><td>'
        + (u.es_admin ? 'administrador' : 'operador') + '</td></tr>').join('')
      + '</tbody></table></div>';
}
$('#btnRefrescarUsuarios').onclick = cargarUsuarios;
$('#btnCrearUsuario').onclick = async () => {
  const email = $('#nuEmail').value.trim();
  if(!email) return avisar('Escribe un email.');
  const boton = $('#btnCrearUsuario');
  boton.disabled = true;
  try{
    const r = await api('/api/usuarios',
      {email, es_admin: $('#nuEsAdmin').checked});
    if(r.error){
      avisar(typeof r.error === 'string' ? r.error : JSON.stringify(r.error));
      return;
    }
    $('#nuevaCredencial').classList.remove('oculto');
    $('#nuevaCredencial').innerHTML =
      '<strong>Cuenta creada.</strong> Copia esta contraseña ahora — no '
      + 'se vuelve a mostrar (solo puedes crear una nueva si se pierde):<br>'
      + 'Usuario: <code>' + esc(r.email) + '</code><br>'
      + 'Contraseña: <code style="font-size:15px">' + esc(r.password) + '</code>';
    $('#nuEmail').value = '';
    $('#nuEsAdmin').checked = false;
    cargarUsuarios();
  } finally {
    boton.disabled = false;
  }
};
cargarUsuarios();

iniciar();
cargarLanzador();
</script>
</body>
</html>
"""


PAGINA_ESTENO = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Módulo estenográfico</title>
<style>
:root{
  --papel:#FAF8F2; --panel:#FFFFFF; --tinta:#20221C; --tinta-suave:#494A41;
  --tenue:#63625A; --verde:#1E5A38; --verde-hondo:#154029; --verde-suave:#E7EFE9;
  --linea:#E4E1D6; --linea-fuerte:#D2CFC2; --ambar:#9C5E18; --ambar-suave:#FAF0DE;
  --rojo:#9B2C2C; --rojo-suave:#F7E7E7; --azul:#1B4C7A; --azul-suave:#E5EEF6;
  --r:9px; --r-sm:6px; --sombra:0 1px 2px rgba(30,40,25,.05),0 8px 24px rgba(30,40,25,.06);
}
*{box-sizing:border-box}
html,body{margin:0;background:var(--papel);color:var(--tinta);
  font:15px/1.55 -apple-system,"Segoe UI",Roboto,"Helvetica Neue",Arial,sans-serif}
:focus-visible{outline:2px solid var(--ambar);outline-offset:2px}
a{color:var(--verde)}
header{display:flex;flex-wrap:wrap;align-items:center;gap:8px 18px;
  padding:13px 24px;background:var(--panel);border-bottom:1px solid var(--linea-fuerte)}
header h1{margin:0;font:700 20px/1.05 Georgia,"Times New Roman",serif;color:var(--verde)}
header .sub{color:var(--tenue);font-size:11px;text-transform:uppercase;letter-spacing:.14em;font-weight:600}
header .der{margin-left:auto;display:flex;gap:10px;align-items:center;flex-wrap:wrap}
.rol{font-size:12px;font-weight:700;padding:3px 10px;border-radius:999px}
.rol.admin{background:var(--azul-suave);color:var(--azul)}
.rol.corr{background:var(--verde-suave);color:var(--verde)}
select,input[type=text],input[type=number]{font:inherit;font-size:14px;padding:7px 10px;
  border:1px solid var(--linea-fuerte);border-radius:var(--r-sm);background:var(--panel);color:var(--tinta)}
main{max-width:1080px;margin:0 auto;padding:20px 24px 80px}
.tarjeta{background:var(--panel);border:1px solid var(--linea);border-radius:var(--r);
  padding:18px 20px;margin:0 0 18px;box-shadow:var(--sombra)}
.tarjeta h2{margin:0 0 4px;font:700 15px/1.2 -apple-system,"Segoe UI",Roboto,sans-serif;color:var(--verde)}
.tarjeta p.ayuda{margin:0 0 14px;color:var(--tenue);font-size:13px}
.fila{display:flex;flex-wrap:wrap;gap:14px;align-items:flex-end}
.campo{display:flex;flex-direction:column;gap:5px;font-size:12px;font-weight:600;color:var(--tenue)}
.campo input,.campo select,.campo textarea{font-weight:400}
.campo textarea{min-width:260px;min-height:74px;font:inherit;padding:7px 10px;
  border:1px solid var(--linea-fuerte);border-radius:var(--r-sm)}
.btn{font:600 14px/1.2 -apple-system,"Segoe UI",Roboto,sans-serif;padding:9px 15px;border-radius:var(--r-sm);
  cursor:pointer;border:1px solid var(--linea-fuerte);background:var(--panel);color:var(--tinta);white-space:nowrap}
.btn:hover{border-color:var(--verde);background:var(--verde-suave);color:var(--verde-hondo)}
.btn.pri{background:var(--verde);border-color:var(--verde);color:#fff}
.btn.pri:hover{background:var(--verde-hondo)}
.btn.peligro{background:var(--rojo-suave);border-color:var(--rojo);color:var(--rojo)}
.btn.chico{padding:5px 10px;font-size:13px}
.btn:disabled{opacity:.5;cursor:default}
.barra-prog{height:9px;border-radius:999px;background:var(--linea);overflow:hidden;margin:8px 0 2px}
.barra-prog>span{display:block;height:100%;background:var(--verde);width:0}
table.plan{width:100%;border-collapse:collapse;font-size:13.5px}
table.plan th,table.plan td{padding:8px 10px;text-align:left;border-bottom:1px solid var(--linea)}
table.plan th{font-size:11px;text-transform:uppercase;letter-spacing:.08em;color:var(--tenue)}
.badge{display:inline-block;font-size:11.5px;font-weight:700;padding:2px 9px;border-radius:999px}
.badge.pend{background:#EFEDE6;color:var(--tenue)}
.badge.edit{background:var(--ambar-suave);color:var(--ambar)}
.badge.term{background:var(--verde-suave);color:var(--verde)}
.badge.mio{background:var(--azul-suave);color:var(--azul)}
.bloques{display:grid;grid-template-columns:repeat(auto-fill,minmax(220px,1fr));gap:12px}
.bloque{border:1px solid var(--linea);border-radius:var(--r);padding:13px 15px;background:var(--panel)}
.bloque.mio{border-color:var(--azul);box-shadow:0 0 0 1px var(--azul) inset}
.bloque.en-curso{opacity:.6;filter:saturate(.5)}
.bloque .rango{font:700 14px/1.2 -apple-system,"Segoe UI",Roboto,sans-serif}
.bloque .met{font-size:12px;color:var(--tenue);margin:4px 0 10px}
.editor-cab{position:sticky;top:0;z-index:10;background:var(--panel);
  border:1px solid var(--linea-fuerte);border-radius:var(--r);padding:12px 16px;
  margin:0 0 16px;box-shadow:var(--sombra)}
.audio-controles{display:flex;flex-wrap:wrap;align-items:center;gap:8px}
.audio-controles .reloj{font-variant-numeric:tabular-nums;font-weight:700;font-size:14px;min-width:118px}
.audio-controles .btn{padding:7px 11px}
.pista{flex:1 1 220px;min-width:180px}
.seg{display:grid;grid-template-columns:92px 1fr;gap:12px;padding:10px 0;border-bottom:1px solid var(--linea)}
.seg .t{font-size:12px;color:var(--tenue);font-variant-numeric:tabular-nums;background:none;border:0;
  border-bottom:1px dashed var(--verde);color:var(--verde);cursor:pointer;padding:0;text-align:left;height:fit-content}
.seg .t:hover{color:var(--ambar);border-color:var(--ambar)}
.seg .campos{display:flex;flex-direction:column;gap:6px;min-width:0}
.seg input.orador{font-weight:700;color:var(--verde);text-transform:uppercase;font-size:12.5px;
  padding:5px 8px;border:1px solid var(--linea);border-radius:var(--r-sm);max-width:340px;flex:1 1 220px}
.orador-fila{display:flex;gap:8px;align-items:center;flex-wrap:wrap}
.seg select.pick{font:inherit;font-size:12.5px;padding:5px 8px;border:1px solid var(--linea-fuerte);
  border-radius:var(--r-sm);background:var(--panel);color:var(--tenue);max-width:230px}
.seg textarea.txt{font:16px/1.6 Georgia,"Times New Roman",serif;padding:8px 10px;border:1px solid var(--linea);
  border-radius:var(--r-sm);resize:vertical;min-height:52px;width:100%}
.seg textarea.txt:focus,.seg input.orador:focus{border-color:var(--verde);outline:none}
/* Editor enriquecido (reemplaza al textarea) */
.seg .editor-rico{font:16px/1.6 Georgia,"Times New Roman",serif;padding:8px 10px;
  border:1px solid var(--linea);border-radius:var(--r-sm);min-height:52px;width:100%;
  outline:none;text-align:justify;background:#fff}
.seg .editor-rico:focus{border-color:var(--verde);outline:2px solid var(--verde-suave)}
/* Barra de formato */
.barra-fmt{display:flex;flex-wrap:wrap;gap:4px;margin-bottom:5px}
.barra-fmt button{font:700 12px/1 -apple-system,"Segoe UI",Roboto,sans-serif;
  padding:4px 8px;border:1px solid var(--linea-fuerte);border-radius:4px;
  background:var(--panel);color:var(--tinta);cursor:pointer;min-width:28px}
.barra-fmt button:hover{background:var(--verde-suave);border-color:var(--verde)}
.barra-fmt button.activo{background:var(--verde);color:#fff;border-color:var(--verde-hondo)}
.barra-fmt select{font:inherit;font-size:12px;padding:3px 6px;border:1px solid var(--linea-fuerte);border-radius:4px;background:var(--panel)}
.barra-fmt .sep{width:1px;background:var(--linea-fuerte);margin:2px 2px;align-self:stretch}
.aviso{position:fixed;left:50%;bottom:22px;transform:translateX(-50%) translateY(140%);
  background:var(--tinta);color:#fff;padding:11px 18px;border-radius:999px;font-size:14px;
  box-shadow:0 8px 26px rgba(0,0,0,.28);transition:transform .28s;z-index:50;max-width:90vw}
.aviso.ver{transform:translateX(-50%) translateY(0)}
.pie-editor{display:flex;flex-wrap:wrap;gap:10px;margin-top:16px;padding-top:14px;border-top:1px solid var(--linea)}
.atajos{font-size:12px;color:var(--tenue);margin-left:auto;align-self:center}
.enlace-corr{display:flex;gap:8px;align-items:center;font-size:13px;margin:4px 0}
.enlace-corr code{background:var(--verde-suave);padding:2px 7px;border-radius:var(--r-sm);color:var(--verde-hondo);font-size:12px}
.cargando{color:var(--tenue);padding:30px;text-align:center}
.pill-audio{font-size:12px;font-weight:700;padding:2px 9px;border-radius:999px}
.pill-audio.si{background:var(--verde-suave);color:var(--verde)}
.pill-audio.no{background:var(--rojo-suave);color:var(--rojo)}
.pie-seg{display:flex;gap:8px;margin-top:8px}
.divisor{margin-top:10px;padding:12px;border:1px solid var(--linea);border-radius:8px;background:var(--papel)}
.divisor .ayuda{margin:0 0 8px;color:var(--tenue);font-size:13px}
.divisor .frases{font:15px/1.8 Georgia,"Times New Roman",serif}
.divisor .frase{cursor:pointer;padding:1px 3px;border-radius:4px}
.divisor .frase:hover{background:var(--verde-suave)}
.divisor .frase.p2{background:#F4E7CE}
.split-oradores{display:flex;flex-wrap:wrap;gap:10px;align-items:center;margin-top:12px}
.split-oradores label{display:flex;gap:6px;align-items:center;font-size:13px;color:var(--tenue)}
.split-oradores input{font:inherit;padding:6px 9px;border:1px solid var(--linea-fuerte);border-radius:6px}
@media(max-width:640px){.seg{grid-template-columns:1fr}main{padding:16px}}
</style>
</head>
<body>
<header>
  <div>
    <h1>Módulo estenográfico</h1>
    <div class="sub">Corrección colaborativa por tramos</div>
  </div>
  <div class="der">
    <span id="rol" class="rol"></span>
    <a id="lnkOtraSesion" href="/esteno" class="btn chico" style="display:none">← Elegir otra sesión</a>
    <a id="lnkPrincipal" href="/" class="btn chico">← Revisión principal</a>
    <a href="/logout" class="btn chico">Cerrar sesión</a>
  </div>
</header>
<main id="app"><p class="cargando">Cargando…</p></main>
<div id="aviso" class="aviso"></div>
<script>
const $ = s => document.querySelector(s);
const app = $('#app');
const params = new URLSearchParams(location.search);
const S = {
  sesion: parseInt(params.get('sesion')||'0',10) || 0,
  // corrector/esAdmin ya NO salen de la URL: se resuelven en iniciar()
  // contra /api/yo, con la cuenta con la que de verdad iniciaste sesión.
  corrector: null, esAdmin: false, usuarios: [], filtroCorrector: '',
  miSlot: null, config:null, bloques:[], total:0, audio:false, sesionInfo:null,
  abierto:null, segmentos:[], timerEstado:null, timerLatido:null, au:null, nombres:[],
};

function avisar(t){ const a=$('#aviso'); a.textContent=t; a.classList.add('ver');
  clearTimeout(avisar._t); avisar._t=setTimeout(()=>a.classList.remove('ver'),4200); }
function esc(t){ const d=document.createElement('div'); d.textContent=(t==null?'':t); return d.innerHTML; }
function hms(seg){ seg=Math.max(0,Math.floor(seg||0));
  const h=Math.floor(seg/3600), m=Math.floor((seg%3600)/60), s=seg%60;
  const p=n=>String(n).padStart(2,'0');
  return (h>0? p(h)+':':'')+p(m)+':'+p(s); }
async function api(path, body){
  const opt = body ? {method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)} : {};
  const r = await fetch(path, opt);
  let j={}; try{ j=await r.json(); }catch(e){}
  if(!r.ok){ throw new Error(j.error || ('Error '+r.status)); }
  return j;
}
function nombreSlot(slot){
  const n = S.config && S.config.nombres || [];
  return (n[slot] && n[slot].trim()) ? n[slot] : ('Corrector '+(slot+1));
}

async function cargarEstado(){
  const d = await api('/api/esteno/estado?sesion='+S.sesion);
  S.config=d.config; S.bloques=d.bloques; S.total=d.total_seg;
  S.audio=d.audio; S.sesionInfo=d.sesion; S.progreso=d.progreso;
  // Resolver mi slot por nombre (si el admin lo registró en la lista)
  if(!S.esAdmin && S.miSlot===null && S.config && S.config.nombres){
    const i = S.config.nombres.findIndex(n => (n||'').trim().toLowerCase()===S.corrector.toLowerCase());
    if(i>=0) S.miSlot=i;
  }
}

/* ---------------- ADMIN ---------------- */
async function vistaAdmin(){
  $('#rol').className='rol admin'; $('#rol').textContent='Administrador';
  let sesiones=[]; try{ sesiones=await api('/api/sesiones'); }catch(e){}
  try{ S.usuarios = await api('/api/usuarios'); }catch(e){ S.usuarios = []; }
  const opts = '<option value="0">— Elige una sesión —</option>'+
    sesiones.map(s=>'<option value="'+s.id+'"'+(s.id===S.sesion?' selected':'')+'>#'+s.id+' — '+esc((s.titulo||'').slice(0,60))+'</option>').join('');
  let html = '<div class="tarjeta"><h2>Sesión</h2>'+
    '<div class="fila"><label class="campo">Sesión a preparar<select id="selSes">'+opts+'</select></label></div></div>';
  if(S.sesion){
    const cfg=S.config||{};
    html += '<div class="tarjeta"><h2>Plan de turnos</h2>'+
      '<p class="ayuda">Define cuántos correctores trabajarán y de cuántos minutos será cada bloque. '+
      'La duración es un <b>mínimo aproximado</b>: si al cumplirse el tiempo hay una intervención en curso, '+
      'el bloque se extiende hasta que ese orador termine (no se parten intervenciones). '+
      'Los bloques se reparten en rotación (Corrector 1: 1º, luego cada N; y así). '+
      'Asigna cada turno a una cuenta real — si a alguien le falta, créala primero en '+
      '"Usuarios" desde la pantalla principal.</p>'+
      '<div class="fila">'+
        '<label class="campo">Nº de correctores<input id="cNum" type="number" min="1" max="50" value="'+(cfg.num_correctores||2)+'" style="width:110px"></label>'+
        '<label class="campo">Minutos por bloque<input id="cMin" type="number" min="1" max="120" value="'+((cfg.bloque_seg||600)/60)+'" style="width:130px"></label>'+
      '</div>'+
      '<div id="filasCorrectores" class="fila" style="margin-top:6px"></div>'+
      '<button class="btn pri" id="bGen" style="margin-top:10px">Generar / actualizar plan</button></div>';

    // Audio
    html += '<div class="tarjeta"><h2>Audio de la sesión</h2>'+
      '<p class="ayuda">Si la transcripción se lanzó con "conservar audio" activo, el audio '+
      'se va enlazando solo conforme llega — no hace falta hacer nada aquí. Esto es solo '+
      'para cuando el audio vive en otro lado y quieres apuntarlo a mano.</p>'+
      '<div class="fila"><span class="pill-audio '+(S.audio?'si':'no')+'">'+(S.audio?'Audio detectado':'Sin audio')+'</span>'+
      '<label class="campo" style="flex:1 1 320px">Ruta del audio (opcional)<input id="aRuta" type="text" placeholder="C:\\\\ruta\\\\al\\\\audio.mp3 o /home/…/audio.mp3"></label>'+
      '<button class="btn" id="bAudio">Vincular audio</button></div></div>';

    // Progreso + plan (en su propio contenedor, para refrescar SOLO esto)
    if(S.config){
      html += '<div class="tarjeta"><h2>Avance</h2>'+
        '<div id="panelPlan">'+panelPlanHTML()+'</div></div>';
      html += '<div class="tarjeta"><h2>Enlace para compartir</h2>'+
        '<p class="ayuda">El mismo enlace sirve para todos — cada quien entra con su propia '+
        'cuenta y ve sus turnos automáticamente.</p>'+
        '<div class="enlace-corr"><code>'+esc(location.origin+'/esteno?sesion='+S.sesion)+'</code>'+
        '<button class="btn chico" data-url="'+esc(location.origin+'/esteno?sesion='+S.sesion)+'">Copiar</button></div></div>';
    }
  }
  app.innerHTML=html;

  $('#selSes').onchange = e => { location.search='?sesion='+e.target.value; };
  pintarFilasCorrectores();
  if($('#cNum')) $('#cNum').oninput = pintarFilasCorrectores;
  if($('#bGen')) $('#bGen').onclick = async ()=>{
    const nombres = Array.from(document.querySelectorAll('#filasCorrectores select[data-slot]'))
      .map(s=>s.value);
    try{
      const r=await api('/api/esteno/configurar',{sesion_id:S.sesion,
        num_correctores:parseInt($('#cNum').value,10)||1,
        bloque_min:parseFloat($('#cMin').value)||10, nombres});
      avisar('Plan generado: '+r.bloques+' bloques'+(r.reinicio?' (estados reiniciados)':'.'));
      await cargarEstado(); vistaAdmin();
    }catch(e){ avisar(e.message); }
  };
  if($('#bAudio')) $('#bAudio').onclick = async ()=>{
    try{ const r=await api('/api/esteno/audio_ruta',{sesion_id:S.sesion,ruta:$('#aRuta').value.trim()});
      avisar(r.existe?'Audio vinculado correctamente.':'Ruta guardada, pero el archivo no existe todavía.');
      await cargarEstado(); vistaAdmin();
    }catch(e){ avisar(e.message); }
  };
  app.querySelectorAll('button[data-url]').forEach(b=> b.onclick=()=>{
    navigator.clipboard.writeText(b.dataset.url); avisar('Enlace copiado.'); });
  cablearReasignar();
}
// Una fila por turno, cada una con un <select> de cuentas reales (no texto
// libre) — así el plan queda atado a usuarios que de verdad pueden entrar.
function pintarFilasCorrectores(){
  const cont = $('#filasCorrectores'); if(!cont) return;
  const num = parseInt($('#cNum').value, 10) || 1;
  const actuales = (S.config && S.config.nombres) || [];
  const opciones = '<option value="">— sin asignar —</option>' +
    (S.usuarios||[]).map(u=>'<option value="'+esc(u.email)+'">'+esc(u.email)
      +(u.es_admin?' (admin)':'')+'</option>').join('');
  let h='';
  for(let i=0;i<num;i++){
    h += '<label class="campo">Turno '+(i+1)+'<select data-slot="'+i+'">'+opciones+'</select></label>';
  }
  cont.innerHTML = h;
  cont.querySelectorAll('select[data-slot]').forEach(sel=>{
    const i = parseInt(sel.dataset.slot, 10);
    if(actuales[i]) sel.value = actuales[i];
  });
}
function panelPlanHTML(){
  const pr=S.progreso||{terminados:0,total:0};
  const pct = pr.total? Math.round(100*pr.terminados/pr.total):0;
  // Quién ha tocado esta sesión: la lista del plan + cualquiera que haya
  // ayudado en un bloque de otro turno (no solo los nombres configurados).
  const personas = new Set();
  (S.config.nombres||[]).forEach(n=> n && personas.add(n));
  S.bloques.forEach(b=>{
    if(b.tomado_por) personas.add(b.tomado_por);
    if(b.terminado_por) personas.add(b.terminado_por);
  });
  const opcsFiltro = '<option value="">— todos —</option>' +
    Array.from(personas).sort().map(n=> '<option value="'+esc(n)+'"'
      + (S.filtroCorrector===n?' selected':'')+'>'+esc(n)+'</option>').join('');
  const urlWord = S.filtroCorrector
    ? '/api/exportar_word?sesion='+S.sesion+'&corrector='+encodeURIComponent(S.filtroCorrector)
    : '';
  const urlWordAutores = '/api/exportar_word?sesion='+S.sesion+'&marcar_autores=1';
  return '<div class="barra-prog"><span style="width:'+pct+'%"></span></div>'+
    '<p class="ayuda">'+pr.terminados+' de '+pr.total+' bloques terminados ('+pct+'%). '+
    'Duración de la sesión: '+hms(S.total)+'. Se actualiza solo (sin borrar tu formulario).</p>'+
    '<div class="fila" style="align-items:flex-end">'+
    '<label class="campo" style="max-width:320px">Ver bloques de'+
    '<select id="filtroCorrector">'+opcsFiltro+'</select></label>'+
    (urlWord ? '<a class="btn chico" href="'+esc(urlWord)+'" target="_blank" '
      + 'title="Word con solo los tramos que tomó o terminó esta persona">'
      + '📄 Descargar Word de esta persona</a>' : '')+
    '<a class="btn chico" href="'+esc(urlWordAutores)+'" target="_blank" '
      + 'title="Word completo (todos los oradores), con una nota antes de '
      + 'cada tramo diciendo quién lo corrigió">'
      + '📄 Descargar Word con autores marcados</a>'+
    '</div>'+
    tablaPlan();
}
function cablearReasignar(){
  app.querySelectorAll('select[data-reasignar]').forEach(sel=> sel.onchange=async ()=>{
    try{ await api('/api/esteno/reasignar',{sesion_id:S.sesion,indice:parseInt(sel.dataset.reasignar,10),slot:parseInt(sel.value,10)});
      await cargarEstado(); refrescarPlan(); }catch(err){ avisar(err.message); }
  });
  const filtro = $('#filtroCorrector');
  if(filtro) filtro.onchange = () => {
    // Redibuja de inmediato: refrescarPlan() se frena a propósito cuando
    // el foco sigue en este mismo select (para que el auto-refresco cada
    // 4 s no te lo cierre a medio uso) — pero aquí SÍ acabas de elegir,
    // así que no hay que esperar al siguiente ciclo.
    S.filtroCorrector = filtro.value;
    const cont = document.getElementById('panelPlan');
    if(cont){ cont.innerHTML = panelPlanHTML(); cablearReasignar(); }
  };
}
// Refresca SOLO el panel de avance/plan; nunca vuelve a dibujar el formulario.
function refrescarPlan(){
  const cont=document.getElementById('panelPlan');
  if(!cont) return;
  const a=document.activeElement;
  if(a && a.matches && (a.matches('select[data-reasignar]') || a.id==='filtroCorrector')) return; // no interrumpir
  cont.innerHTML = panelPlanHTML();
  cablearReasignar();
}
function tablaPlan(){
  const N=S.config.num_correctores;
  const filtro = S.filtroCorrector;
  const bloques = filtro
    ? S.bloques.filter(b=> b.tomado_por===filtro || b.terminado_por===filtro)
    : S.bloques;
  if(filtro && !bloques.length){
    return '<p class="ayuda">'+esc(filtro)+' todavía no ha tocado ningún bloque de esta sesión.</p>';
  }
  let r='<table class="plan"><thead><tr><th>#</th><th>Tramo</th><th>Asignado</th><th>Estado</th><th>Segmentos</th></tr></thead><tbody>';
  bloques.forEach(b=>{
    let opts=''; for(let i=0;i<N;i++) opts+='<option value="'+i+'"'+(i===b.slot?' selected':'')+'>'+esc(nombreSlot(i))+'</option>';
    let est='<span class="badge pend">pendiente</span>';
    if(b.estado==='editando') est='<span class="badge edit">editando · '+esc(b.tomado_por||'')+(b.vencido?' (inactivo)':'')+'</span>';
    if(b.estado==='terminado') est='<span class="badge term">terminado · '+esc(b.terminado_por||'')+'</span>';
    r+='<tr><td>'+(b.indice+1)+'</td><td>'+hms(b.inicio_seg)+' – '+hms(b.fin_seg)+'</td>'+
       '<td><select data-reasignar="'+b.indice+'">'+opts+'</select></td><td>'+est+'</td><td>'+b.segmentos+'</td></tr>';
  });
  return r+'</tbody></table>';
}

/* ---------------- CORRECTOR ---------------- */
function vistaCorrector(){
  $('#rol').className='rol corr'; $('#rol').textContent='Corrector: '+(S.corrector||'—');
  if(S.abierto!==null){ vistaEditor(); return; }
  if(!S.config){ app.innerHTML='<div class="tarjeta"><h2>Aún no hay plan</h2><p class="ayuda">El administrador todavía no ha configurado los turnos de esta sesión. Vuelve a intentar en un momento.</p></div>'; return; }
  let selSlot='';
  if(S.miSlot===null){
    let o=''; for(let i=0;i<S.config.num_correctores;i++) o+='<option value="'+i+'">'+esc(nombreSlot(i))+'</option>';
    selSlot='<div class="tarjeta"><h2>¿Qué turno tienes?</h2><p class="ayuda">El administrador todavía no te asignó un turno con tu cuenta ('+esc(S.corrector||'')+'); elige cuál te toca para ver primero tus bloques (igual puedes ayudar en cualquier otro).</p>'+
      '<div class="fila"><label class="campo">Mi turno<select id="miSlot">'+o+'</select></label><button class="btn pri" id="bSlot">Confirmar</button></div></div>';
  }
  const mios = S.miSlot===null? [] : S.bloques.filter(b=>b.slot===S.miSlot);
  const otros = S.miSlot===null? S.bloques : S.bloques.filter(b=>b.slot!==S.miSlot);
  if(!S.bloques.length){
    app.innerHTML = selSlot + '<div class="tarjeta"><h2>El plan aún no tiene bloques</h2>'+
      '<p class="ayuda">La sesión todavía no registra duración (quizá acaba de empezar o aún no hay texto transcrito). '+
      'Los bloques aparecerán solos en cuanto la transcripción avance; esta página se actualiza cada pocos segundos.</p></div>';
    if($('#bSlot')) $('#bSlot').onclick=()=>{ S.miSlot=parseInt($('#miSlot').value,10); vistaCorrector(); };
    return;
  }
  let html=selSlot;
  if(S.miSlot!==null){
    html+='<div class="tarjeta"><h2>Mis bloques'+(S.audio?'':' · <span class="pill-audio no">sin audio</span>')+'</h2>'+
      '<p class="ayuda">Abre un bloque para corregirlo escuchando su tramo. Al terminar, guarda para afectar el registro final.</p>'+
      '<div class="bloques">'+mios.map(b=>tarjBloque(b,true)).join('')+'</div></div>';
  }
  html+='<div class="tarjeta"><h2>Otros bloques</h2><p class="ayuda">Puedes ayudar con un bloque libre de otro turno.</p>'+
    '<div class="bloques">'+otros.map(b=>tarjBloque(b,false)).join('')+'</div></div>';
  app.innerHTML=html;
  if($('#bSlot')) $('#bSlot').onclick=()=>{ S.miSlot=parseInt($('#miSlot').value,10); vistaCorrector(); };
  app.querySelectorAll('button[data-abrir]').forEach(b=> b.onclick=()=>abrirBloque(parseInt(b.dataset.abrir,10)));
}
function tarjBloque(b,mio){
  let est='<span class="badge pend">pendiente</span>', accion='';
  const libre = b.estado!=='editando' || b.vencido || b.tomado_por===S.corrector;
  if(b.estado==='editando') est='<span class="badge '+(b.tomado_por===S.corrector?'mio':'edit')+'">'+(b.tomado_por===S.corrector?'tú lo tienes':'editando · '+esc(b.tomado_por||'')+(b.vencido?' (inactivo)':''))+'</span>';
  if(b.estado==='terminado') est='<span class="badge term">terminado</span>';
  if(!b.listo){
    accion='<span class="badge" style="background:#FEF3C7;color:#92400E;font-size:11px">⏳ grabando…</span>';
  } else if(libre){
    accion='<button class="btn '+(mio?'pri':'')+' chico" data-abrir="'+b.indice+'">'+(b.estado==='terminado'?'Revisar de nuevo':'Abrir y corregir')+'</button>';
  } else {
    accion='<button class="btn chico" disabled>En uso</button>';
  }
  return '<div class="bloque'+(mio?' mio':'')+(b.listo?'':' en-curso')+'">'+
    '<div class="rango">Bloque '+(b.indice+1)+'</div>'+
    '<div class="met">'+hms(b.inicio_seg)+' – '+hms(b.fin_seg)+' · '+b.segmentos+' segmentos<br>'+est+'</div>'+
    accion+'</div>';
}

async function abrirBloque(indice){
  try{
    const r=await api('/api/esteno/tomar',{sesion_id:S.sesion,indice,corrector:S.corrector||'anónimo'});
    S.abierto=r.bloque; S.segmentos=r.segmentos;
    detenerPoll();
    vistaEditor();
    // latido para conservar el bloqueo mientras edita
    S.timerLatido=setInterval(()=>{ api('/api/esteno/latido',{sesion_id:S.sesion,indice,corrector:S.corrector||'anónimo'}).catch(()=>{}); }, 30000);
  }catch(e){ avisar(e.message); await cargarEstado(); vistaCorrector(); }
}

function vistaEditor(){
  const b=S.abierto;
  const turnos = agruparTurnos(S.segmentos);
  const opcNombres = '<option value="">— elegir integrante —</option>' +
    (S.nombres||[]).map(n=>'<option>'+esc(n)+'</option>').join('');
  const datalist = '<datalist id="dlNombres">' +
    (S.nombres||[]).map(n=>'<option value="'+esc(n)+'">').join('') + '</datalist>';
  const barraFmt =
    '<div class="barra-fmt">'+
    '<button data-cmd="bold" title="Negrita (Ctrl+B)"><b>N</b></button>'+
    '<button data-cmd="italic" title="Cursiva (Ctrl+I)"><i>C</i></button>'+
    '<button data-cmd="underline" title="Subrayado (Ctrl+U)"><u>S</u></button>'+
    '<button data-cmd="strikeThrough" title="Tachado"><s>T</s></button>'+
    '<span class="sep"></span>'+
    '<button data-cmd="justifyFull" title="Justificar">Just.</button>'+
    '<button data-cmd="justifyLeft" title="Izquierda">Izq.</button>'+
    '<button data-cmd="justifyCenter" title="Centrar">Cent.</button>'+
    '<button data-cmd="justifyRight" title="Derecha">Der.</button>'+
    '<span class="sep"></span>'+
    '<button data-cmd="insertUnorderedList" title="Lista con viñetas">• Lista</button>'+
    '<button data-cmd="insertOrderedList" title="Lista numerada">1. Lista</button>'+
    '<span class="sep"></span>'+
    '<button data-cmd="removeFormat" title="Quitar formato">Sin fmt</button>'+
    '</div>';
  let segs = turnos.map((t,i) => {
    const contenidoHTML = t.texto || '';
    const btnUnir = i < turnos.length-1
      ? '<button class="btn chico b-unir" title="Pegar este turno con el siguiente">🔗 Unir con el siguiente</button>'
      : '';
    return '<div class="seg" data-ids="'+t.ids.join(',')+'">'+
      '<button class="t" data-t="'+t.inicio_seg+'">'+esc(t.inicio_hms||hms(t.inicio_seg))+' ▸</button>'+
      '<div class="campos">'+
        '<div class="orador-fila">'+
          '<input class="orador" list="dlNombres" value="'+esc(t.orador||'')+'" placeholder="Orador">'+
          '<select class="pick" title="Elegir de la lista de integrantes">'+opcNombres+'</select>'+
        '</div>'+
        barraFmt+
        '<div class="editor-rico" contenteditable="true" spellcheck="true">'+contenidoHTML+'</div>'+
        '<div class="pie-seg">'+btnUnir+
          '<button class="btn chico b-dividir" title="Partir este turno en dos oradores">✂ Dividir</button>'+
        '</div>'+
      '</div></div>';
  }).join('');
  segs = datalist + segs;
  app.innerHTML =
    '<div class="editor-cab">'+
      '<div style="display:flex;flex-wrap:wrap;gap:6px 14px;align-items:center;margin-bottom:10px">'+
        '<button class="btn chico" id="bVolver" title="Suelta este bloque sin guardar y regresa a la lista">← Volver a mis bloques</button>'+
        '<strong>Bloque '+(b.indice+1)+'</strong>'+
        '<span style="color:var(--tenue);font-size:13px">'+hms(b.inicio_seg)+' – '+hms(b.fin_seg)+'</span>'+
        (S.audio?'':'<span class="pill-audio no">sin audio</span>')+
      '</div>'+
      '<div class="audio-controles">'+
        '<button class="btn" id="aRet" title="Atrasar 5 s (F1)">« 5s</button>'+
        '<button class="btn pri" id="aPlay" title="Reproducir / pausar (F2)">▶ Reproducir</button>'+
        '<button class="btn" id="aAde" title="Adelantar 5 s (F3)">5s »</button>'+
        '<button class="btn" id="aIni" title="Ir al inicio del bloque">⟲ inicio</button>'+
        '<span class="reloj" id="aReloj">00:00 / 00:00</span>'+
        '<input class="pista" id="aPista" type="range" min="0" max="1000" value="0">'+
        '<label style="font-size:12px;color:var(--tenue);display:flex;gap:5px;align-items:center"><input type="checkbox" id="aLoop" checked> solo mi tramo</label>'+
        '<label style="font-size:12px;color:var(--tenue);display:flex;gap:5px;align-items:center">vel.'+
          '<select id="aVel"><option>0.75</option><option selected>1</option><option>1.25</option><option>1.5</option></select></label>'+
      '</div>'+
    '</div>'+
    '<div id="listaSeg">'+(segs||'<p class="ayuda">Este bloque no tiene segmentos.</p>')+'</div>'+
    '<div class="pie-editor">'+
      '<button class="btn pri" id="bTerm">✓ Guardar y terminar</button>'+
      '<button class="btn" id="bAvance">Guardar avance</button>'+
      '<button class="btn peligro" id="bCerrar">Cerrar sin guardar</button>'+
      '<span class="atajos">Atajos: F1 atrasar · F2 play/pausa · F3 adelantar</span>'+
    '</div>';
  montarAudio();
  $('#bTerm').onclick=()=>guardar(true);
  $('#bAvance').onclick=()=>guardar(false);
  $('#bCerrar').onclick=cerrar;
  $('#bVolver').onclick=cerrar;
  app.querySelectorAll('.seg .t').forEach(t=> t.onclick=()=>{ if(S.au){ S.au.currentTime=parseFloat(t.dataset.t); S.au.play(); actualizarPlay(); } });
  app.querySelectorAll('.seg .pick').forEach(sel=> sel.onchange=()=>{
    if(!sel.value) return;
    sel.closest('.campos').querySelector('.orador').value=sel.value; sel.value='';
  });
  app.querySelectorAll('.seg .b-unir').forEach(b=> b.onclick=()=>unirConSiguiente(b));
  app.querySelectorAll('.seg .b-dividir').forEach(b=> b.onclick=()=>dividirAqui(b));
  // Barra de formato: execCommand sobre el editor enfocado
  app.querySelectorAll('.barra-fmt button[data-cmd]').forEach(btn=>{
    btn.onmousedown=e=>{
      e.preventDefault();  // no perder el foco del editor
      document.execCommand(btn.dataset.cmd, false, null);
    };
  });
  // Actualiza qué botones aparecen activos según la selección actual
  function actualizarBarraFmt(){
    const cmds=['bold','italic','underline','strikeThrough',
                'justifyFull','justifyLeft','justifyCenter','justifyRight',
                'insertUnorderedList','insertOrderedList'];
    cmds.forEach(cmd=>{
      app.querySelectorAll('[data-cmd="'+cmd+'"]').forEach(b=>{
        b.classList.toggle('activo', document.queryCommandState(cmd));
      });
    });
  }
  app.querySelectorAll('.editor-rico').forEach(ed=>{
    ed.addEventListener('keyup', actualizarBarraFmt);
    ed.addEventListener('mouseup', actualizarBarraFmt);
    ed.addEventListener('focus', actualizarBarraFmt);
  });
}

function montarAudio(){
  const b=S.abierto;
  const au=new Audio(); S.au=au;
  au.preload='metadata';
  if(S.audio) au.src='/api/esteno/audio?sesion='+S.sesion;
  au.playbackRate=parseFloat($('#aVel').value);
  let listo=false;
  au.addEventListener('loadedmetadata',()=>{ listo=true; au.currentTime=b.inicio_seg; pintarReloj(); });
  au.addEventListener('timeupdate',()=>{
    if($('#aLoop') && $('#aLoop').checked && au.currentTime>=b.fin_seg){ au.currentTime=b.inicio_seg; if(!au.paused) au.play(); }
    pintarReloj();
  });
  au.addEventListener('ended',actualizarPlay);
  au.addEventListener('play',actualizarPlay); au.addEventListener('pause',actualizarPlay);
  function pintarReloj(){
    const cur=Math.max(0,au.currentTime-b.inicio_seg), dur=Math.max(0,b.fin_seg-b.inicio_seg);
    $('#aReloj').textContent=hms(cur)+' / '+hms(dur);
    if(!$('#aPista').matches(':active')) $('#aPista').value=dur? Math.round(1000*cur/dur):0;
  }
  $('#aPlay').onclick=()=>{ if(!S.audio) return avisar('No hay audio vinculado a esta sesión.'); au.paused?au.play():au.pause(); };
  $('#aRet').onclick=()=>{ au.currentTime=Math.max(b.inicio_seg,au.currentTime-5); };
  $('#aAde').onclick=()=>{ au.currentTime=Math.min(b.fin_seg,au.currentTime+5); };
  $('#aIni').onclick=()=>{ au.currentTime=b.inicio_seg; };
  $('#aVel').onchange=()=>{ au.playbackRate=parseFloat($('#aVel').value); };
  $('#aPista').oninput=()=>{ const dur=Math.max(0,b.fin_seg-b.inicio_seg); au.currentTime=b.inicio_seg+dur*($('#aPista').value/1000); pintarReloj(); };
}
function actualizarPlay(){ const btn=$('#aPlay'); if(!btn||!S.au) return;
  btn.textContent = S.au.paused? '▶ Reproducir':'❚❚ Pausa'; }

// Atajos de teclado (pensados también para el pedal más adelante)
document.addEventListener('keydown', e=>{
  if(S.abierto===null || !S.au) return;
  const b=S.abierto;
  if(e.key==='F1'){ e.preventDefault(); S.au.currentTime=Math.max(b.inicio_seg,S.au.currentTime-5); }
  else if(e.key==='F2'){ e.preventDefault(); S.audio && (S.au.paused?S.au.play():S.au.pause()); }
  else if(e.key==='F3'){ e.preventDefault(); S.au.currentTime=Math.min(b.fin_seg,S.au.currentTime+5); }
});

// Agrupa segmentos consecutivos del mismo orador en un solo turno editable,
// respetando el orden de aparición. Une el texto en un párrafo corrido.
function agruparTurnos(segs){
  const turnos=[];
  segs.forEach(s=>{
    const last=turnos[turnos.length-1];
    const orad=(s.orador||'').trim();
    if(last && (last.orador||'').trim()===orad){
      last.ids.push(s.id);
      if((s.texto||'').trim()) last.partes.push((s.texto||'').trim());
    }else{
      turnos.push({ids:[s.id], orador:s.orador||'',
        partes:(s.texto||'').trim()?[(s.texto||'').trim()]:[],
        inicio_seg:s.inicio_seg, inicio_hms:s.inicio_hms});
    }
  });
  turnos.forEach(t=> t.texto=t.partes.join(' '));
  return turnos;
}

// Pega este turno con el siguiente: sirve para el caso típico de un turno
// vacío (Whisper cortó sin detectar orador) que en realidad es del mismo
// que sigue. Es puro DOM — no toca el servidor hasta que se guarde.
function unirConSiguiente(boton){
  const cont = boton.closest('.seg');
  const sig = cont.nextElementSibling;
  if(!sig || !sig.classList.contains('seg')) return;
  cont.dataset.ids = cont.dataset.ids.split(',')
    .concat(sig.dataset.ids.split(',')).join(',');
  const orA = cont.querySelector('.orador');
  const orB = sig.querySelector('.orador');
  if(!orA.value.trim()) orA.value = orB.value;
  const edA = cont.querySelector('.editor-rico');
  const edB = sig.querySelector('.editor-rico');
  const htmlA = edA.innerHTML.trim(), htmlB = edB.innerHTML.trim();
  edA.innerHTML = (htmlA && htmlB) ? (htmlA + ' ' + htmlB) : (htmlA || htmlB);
  // Si el turno pegado ya era el último, este pasa a serlo.
  if(!sig.nextElementSibling) cont.querySelector('.b-unir').remove();
  sig.remove();
  avisar('Turnos unidos — no olvides guardar.');
}

// Vuelve a pedir el bloque abierto (vía /api/esteno/tomar, que ya lo tiene
// tomado, así que solo refresca) para traer los segmentos tal como quedaron
// después de un /api/dividir en el servidor.
async function refrescarSegmentosAbiertos(){
  const b = S.abierto;
  const r = await api('/api/esteno/tomar',
    {sesion_id:S.sesion, indice:b.indice, corrector:S.corrector||'anónimo'});
  S.abierto = r.bloque; S.segmentos = r.segmentos;
}

// Divide un turno en dos oradores, apoyándose en /api/dividir (la misma
// ruta que usa la pantalla principal) — necesita el texto SIN formato de
// los segmentos originales, porque el corte se calcula por caracteres
// sobre lo que hay guardado en la base. Si el turno ya tiene formato
// (negritas, etc.) de una edición previa, no se puede dividir así de
// seguro: se le pide al corrector que guarde primero y reabra el bloque.
function dividirAqui(boton){
  document.querySelectorAll('.divisor').forEach(e=> e.remove());
  const cont = boton.closest('.seg');
  const ids = cont.dataset.ids.split(',').map(Number);
  const segs = ids.map(id => S.segmentos.find(s=> s.id===id)).filter(Boolean);
  if(!segs.length) return;
  if(segs.some(s=> (s.texto||'').indexOf('<') !== -1)){
    avisar('Este turno ya tiene formato aplicado; guarda el avance y vuelve '
          + 'a abrir el bloque para dividirlo mientras siga en texto plano.');
    return;
  }
  const full = segs.map(s=> s.texto||'').join(' ');
  const frases = []; let m; const rex = /[^.?!]+[.?!]*\\s*/g;
  while((m = rex.exec(full)) !== null){
    if(m[0].trim()) frases.push({txt:m[0].trim(), off:m.index});
    if(rex.lastIndex === m.index) rex.lastIndex++;
  }
  if(frases.length < 2){
    avisar('Muy corto para dividir por frases.');
    return;
  }
  const panel = document.createElement('div');
  panel.className = 'divisor';
  panel.innerHTML =
    '<p class="ayuda">Haz clic en la frase donde empieza el '
    + '<strong>segundo orador</strong>:</p>'
    + '<div class="frases">'
    + frases.map((f,k)=> '<span class="frase" data-k="'+k+'" data-off="'
        + f.off+'">'+esc(f.txt)+'</span>').join(' ')
    + '</div>'
    + '<div class="split-oradores" hidden>'
    + '<label>1ª parte <input class="o1" list="dlNombres"></label>'
    + '<label>2ª parte <input class="o2" list="dlNombres"></label>'
    + '<button class="btn pri chico b-div-ok">Dividir aquí</button>'
    + '<button class="btn chico b-div-cancelar">Cancelar</button></div>';
  cont.querySelector('.campos').appendChild(panel);
  let corte = null;
  const panelO = panel.querySelector('.split-oradores');
  panel.querySelectorAll('.frase').forEach(sp => sp.onclick = () => {
    const k = +sp.dataset.k;
    if(k === 0){ avisar('Elige una frase posterior a la primera.'); return; }
    corte = +sp.dataset.off;
    panel.querySelectorAll('.frase').forEach((s,j)=> s.classList.toggle('p2', j>=k));
    panelO.hidden = false;
    panelO.querySelector('.o1').value = cont.querySelector('.orador').value;
    panelO.querySelector('.o2').focus();
  });
  panel.querySelector('.b-div-cancelar').onclick = () => panel.remove();
  panel.querySelector('.b-div-ok').onclick = async () => {
    const o1 = panelO.querySelector('.o1').value.trim();
    const o2 = panelO.querySelector('.o2').value.trim();
    if(corte === null || !o1 || !o2){
      avisar('Marca la frase y escribe los dos oradores.'); return;
    }
    try{
      await api('/api/dividir',
        {sesion_id:S.sesion, ids, corte, orador1:o1, orador2:o2});
      avisar('Turno dividido.');
      await refrescarSegmentosAbiertos();
      vistaEditor();
    }catch(e){ avisar(e.message); }
  };
}

// Al guardar: el orador se aplica a TODOS los segmentos del turno; el texto
// (ya unido y corregido) se guarda en el PRIMER segmento y los demás quedan
// vacíos, para no duplicar. Se conservan ids y tiempos (audio intacto).
function recogerTurnos(){
  const payload=[];
  Array.from(app.querySelectorAll('.seg')).forEach(d=>{
    const ids=(d.dataset.ids||'').split(',').map(x=>parseInt(x,10)).filter(n=>!isNaN(n));
    const orador=d.querySelector('.orador').value;
    const ed=d.querySelector('.editor-rico');
    const texto = ed ? ed.innerHTML : (d.querySelector('.txt')||{value:''}).value;
    ids.forEach((id,k)=> payload.push({id, orador, texto: k===0? texto : ''}));
  });
  return payload;
}
async function guardar(terminar){
  const b=S.abierto;
  try{
    const r=await api('/api/esteno/guardar',{sesion_id:S.sesion,indice:b.indice,
      corrector:S.corrector||'anónimo', segmentos:recogerTurnos(), terminar});
    avisar(terminar? ('Bloque '+(b.indice+1)+' terminado.') : ('Avance guardado.'));
    if(terminar){ limpiarEditor(); await cargarEstado(); vistaCorrector(); iniciarPoll(); }
  }catch(e){ avisar(e.message); }
}
async function cerrar(){
  const b=S.abierto;
  try{ await api('/api/esteno/soltar',{sesion_id:S.sesion,indice:b.indice,corrector:S.corrector||'anónimo'}); }catch(e){}
  limpiarEditor(); await cargarEstado(); vistaCorrector(); iniciarPoll();
}
function limpiarEditor(){
  if(S.au){ try{S.au.pause();}catch(e){} S.au=null; }
  clearInterval(S.timerLatido); S.timerLatido=null; S.abierto=null; S.segmentos=[];
}
// Si cierra la pestaña con un bloque abierto, liberarlo.
window.addEventListener('beforeunload', ()=>{
  if(S.abierto!==null){
    try{ navigator.sendBeacon('/api/esteno/soltar', new Blob([JSON.stringify({sesion_id:S.sesion,indice:S.abierto.indice,corrector:S.corrector||'anónimo'})],{type:'application/json'})); }catch(e){}
  }
});

/* ---------------- Arranque + refresco ---------------- */
function iniciarPoll(){
  detenerPoll();
  S.timerEstado=setInterval(async ()=>{
    if(S.abierto!==null) return;             // no molestar mientras se edita
    if(document.hidden) return;
    // No refrescar si el usuario está escribiendo/eligiendo en un campo.
    const a=document.activeElement;
    if(a && a.matches && a.matches('input,select,textarea')) return;
    try{
      await cargarEstado();
      if(S.esAdmin){
        // Solo actualiza el panel de avance; el formulario NO se toca.
        if(document.getElementById('panelPlan')) refrescarPlan();
      }else{
        vistaCorrector();
      }
    }catch(e){}
  }, 4000);
}
function detenerPoll(){ clearInterval(S.timerEstado); S.timerEstado=null; }

async function iniciar(){
  // Quién eres y si eres admin lo dice el login, no la URL.
  try{
    const yo = await api('/api/yo');
    S.corrector = yo.email; S.esAdmin = !!yo.es_admin;
    // Para un corrector, "/" solo lo regresa aquí mismo (no tiene acceso) —
    // ese link no pinta nada; en su lugar le damos uno para cambiar de
    // sesión (el admin ya tiene su propio selector en la pantalla).
    if(!S.esAdmin){
      const l=$('#lnkPrincipal'); if(l) l.style.display='none';
      const o=$('#lnkOtraSesion'); if(o) o.style.display='';
    }
  }catch(e){
    app.innerHTML = '<div class="tarjeta"><h2>No se pudo verificar tu sesión</h2>'
      + '<p class="ayuda">Vuelve a <a href="/login">iniciar sesión</a>.</p></div>';
    return;
  }
  if(!S.sesion && !S.esAdmin){ return vistaElegirSesion(); }
  if(S.esAdmin){
    try{ if(S.sesion) await cargarEstado(); }catch(e){ avisar(e.message); }
    vistaAdmin(); if(S.sesion) iniciarPoll();
  }else{
    try{ await cargarEstado(); }catch(e){ avisar(e.message); }
    try{ const r=await api('/api/esteno/nombres?sesion='+S.sesion); S.nombres=r.nombres||[]; }catch(e){}
    vistaCorrector(); iniciarPoll();
  }
}
// Landing del corrector cuando entra sin ?sesion= (p. ej. justo tras el
// login): elige de la misma lista de sesiones que ve el admin.
async function vistaElegirSesion(){
  $('#rol').className='rol corr'; $('#rol').textContent='Corrector: '+(S.corrector||'—');
  let sesiones=[]; try{ sesiones=await api('/api/sesiones'); }catch(e){}
  if(!sesiones.length){
    app.innerHTML = '<div class="tarjeta"><h2>Aún no hay sesiones</h2>'
      + '<p class="ayuda">Cuando haya una sesión en curso o terminada, aparecerá aquí.</p></div>';
    return;
  }
  const opts = '<option value="0">— Elige una sesión —</option>' +
    sesiones.map(s=>'<option value="'+s.id+'">#'+s.id+' — '+esc((s.titulo||'').slice(0,70))+'</option>').join('');
  app.innerHTML = '<div class="tarjeta"><h2>¿Qué sesión vas a corregir?</h2>'
    + '<div class="fila"><label class="campo">Sesión<select id="selSesCorr">'+opts+'</select></label></div></div>';
  $('#selSesCorr').onchange = e => { if(e.target.value) location.search='?sesion='+e.target.value; };
}
iniciar();
</script>
</body>
</html>
"""


# ---------------------------------------------------------------------------
# Compactación: une físicamente los registros consecutivos del mismo orador
# ---------------------------------------------------------------------------

def hms(segundos):
    segundos = int(segundos or 0)
    h, r = divmod(segundos, 3600)
    m, s = divmod(r, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


# ---------------------------------------------------------------------------
# Conversor HTML → python-docx
# ---------------------------------------------------------------------------

class _NodoHTML:
    """Árbol simplificado del HTML: tag, atributos, hijos (texto o nodos)."""
    def __init__(self, tag="", attrs=None):
        self.tag = tag.lower() if tag else ""
        self.attrs = dict(attrs or [])
        self.children = []

    def texto_plano(self):
        out = []
        for c in self.children:
            if isinstance(c, str):
                out.append(c)
            else:
                out.append(c.texto_plano())
        return "".join(out)


class _Parser(html.parser.HTMLParser):
    VOID = {"br", "hr", "img", "input", "meta", "link"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = _NodoHTML("root")
        self.stack = [self.root]

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        nodo = _NodoHTML(tag, attrs)
        self.stack[-1].children.append(nodo)
        if tag not in self.VOID:
            self.stack.append(nodo)

    def handle_startendtag(self, tag, attrs):
        self.stack[-1].children.append(_NodoHTML(tag.lower(), attrs))

    def handle_endtag(self, tag):
        tag = tag.lower()
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                self.stack = self.stack[:i]
                break

    def handle_data(self, data):
        self.stack[-1].children.append(data)


def _es_bloque(tag):
    return tag in {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
                   "blockquote", "ul", "ol", "li", "br", "hr"}


def html_a_word(doc, texto_html, orador=None,
                fuente_nombre="Arial", fuente_pt=11,
                justify=True, first_indent_cm=None):
    """Inserta en `doc` los párrafos y runs equivalentes al HTML dado.
    Si `orador` no es None, el primer párrafo lleva el nombre en negrita.
    Devuelve el número de párrafos añadidos."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not texto_html or not texto_html.strip():
        return 0

    # Si el texto no contiene etiquetas HTML, lo tratamos como texto plano
    # (retrocompatibilidad con registros anteriores al formato enriquecido).
    if "<" not in texto_html:
        bloques = [b.strip() for b in texto_html.replace("\r", "").split("\n")
                   if b.strip()]
        if not bloques:
            return 0
        for k, bloque in enumerate(bloques):
            p = doc.add_paragraph()
            if justify:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if k == 0 and orador:
                r = p.add_run(f"{orador.upper()}: ")
                r.bold = True
                r.font.name = fuente_nombre
                r.font.size = Pt(fuente_pt)
            elif first_indent_cm and k > 0:
                p.paragraph_format.first_line_indent = Cm(first_indent_cm)
            r2 = p.add_run(bloque)
            r2.font.name = fuente_nombre
            r2.font.size = Pt(fuente_pt)
        return len(bloques)

    # Parsear el HTML.
    parser = _Parser()
    parser.feed(texto_html)
    root = parser.root

    parrafos_añadidos = 0
    primer = True

    def nuevo_parrafo(align=None):
        nonlocal parrafos_añadidos, primer
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        elif justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if primer and orador:
            r = p.add_run(f"{orador.upper()}: ")
            r.bold = True
            r.font.name = fuente_nombre
            r.font.size = Pt(fuente_pt)
            primer = False
        elif first_indent_cm and parrafos_añadidos > 0:
            p.paragraph_format.first_line_indent = Cm(first_indent_cm)
        parrafos_añadidos += 1
        return p

    def alinear(nodo):
        style = nodo.attrs.get("style", "")
        if "center" in style:
            return WD_ALIGN_PARAGRAPH.CENTER
        if "right" in style:
            return WD_ALIGN_PARAGRAPH.RIGHT
        cls = nodo.attrs.get("class", "")
        if "text-center" in cls:
            return WD_ALIGN_PARAGRAPH.CENTER
        if "text-right" in cls:
            return WD_ALIGN_PARAGRAPH.RIGHT
        return None

    def volcar_inline(p, nodo, bold=False, italic=False,
                      underline=False, strike=False, color=None):
        """Recorre un nodo inline y agrega runs al párrafo p."""
        for c in nodo.children:
            if isinstance(c, str):
                if c:
                    r = p.add_run(c)
                    r.bold = bold
                    r.italic = italic
                    r.underline = underline
                    r.font.strike = strike
                    r.font.name = fuente_nombre
                    r.font.size = Pt(fuente_pt)
                    if color:
                        try:
                            hex_c = color.lstrip("#")
                            r.font.color.rgb = RGBColor(
                                int(hex_c[0:2], 16),
                                int(hex_c[2:4], 16),
                                int(hex_c[4:6], 16))
                        except Exception:
                            pass
            else:
                t = c.tag
                nb = bold or t in ("b", "strong")
                ni = italic or t in ("i", "em")
                nu = underline or t == "u"
                ns = strike or t in ("s", "strike", "del")
                nc = color
                # color inline: style="color:#rrggbb"
                st = c.attrs.get("style", "")
                if "color:" in st:
                    try:
                        nc = st.split("color:")[1].split(";")[0].strip()
                    except Exception:
                        pass
                if t == "br":
                    p.add_run("\n")
                elif t in ("b", "strong", "i", "em", "u", "s", "strike",
                            "del", "span", "a", "mark"):
                    volcar_inline(p, c, nb, ni, nu, ns, nc)
                elif _es_bloque(t):
                    # bloque dentro de inline: ignoramos estructura, volcamos texto
                    volcar_inline(p, c, nb, ni, nu, ns, nc)
                else:
                    volcar_inline(p, c, nb, ni, nu, ns, nc)

    def procesar(nodo, en_lista=None, nivel=0):
        # Texto/tags inline (b, i, ...) que aparecen sueltos en la raíz, sin
        # un <p>/<div> que los envuelva (típico de un contenteditable de una
        # sola línea: el navegador no arma <div> hasta que se pulsa Enter).
        # Se van acumulando en un mismo párrafo en vez de uno nuevo por nodo,
        # y cada tag inline aplica su propio estilo (antes se perdía).
        parrafo_suelto = None
        for c in nodo.children:
            if isinstance(c, str):
                if c.strip():
                    # Solo se recorta el espacio inicial si arranca párrafo;
                    # el resto se conserva (separa palabras de tags vecinos,
                    # p. ej. "texto <b>negrita</b> más texto").
                    txt = c.lstrip() if parrafo_suelto is None else c
                    if parrafo_suelto is None:
                        parrafo_suelto = nuevo_parrafo()
                    r = parrafo_suelto.add_run(txt)
                    r.font.name = fuente_nombre
                    r.font.size = Pt(fuente_pt)
            elif isinstance(c, _NodoHTML):
                t = c.tag
                aln = alinear(c)

                if t in ("b", "strong", "i", "em", "u", "s",
                            "strike", "span", "a", "mark"):
                    if parrafo_suelto is None:
                        parrafo_suelto = nuevo_parrafo(align=aln)
                    comodin = _NodoHTML("_wrap")
                    comodin.children = [c]
                    volcar_inline(parrafo_suelto, comodin)
                    continue

                parrafo_suelto = None   # cualquier otro tag cierra el suelto

                if t in ("p", "div", "blockquote"):
                    # Verificar si tiene hijos bloque para no crear párrafo vacío
                    tiene_bloque = any(isinstance(h, _NodoHTML)
                                       and _es_bloque(h.tag)
                                       for h in c.children)
                    if tiene_bloque:
                        procesar(c, en_lista, nivel)
                    else:
                        p = nuevo_parrafo(align=aln)
                        volcar_inline(p, c)

                elif t in ("h1", "h2", "h3", "h4"):
                    p = nuevo_parrafo(align=aln or WD_ALIGN_PARAGRAPH.LEFT)
                    for h in c.children:
                        if isinstance(h, str):
                            r = p.add_run(h)
                        else:
                            r = p.add_run(h.texto_plano())
                        r.bold = True
                        tamaños = {"h1": 16, "h2": 14, "h3": 12, "h4": 11}
                        r.font.size = Pt(tamaños.get(t, fuente_pt))
                        r.font.name = fuente_nombre

                elif t == "br":
                    p = nuevo_parrafo()

                elif t == "hr":
                    p = nuevo_parrafo()
                    p.add_run("─" * 60)

                elif t in ("ul", "ol"):
                    for i, li in enumerate(c.children):
                        if isinstance(li, _NodoHTML) and li.tag == "li":
                            p = nuevo_parrafo(align=aln)
                            if t == "ul":
                                prefijo = "  " * nivel + "• "
                            else:
                                prefijo = "  " * nivel + f"{i+1}. "
                            r = p.add_run(prefijo)
                            r.font.name = fuente_nombre
                            r.font.size = Pt(fuente_pt)
                            volcar_inline(p, li)
                        elif isinstance(li, _NodoHTML) and li.tag in ("ul","ol"):
                            procesar(_NodoHTML("_wrap",
                                               []), en_lista, nivel + 1)

                else:
                    procesar(c, en_lista, nivel)

    procesar(root)

    # Si nunca se creó ningún párrafo (HTML vacío o solo etiquetas sin texto)
    if parrafos_añadidos == 0 and orador:
        p = doc.add_paragraph()
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"{orador.upper()}: ")
        r.bold = True
        r.font.name = fuente_nombre
        r.font.size = Pt(fuente_pt)
        parrafos_añadidos = 1

    return parrafos_añadidos


def _ruta_membrete():
    """Devuelve la ruta de la imagen del membrete o None si no se encuentra.
    Prioridad: variable de entorno MEMBRETE_IMG; luego varios nombres
    habituales junto al script y en la carpeta actual."""
    env = os.environ.get("MEMBRETE_IMG", "").strip()
    if env and os.path.isfile(env):
        return env
    base = os.path.dirname(os.path.abspath(__file__))
    candidatos = [
        "membrete.jpeg", "membrete.jpg", "membrete.png",
        "membrete sap_Mesa de trabajo 1.jpg.jpeg",
    ]
    for nombre in candidatos:
        for ruta in (os.path.join(base, nombre), nombre):
            if os.path.isfile(ruta):
                return ruta
    return None


MESES_ES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre",
            "diciembre"]


def _fecha_legible(valor):
    """Convierte un timestamp ISO (u otro texto) en una fecha en español
    del estilo '15 de marzo de 2026'. Si no puede, devuelve el texto tal cual."""
    if not valor:
        return "N/A"
    texto = str(valor).strip()
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            d = datetime.strptime(texto[:len(fmt) + 2], fmt)
            return f"{d.day} de {MESES_ES[d.month]} de {d.year}"
        except ValueError:
            continue
    return texto


def _campo(fila, nombre, defecto=None):
    return fila[nombre] if nombre in fila.keys() else defecto


def _generar_password_aleatoria(longitud=10):
    """Contraseña al azar para cuentas de operador (agente de captura,
    correctores): sin 0/O ni 1/l/I, para que se pueda copiar o dictar por
    teléfono sin confusiones."""
    alfabeto = "ABCDEFGHJKLMNPQRSTUVWXYZabcdefghijkmnpqrstuvwxyz23456789"
    return "".join(secrets.choice(alfabeto) for _ in range(longitud))


# ===========================================================================
#  Módulo estenográfico colaborativo: helpers
# ===========================================================================
ESTENO_LOCK_STALE = 180          # seg. sin latido tras los que un bloque se
                                 # considera abandonado y otro puede retomarlo
ESTENO_AUDIO_EXTS = ("m4a", "mp3", "wav", "ogg", "oga", "opus", "webm",
                     "mka", "flac", "aac")


def _esteno_num(v):
    try:
        return float(v)
    except (TypeError, ValueError):
        return 0.0


def _esteno_dur_sesion(con, sid):
    """Duración total conocida de la sesión, en segundos.
    Usa fin_seg si el transcriptor lo llena; si no (queda en NULL/0), se
    apoya en el inicio del último segmento (+1 s) para que ese segmento
    quede dentro de un bloque. Devuelve 0 si la sesión está vacía."""
    r = con.execute("SELECT MAX(fin_seg) AS f, MAX(inicio_seg) AS i "
                    "FROM participaciones WHERE sesion_id=?", (sid,)).fetchone()
    f = _esteno_num(r["f"]) if r else 0.0
    i = _esteno_num(r["i"]) if r else 0.0
    if f <= 0 and i <= 0:
        return 0.0
    return max(f, i + 1.0)


def _esteno_config(con, sid):
    return con.execute("SELECT * FROM esteno_config WHERE sesion_id=?",
                       (sid,)).fetchone()


def _esteno_generar(con, sid, num, dur, reset=False):
    """Crea/actualiza los bloques de tiempo de una sesión SIN partir
    intervenciones: cada bloque dura al menos `dur`, pero cierra al terminar
    la intervención (racha del mismo orador) que esté en curso al cumplirse
    el tiempo. Así, si al minuto alguien sigue hablando, ese bloque se
    extiende hasta que cambie de orador; el siguiente turno empieza limpio.

    Es 'append-only' y estable: los bloques ya 'editando'/'terminado' (o ya
    completos) quedan CONGELADOS y nunca se reacomodan; solo se recalcula la
    cola pendiente conforme la sesión crece en vivo.
    - reset=True: borra todo y recrea (se usa al cambiar la duración)."""
    num = max(1, int(num))
    dur = max(5, int(dur))
    if reset:
        con.execute("DELETE FROM esteno_bloques WHERE sesion_id=?", (sid,))
    existentes = list(con.execute("SELECT * FROM esteno_bloques WHERE "
                                  "sesion_id=? ORDER BY indice", (sid,)))
    # Congelamos los bloques tocados o ya completos; recalculamos solo la cola.
    idx_congelado = -1
    for b in existentes:
        completo = (float(b["fin_seg"]) - float(b["inicio_seg"])) >= dur - 0.5
        if b["estado"] in ("editando", "terminado") or completo:
            idx_congelado = b["indice"]
    con.execute("DELETE FROM esteno_bloques WHERE sesion_id=? AND indice>?",
                (sid, idx_congelado))
    conservados = [b for b in existentes if b["indice"] <= idx_congelado]
    if conservados:
        cubierto = float(conservados[-1]["fin_seg"])
        idx = conservados[-1]["indice"] + 1
    else:
        cubierto = 0.0
        idx = 0

    total = _esteno_dur_sesion(con, sid)
    if total <= cubierto:
        con.commit()
        return len(conservados)

    # Turnos = rachas consecutivas del mismo orador, a partir de 'cubierto'.
    segs = list(con.execute(
        "SELECT id, orador, inicio_seg FROM participaciones WHERE sesion_id=? "
        "AND inicio_seg >= ? ORDER BY inicio_seg, id", (sid, cubierto)))
    turnos = []
    for s in segs:
        o = (s["orador"] or "").strip()
        ini = float(s["inicio_seg"] or 0)
        if turnos and turnos[-1]["orador"] == o:
            turnos[-1]["ids"].append(s["id"])
        else:
            turnos.append({"orador": o, "inicio": ini, "ids": [s["id"]]})
    if not turnos:
        con.commit()
        return len(conservados)
    # Fin de cada turno = inicio del siguiente (el último, la duración total).
    for i, t in enumerate(turnos):
        t["fin"] = turnos[i + 1]["inicio"] if i + 1 < len(turnos) else total

    # Empacar turnos en bloques de >= dur, cerrando en cambio de orador.
    i = 0
    ini_bloque = cubierto
    while i < len(turnos):
        j = i
        while j < len(turnos) and (turnos[j]["fin"] - ini_bloque) < dur:
            j += 1
        if j >= len(turnos):
            j = len(turnos) - 1        # cola: no alcanzó 'dur' todavía
        fin_bloque = turnos[j]["fin"]
        con.execute("INSERT INTO esteno_bloques (sesion_id,indice,inicio_seg,"
                    "fin_seg,slot,estado) VALUES (?,?,?,?,?, 'pendiente')",
                    (sid, idx, round(ini_bloque, 3), round(fin_bloque, 3),
                     idx % num))
        ini_bloque = fin_bloque
        idx += 1
        i = j + 1
    con.commit()
    return idx


def _esteno_extender(con, sid):
    """Agrega bloques nuevos si la sesión creció (transcripción en vivo),
    sin tocar los estados de los existentes. No hace nada si no hay config."""
    cfg = _esteno_config(con, sid)
    if not cfg:
        return 0
    return _esteno_generar(con, sid, cfg["num_correctores"],
                           cfg["bloque_seg"], reset=False)


def _esteno_segmentos(con, sid, ini, fin):
    """Segmentos cuyo inicio cae dentro de [ini, fin) del bloque.
    Selecciona * para no depender de columnas opcionales (p. ej. fin_hms,
    que no existe en todos los esquemas)."""
    return list(con.execute(
        "SELECT * FROM participaciones WHERE sesion_id=? "
        "AND inicio_seg >= ? AND inicio_seg < ? ORDER BY inicio_seg, id",
        (sid, ini, fin)))


def _esteno_dir_audio():
    env = os.environ.get("ESTENO_AUDIO_DIR", "").strip()
    if env:
        return env
    base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, "audio")


def _esteno_ruta_audio(con, sid):
    """Ruta del archivo de audio de la sesión, o None. Prioridad:
    1) ruta fijada a mano por el admin (tabla esteno_audio);
    2) audio/sesion_<id>.<ext> en la carpeta de audio (configurable);
    3) audio/<id>/<cualquier archivo> ."""
    try:
        r = con.execute("SELECT ruta FROM esteno_audio WHERE sesion_id=?",
                        (sid,)).fetchone()
        if r and r["ruta"] and os.path.isfile(r["ruta"]):
            return r["ruta"]
    except sqlite3.Error:
        pass
    carpeta = _esteno_dir_audio()
    for ext in ESTENO_AUDIO_EXTS:
        p = os.path.join(carpeta, f"sesion_{sid}.{ext}")
        if os.path.isfile(p):
            return p
    sub = os.path.join(carpeta, str(sid))
    if os.path.isdir(sub):
        for nombre in sorted(os.listdir(sub)):
            if nombre.lower().rsplit(".", 1)[-1] in ESTENO_AUDIO_EXTS:
                return os.path.join(sub, nombre)
    return None


def compactar(con, sesion_id, solo_ids=None):
    """Une en un solo registro las filas consecutivas del mismo orador
    (texto concatenado, tiempos extendidos, mejor evidencia de voz).
    Con solo_ids, solo compacta los grupos que tocan esos ids; sin él,
    compacta la sesión completa. Devuelve cuántas filas se fusionaron."""
    filas = list(con.execute(
        "SELECT * FROM participaciones WHERE sesion_id=? "
        "ORDER BY inicio_seg, id", (sesion_id,)))
    grupos, actual = [], []
    for f in filas:
        if actual and actual[-1]["orador"] == f["orador"]:
            actual.append(f)
        else:
            actual = [f]
            grupos.append(actual)
    solo = set(solo_ids) if solo_ids else None
    unidos = 0
    for g in grupos:
        if len(g) < 2:
            continue
        if solo is not None and not any(f["id"] in solo for f in g):
            continue
        base = g[0]
        texto = " ".join((f["texto"] or "").strip() for f in g).strip()
        fin = g[-1]["fin_seg"]
        sets = ["texto=?", "fin_seg=?", "fin_hms=?"]
        vals = [texto, fin, hms(fin)]
        if "voz_orador" in base.keys():
            mejor = max(g, key=lambda f: (_campo(f, "voz_similitud") or 0))
            sets += ["voz_orador=?", "voz_similitud=?"]
            vals += [_campo(mejor, "voz_orador"),
                     _campo(mejor, "voz_similitud")]
        if "revisado_ia" in base.keys():
            # se conserva el veredicto de mayor rango del grupo
            orden = {"validado": 3, "media": 2, "descartado": 1}
            rev = mot = None
            mejor = 0
            for f in g:
                rr = _campo(f, "revisado_ia")
                if orden.get(rr, 0) > mejor:
                    mejor, rev, mot = orden.get(rr, 0), rr, _campo(f, "motivo_ia")
            sets += ["revisado_ia=?", "motivo_ia=?"]
            vals += [rev, mot]
        con.execute("UPDATE participaciones SET " + ", ".join(sets)
                    + " WHERE id=?", vals + [base["id"]])
        marcas = ",".join("?" * (len(g) - 1))
        con.execute(f"DELETE FROM participaciones WHERE id IN ({marcas})",
                    [f["id"] for f in g[1:]])
        unidos += len(g) - 1
    con.commit()
    return unidos


# ---------------------------------------------------------------------------
# Resumen ejecutivo de intervenciones (prompt de analista parlamentario)
# ---------------------------------------------------------------------------

PROMPT_RESUMEN = """Actúa como analista parlamentario especializado en actividad legislativa.

Analiza la siguiente intervención de un diputado o diputada y genera un resumen ejecutivo objetivo, claro y conciso.

Instrucciones:

1. Identifica el nombre del legislador/a.
2. Resume la participación en un máximo de 500 palabras.
3. Conserva únicamente las ideas principales, propuestas, posicionamientos, críticas o argumentos relevantes.
4. Elimina saludos, agradecimientos, fórmulas protocolarias, repeticiones y expresiones retóricas.
5. Mantén un lenguaje neutral y sin opiniones.
6. Identifica el tema principal de la intervención.
7. Señala, cuando exista:
- Propuesta presentada.
- Problema señalado.
- Solicitud realizada.
- Postura a favor o en contra de algún asunto.
8. Si la intervención no contiene propuestas o posicionamientos claros, indícalo expresamente.

Formato de salida:

Diputado(a): [Nombre]

Tema:
[Tema principal]

Resumen:
[Resumen ejecutivo]

Puntos clave:
• [Punto 1]
• [Punto 2]
• [Punto 3]

Propuesta o solicitud:
[Texto o "No se identifica"]

Posicionamiento:
[A favor / En contra / Informativo / No identificado]

Intervención a analizar:
Orador registrado: {orador}

{texto}"""


def generar_resumen(orador, texto):
    """Con clave API (variable de entorno ANTHROPIC_API_KEY) genera el
    resumen automáticamente llamando a la API de Anthropic; sin clave,
    devuelve el prompt completo listo para copiarse y pegarse en Claude."""
    prompt = PROMPT_RESUMEN.format(orador=orador or "No registrado",
                                   texto=texto)
    clave = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not clave:
        return {"modo": "copiar", "prompt": prompt}
    import urllib.request
    cuerpo = json.dumps({
        "model": os.environ.get("RESUMEN_MODELO", "claude-sonnet-5"),
        "max_tokens": 1500,
        "messages": [{"role": "user", "content": prompt}],
    }).encode("utf-8")
    solicitud = urllib.request.Request(
        "https://api.anthropic.com/v1/messages", data=cuerpo,
        headers={"Content-Type": "application/json", "x-api-key": clave,
                 "anthropic-version": "2023-06-01"})
    try:
        with urllib.request.urlopen(solicitud, timeout=90) as r:
            data = json.loads(r.read().decode("utf-8"))
        partes = [b.get("text", "") for b in data.get("content", [])
                  if b.get("type") == "text"]
        resumen = "\n".join(p for p in partes if p).strip()
        if not resumen:
            raise ValueError("respuesta vacía")
        return {"modo": "auto", "resumen": resumen}
    except Exception as e:
        return {"modo": "copiar", "prompt": prompt,
                "error": "No se pudo generar automáticamente "
                         f"({type(e).__name__}); se copió el prompt para "
                         "usarlo manualmente en Claude."}


PROMPT_CORRECCION = """Actúa como corrector de estilo y editor profesional para versiones estenográficas legislativas.
Tu tarea es pulir el siguiente fragmento aplicando rigurosamente estas reglas:

1. ORTOGRAFÍA Y MAYÚSCULAS: Corrige nombres propios, lugares e instituciones. Capitaliza SIEMPRE los cargos públicos, honoríficos y legislativos (ej. Diputado, Diputada, Presidente, Presidenta, Secretario, Secretaria, Gobernadora, Alcalde, Pleno, Legislatura, Congreso).
2. PUNTUACIÓN Y PÁRRAFOS: Mejora la legibilidad rompiendo oraciones interminables (usa puntos y seguido en lugar de un exceso de comas). Divide los bloques de texto muy largos en párrafos más cortos y digeribles, estructurando las ideas de forma clara y ejecutiva (máximo 4-5 líneas por párrafo si es posible). Separa cada párrafo con un salto de línea (\\n).
3. NOMBRES DE DIPUTADOS: Si en el texto aparece el nombre de un legislador o legisladora escrito de forma incorrecta o incompleta, y se corresponde claramente con uno de la lista de referencia, corrígelo a la forma EXACTA de esa lista. No inventes ni fuerces coincidencias dudosas.
4. REGLA ESTRICTA: NO inventes información, no resumas y NO modifiques el sentido ni las ideas del orador. Respeta la naturalidad del discurso, solo dale un formato profesional y ortográficamente impecable.

Devuelve ÚNICAMENTE el texto corregido, sin introducciones, saludos ni comentarios.{lista_diputados}

Texto original:
{texto}"""


def corregir_texto_ia(texto, catalogo=None):
    """Corrige ortografía, gramática, estilo y nombres de un fragmento usando
    la API de Anthropic. `catalogo` es una lista opcional de nombres de
    diputados que se usa como referencia para uniformar los nombres."""
    clave = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not clave:
        return {"error": "No hay clave API configurada (variable "
                         "ANTHROPIC_API_KEY)."}

    lista = ""
    if catalogo:
        # Limitamos para no inflar el prompt en catálogos enormes.
        nombres = "\n".join("- " + n for n in catalogo[:400])
        lista = ("\n\nLista de referencia de nombres correctos de "
                 "diputados y diputadas:\n" + nombres)
    prompt = PROMPT_CORRECCION.format(texto=texto, lista_diputados=lista)

    # Escalamos el límite de salida al tamaño del texto para no truncar
    # intervenciones largas (aprox. 1 token ~ 4 caracteres, con holgura).
    max_tokens = max(1024, min(8192, int(len(texto) / 2) + 512))

    import urllib.request
    cuerpo = json.dumps({
        "model": os.environ.get("CORRECCION_MODELO", "claude-haiku-4-5"),
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": prompt}],
    }).encode("utf-8")

    solicitud = urllib.request.Request(
        "https://api.anthropic.com/v1/messages", data=cuerpo,
        headers={"Content-Type": "application/json", "x-api-key": clave,
                 "anthropic-version": "2023-06-01"})
    try:
        with urllib.request.urlopen(solicitud, timeout=120) as r:
            data = json.loads(r.read().decode("utf-8"))
        partes = [b.get("text", "") for b in data.get("content", [])
                  if b.get("type") == "text"]
        texto_corregido = "\n".join(p for p in partes if p).strip()
        if not texto_corregido:
            return {"error": "La IA devolvió una respuesta vacía."}
        return {"texto_corregido": texto_corregido}
    except Exception as e:
        return {"error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Edición de contenido: numeración de los puntos del orden del día
# ---------------------------------------------------------------------------

PROMPT_ESTRUCTURA = """Eres editor parlamentario. Recibes las intervenciones de una sesión legislativa, cada una con un identificador [ID], el nombre del orador y su texto, en orden cronológico.

Al inicio de la reunión se LEE el ORDEN DEL DÍA: la lista de puntos a tratar. Tu tarea es:
1. Encontrar esa lectura y COPIAR el texto COMPLETO de cada punto, tal cual se enuncia.
2. Indicar en qué intervención EMPIEZA el desahogo (la discusión) de cada punto.

Devuelve ÚNICAMENTE un objeto JSON válido, sin texto adicional ni ```json, con esta forma exacta:
{{
  "secciones": [
    {{"ancla_id": 145, "titulo": "Análisis de la Iniciativa de Decreto por el que se expide la Ley ..., presentada por ..., en su caso, intervención de personas servidoras públicas del Gobierno del Estado."}},
    {{"ancla_id": 320, "titulo": "Clausura de la reunión."}}
  ]
}}

Reglas estrictas:
- El "titulo" debe ser el TEXTO ÍNTEGRO Y LITERAL del punto, completo, tal como aparece en la lectura del orden del día. NO lo acortes, NO lo resumas, NO lo parafrasees. Conserva la redacción formal completa, incluidos nombres propios, cargos y honoríficos.
- NO incluyas prefijos de numeración en el "titulo" ("1.", "Punto número uno", "Primer punto", etc.); la numeración se agrega aparte.
- "ancla_id" DEBE ser uno de los [ID] de abajo: la intervención donde EMPIEZA el desahogo de ese punto (no donde solo se lee la lista).
- Devuelve los puntos en el ORDEN del orden del día.
- Si un punto se lee pero no alcanzas a ver su texto completo, usa la versión más completa que aparezca en las intervenciones.
- NO inventes puntos. NO detectes acuerdos ni votaciones. NO modifiques el texto de las intervenciones.

Intervenciones:
{intervenciones}"""


def estructurar_ia(turnos):
    """Recibe una lista de turnos [{id, orador, texto}] y pide a la IA la
    numeración de los puntos del orden del día. Devuelve un dict con
    'secciones' (cada una con ancla_id, numero y titulo), o {'error': ...}."""
    clave = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not clave:
        return {"error": "No hay clave API configurada (variable "
                         "ANTHROPIC_API_KEY)."}
    if not turnos:
        return {"error": "La sesión no tiene intervenciones."}

    # La IA necesita VER el texto completo del punto para copiarlo íntegro.
    # El orden del día se lee al inicio de la reunión, así que damos un
    # extracto generoso a las primeras intervenciones (donde suele leerse) y
    # uno más corto al resto (basta para reconocer dónde empieza cada punto).
    lineas = []
    for i, t in enumerate(turnos):
        limite = 2500 if i < 40 else 400
        extracto = " ".join((t["texto"] or "").split())[:limite]
        lineas.append(f"[ID {t['id']}] {t['orador']}: {extracto}")
    cuerpo_prompt = "\n".join(lineas)
    # Guarda de tamaño: si la sesión es enorme, recomprimimos para no
    # desbordar el contexto (el orden del día ya quedó en las primeras).
    if len(cuerpo_prompt) > 180000:
        lineas = []
        for i, t in enumerate(turnos):
            limite = 2500 if i < 40 else 200
            extracto = " ".join((t["texto"] or "").split())[:limite]
            lineas.append(f"[ID {t['id']}] {t['orador']}: {extracto}")
        cuerpo_prompt = "\n".join(lineas)
    prompt = PROMPT_ESTRUCTURA.format(intervenciones=cuerpo_prompt)

    import urllib.request
    cuerpo = json.dumps({
        "model": os.environ.get("ESTRUCTURA_MODELO", "claude-sonnet-5"),
        "max_tokens": 8192,
        "messages": [{"role": "user", "content": prompt}],
    }).encode("utf-8")
    solicitud = urllib.request.Request(
        "https://api.anthropic.com/v1/messages", data=cuerpo,
        headers={"Content-Type": "application/json", "x-api-key": clave,
                 "anthropic-version": "2023-06-01"})
    try:
        with urllib.request.urlopen(solicitud, timeout=180) as r:
            data = json.loads(r.read().decode("utf-8"))
        partes = [b.get("text", "") for b in data.get("content", [])
                  if b.get("type") == "text"]
        crudo = "\n".join(p for p in partes if p).strip()
        # Quitamos posibles vallas de código ```json ... ```
        if crudo.startswith("```"):
            crudo = crudo.split("\n", 1)[1] if "\n" in crudo else crudo
            crudo = crudo.rsplit("```", 1)[0]
        # Aislamos el objeto JSON por si viene con texto alrededor.
        ini, fin = crudo.find("{"), crudo.rfind("}")
        if ini != -1 and fin != -1:
            crudo = crudo[ini:fin + 1]
        obj = json.loads(crudo)
        # Orden cronológico de los turnos, para numerar los puntos en orden.
        posicion = {t["id"]: i for i, t in enumerate(turnos)}
        ids_validos = set(posicion)

        crudas = []
        vistas = set()
        for s in obj.get("secciones", []):
            try:
                aid = int(s.get("ancla_id"))
            except (TypeError, ValueError):
                continue
            if aid not in ids_validos or aid in vistas:
                continue
            titulo = (s.get("titulo") or "").strip()
            # Quitamos prefijos de numeración que la IA pudiera dejar; la
            # numeración la ponemos nosotros para que sea consistente.
            #  - "1. ", "2) ", "3 - " …
            titulo = re.sub(r"^\s*\d+\s*[.)-]\s*", "", titulo).strip()
            #  - "Punto número uno:", "Primer punto.", "Punto 1 -" …
            titulo = re.sub(
                r"^\s*(punto\s+(n[uú]mero\s+)?[\w]+|"
                r"(primer|segund|tercer|cuart|quint|sext|s[eé]ptim|octav|"
                r"noven|d[eé]cim)[oa]\s+punto)\s*[.:)-]?\s*",
                "", titulo, flags=re.IGNORECASE).strip()
            if not titulo:
                continue
            vistas.add(aid)
            crudas.append((posicion[aid], aid, titulo))

        # Ordenamos por aparición y numeramos 1., 2., 3., …
        crudas.sort(key=lambda x: x[0])
        secciones = []
        for numero, (_, aid, titulo) in enumerate(crudas, start=1):
            secciones.append({
                "ancla_id": aid,
                "numero": numero,
                "titulo": f"{numero}. {titulo}",
            })
        return {"secciones": secciones}
    except json.JSONDecodeError:
        return {"error": "La IA no devolvió un JSON válido; intenta de nuevo."}
    except Exception as e:
        return {"error": f"{type(e).__name__}: {e}"}


def turnos_de_sesion(con, sesion_id):
    """Agrupa las participaciones en turnos consecutivos por orador y
    devuelve [{id, orador, texto}] usando el id de la primera fila del turno
    como ancla."""
    filas = con.execute(
        "SELECT id, orador, texto FROM participaciones "
        "WHERE sesion_id=? ORDER BY inicio_seg, id", (sesion_id,)).fetchall()
    turnos, actual = [], None
    for f in filas:
        if actual is None or actual["orador"] != f["orador"]:
            actual = {"id": f["id"], "orador": f["orador"],
                      "texto": (f["texto"] or "")}
            turnos.append(actual)
        else:
            actual["texto"] += " " + (f["texto"] or "")
    return turnos


# ---------------------------------------------------------------------------
# Servidor
# ---------------------------------------------------------------------------

class Manejador(BaseHTTPRequestHandler):
    ruta_db = "sesiones.db"
    # Lanzador de transcripción (costura 2)
    ruta_transcriptor = "transcribir_en_vivo_c3.py"
    ruta_contextos = "contextos.json"
    # URL interna de la API (api/) para crear trabajos "desde evento real"
    # del sistema de registro parlamentario. Solo se usa dentro de Docker,
    # donde "api" es el nombre del servicio hermano en el mismo compose.
    api_interna = "http://api:8000"
    proc_activo = None          # subprocess.Popen en curso (o None)
    log_transcripcion = None    # ruta del log de la transcripción en curso
    _lock_proc = threading.Lock()
    requiere_login = False

    def log_message(self, *args):        # silenciar el registro por consola
        pass

    def _db(self):
        # timeout: espera hasta 30 s si otro proceso (el transcriptor en
        # vivo) está escribiendo, en vez de fallar con "database is locked".
        con = sqlite3.connect(self.ruta_db, timeout=30)
        con.row_factory = sqlite3.Row
        # WAL permite leer mientras otro escribe; busy_timeout refuerza la
        # espera a nivel del motor SQLite.
        try:
            con.execute("PRAGMA journal_mode=WAL")
            con.execute("PRAGMA busy_timeout=30000")
            con.execute("PRAGMA synchronous=NORMAL")
        except sqlite3.OperationalError:
            pass
        # Tabla de resúmenes guardados (uno por bloque, identificado por el
        # primer id de sus segmentos). Se crea sola la primera vez.
        con.execute("""
            CREATE TABLE IF NOT EXISTS resumenes (
                id         INTEGER PRIMARY KEY AUTOINCREMENT,
                sesion_id  INTEGER,
                ancla_id   INTEGER UNIQUE,
                orador     TEXT,
                resumen    TEXT,
                creado     TEXT
            )""")
        # Estructura editorial de la sesión: encabezados de sección (orden
        # del día) numerados. Se guarda como JSON, uno por sesión.
        con.execute("""
            CREATE TABLE IF NOT EXISTS estructura (
                sesion_id  INTEGER PRIMARY KEY,
                datos      TEXT,
                creado     TEXT
            )""")
        # --- Módulo estenográfico colaborativo (corrección por tramos) ---
        # Configuración por sesión: cuántos correctores y duración del bloque.
        con.execute("""
            CREATE TABLE IF NOT EXISTS esteno_config (
                sesion_id       INTEGER PRIMARY KEY,
                num_correctores INTEGER NOT NULL,
                bloque_seg      INTEGER NOT NULL,
                nombres         TEXT,
                creado          TEXT,
                actualizado     TEXT
            )""")
        # Un renglón por bloque de tiempo, con su asignación y estado/bloqueo.
        con.execute("""
            CREATE TABLE IF NOT EXISTS esteno_bloques (
                sesion_id     INTEGER,
                indice        INTEGER,
                inicio_seg    REAL,
                fin_seg       REAL,
                slot          INTEGER,
                estado        TEXT DEFAULT 'pendiente',
                tomado_por    TEXT,
                tomado_en     TEXT,
                terminado_por TEXT,
                terminado_en  TEXT,
                PRIMARY KEY (sesion_id, indice)
            )""")
        # Ruta explícita del audio de una sesión (si el descubrimiento
        # automático no lo encuentra, el admin la fija a mano).
        con.execute("""
            CREATE TABLE IF NOT EXISTS esteno_audio (
                sesion_id INTEGER PRIMARY KEY,
                ruta      TEXT
            )""")
        con.commit()
        return con

    def _responder(self, cuerpo, tipo="application/json; charset=utf-8",
                   codigo=200):
        datos = (cuerpo if isinstance(cuerpo, bytes)
                 else json.dumps(cuerpo, ensure_ascii=False).encode("utf-8"))
        try:
            self.send_response(codigo)
            self.send_header("Content-Type", tipo)
            self.send_header("Content-Length", str(len(datos)))
            # Evita que el navegador muestre una versión cacheada tras
            # actualizar el archivo (la causa típica de "lo veo igual").
            self.send_header("Cache-Control", "no-store, must-revalidate")
            self.end_headers()
            self.wfile.write(datos)
        except OSError:
            # El cliente cerró la conexión antes de recibir toda la
            # respuesta (recarga, cambio de sesión, o el auto-refresh "En
            # vivo" que lanza una petición nueva y aborta la anterior). En
            # Windows esto llega como ConnectionAbortedError (WinError
            # 10053); en Linux/Mac como BrokenPipe/ConnectionReset. Todos
            # son inofensivos: se ignoran para no ensuciar la consola.
            pass

    # -- Autenticación opcional (--requiere-login) --------------------

    def _usuario_autenticado(self):
        if not self.requiere_login:
            return True
        m = re.search(rf"{COOKIE_SESION}=([^;]+)",
                      self.headers.get("Cookie", ""))
        if not m:
            return False
        try:
            _jwt.decode(m.group(1), _jwt_secret, algorithms=[_jwt_algoritmo])
            return True
        except Exception:
            return False

    def _redirigir_login(self):
        self.send_response(302)
        self.send_header("Location", "/login")
        self.end_headers()

    def _pagina_login(self):
        self._responder(PAGINA_LOGIN.encode("utf-8"),
                        "text/html; charset=utf-8")

    def _procesar_login(self, datos):
        email = (datos.get("email") or "").strip().lower()
        password = datos.get("password") or ""
        usuario = _obtener_usuario(email) if email else None
        if not usuario or not _verificar_password(password,
                                                   usuario["password_hash"]):
            return self._responder({"error": "Email o contraseña incorrectos"},
                                   codigo=401)
        token = _crear_token(usuario["email"])
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header(
            "Set-Cookie",
            f"{COOKIE_SESION}={token}; Path=/; HttpOnly; SameSite=Lax")
        cuerpo = json.dumps({"ok": True}).encode("utf-8")
        self.send_header("Content-Length", str(len(cuerpo)))
        self.end_headers()
        self.wfile.write(cuerpo)

    def _cerrar_sesion(self):
        self.send_response(302)
        self.send_header("Location", "/login")
        self.send_header(
            "Set-Cookie", f"{COOKIE_SESION}=; Path=/; Max-Age=0")
        self.end_headers()

    _args_transcriptor_cache = None
    _args_transcriptor_leido = False

    @classmethod
    def _args_transcriptor(cls):
        """Conjunto de opciones largas (--x) que acepta el transcriptor,
        leyendo su ayuda con --help. Se cachea una sola vez. Si no se pudo
        leer, devuelve None y el lanzador pasa todos los argumentos (como
        antes), para no romper transcriptores que no imprimen ayuda."""
        if cls._args_transcriptor_leido:
            return cls._args_transcriptor_cache
        soportadas = None
        try:
            out = subprocess.run(
                [sys.executable, cls.ruta_transcriptor, "--help"],
                capture_output=True, text=True, timeout=30)
            texto = (out.stdout or "") + (out.stderr or "")
            encontradas = set(re.findall(r"--[a-zA-Z][\w-]*", texto))
            # Solo confiamos en la detección si la ayuda se leyó de verdad.
            soportadas = encontradas or None
        except Exception:
            soportadas = None
        cls._args_transcriptor_cache = soportadas
        cls._args_transcriptor_leido = True
        return soportadas

    def _lanzar_transcripcion(self, datos):
        """Arranca transcribir_en_vivo_c3.py como proceso aparte, con el
        contexto elegido en los selects. Solo uno a la vez."""
        url = (datos.get("url") or "").strip()
        tipo = (datos.get("tipo") or "pleno").strip()
        comisiones = [c for c in (datos.get("comisiones") or []) if c]
        fecha = (datos.get("fecha") or "").strip()
        modelo = (datos.get("modelo") or "").strip()

        if not url.startswith("http"):
            return self._responder(
                {"error": "Pega una URL de video válida (http…)."}, codigo=400)
        if tipo == "comision" and not comisiones:
            return self._responder(
                {"error": "Elige al menos una comisión (o cambia a Pleno)."},
                codigo=400)
        if not os.path.isfile(self.ruta_transcriptor):
            return self._responder(
                {"error": f"No encuentro el transcriptor "
                          f"({self.ruta_transcriptor}). Ponlo en esta carpeta "
                          "o pásalo con --transcriptor."}, codigo=400)

        with Manejador._lock_proc:
            if Manejador.proc_activo and Manejador.proc_activo.poll() is None:
                return self._responder(
                    {"error": "Ya hay una transcripción en curso. Deténla "
                              "antes de iniciar otra."}, codigo=409)

            # Solo pasamos los argumentos que el transcriptor realmente acepta,
            # así funciona con versiones viejas que no conocen --tipo/--comision.
            soportadas = self._args_transcriptor()

            def acepta(flag):
                # Si no se pudo leer la ayuda (None), se asume que sí (previo).
                return soportadas is None or flag in soportadas

            cmd = [sys.executable, "-u", self.ruta_transcriptor, url]
            omitidas = []
            if acepta("--voz"):
                cmd += ["--voz"]
            # Guardar el audio permite corregir escuchando en el módulo
            # estenográfico. Si el transcriptor no lo soporta, se avisa.
            if acepta("--conservar-audio"):
                cmd += ["--conservar-audio"]
            else:
                omitidas.append("--conservar-audio")
            if acepta("--db"):
                cmd += ["--db", self.ruta_db]
            if acepta("--contextos"):
                cmd += ["--contextos", self.ruta_contextos]
            else:
                omitidas.append("--contextos")
            if modelo and acepta("--modelo"):
                cmd += ["--modelo", modelo]
            if acepta("--tipo"):
                cmd += ["--tipo", tipo]
            else:
                omitidas.append("--tipo")
            if acepta("--comision"):
                for c in comisiones:
                    cmd += ["--comision", c]
            elif comisiones:
                omitidas.append("--comision")
            if fecha:
                if acepta("--fecha"):
                    cmd += ["--fecha", fecha]
                else:
                    omitidas.append("--fecha")

            carpeta = os.path.dirname(os.path.abspath(self.ruta_transcriptor))
            log = os.path.join(carpeta, "transcripcion_en_curso.log")
            entorno = dict(os.environ, PYTHONIOENCODING="utf-8",
                           PYTHONUNBUFFERED="1")
            try:
                flog = open(log, "wb")
                if omitidas:
                    flog.write((
                        "[aviso] Tu transcriptor no acepta estos argumentos y "
                        "se omitieron: " + ", ".join(omitidas) + ".\n"
                        "        La transcripción arrancará igual, pero la "
                        "sesión no quedará etiquetada con esos datos.\n"
                        "        Para usarlos, actualiza "
                        + os.path.basename(self.ruta_transcriptor)
                        + " para que los acepte.\n\n").encode("utf-8"))
                flog.write(("$ " + " ".join(
                    (f'"{a}"' if " " in a else a) for a in cmd)
                    + "\n\n").encode("utf-8"))
                flog.flush()
                proc = subprocess.Popen(
                    cmd, stdout=flog, stderr=subprocess.STDOUT,
                    cwd=carpeta or None, env=entorno)
            except Exception as e:
                return self._responder(
                    {"error": f"No pude arrancar la transcripción: {e}"},
                    codigo=500)
            Manejador.proc_activo = proc
            Manejador.log_transcripcion = log

        # Comando legible para mostrar en la interfaz (transparencia)
        legible = " ".join((f'"{a}"' if " " in a else a) for a in cmd)
        respuesta = {"ok": True, "pid": proc.pid, "comando": legible}
        if omitidas:
            respuesta["aviso"] = (
                "Tu transcriptor no acepta " + ", ".join(omitidas)
                + "; se omitieron y la sesión no quedará etiquetada con esos "
                "datos. Actualiza el transcriptor para usarlos.")
        return self._responder(respuesta)

    def _detener_transcripcion(self):
        """Pide al proceso en curso que termine."""
        proc = Manejador.proc_activo
        if not (proc and proc.poll() is None):
            return self._responder(
                {"ok": True, "nota": "No había ninguna transcripción en "
                                     "curso."})
        try:
            proc.terminate()
            try:
                proc.wait(timeout=8)
            except subprocess.TimeoutExpired:
                proc.kill()
        except Exception as e:
            return self._responder(
                {"error": f"No pude detenerla: {e}"}, codigo=500)
        return self._responder({"ok": True, "detenida": True})

    # ===================================================================
    #  Módulo estenográfico colaborativo: endpoints
    # ===================================================================
    def _sid(self, q):
        try:
            return int(q.get("sesion", ["0"])[0] or 0)
        except (TypeError, ValueError):
            return 0

    def _esteno_nombres(self, q):
        """Lista de nombres para el selector del editor: diputados del
        catálogo (pleno), mesa directiva e integrantes de comisiones del
        contextos.json, y los oradores ya usados en la sesión."""
        nombres = set()
        for n in cargar_catalogo():
            if n and n.strip():
                nombres.add(n.strip())

        def agregar(v):
            if isinstance(v, list):
                for x in v:
                    agregar(x)
            elif isinstance(v, dict):
                # {cargo: nombre} o {clave: {...}} o {"nombre": "..."}
                for k in ("nombre", "name", "diputado", "integrante"):
                    if isinstance(v.get(k), str) and v[k].strip():
                        nombres.add(v[k].strip())
                        return
                for val in v.values():
                    agregar(val)
            elif isinstance(v, str) and v.strip():
                nombres.add(v.strip())

        try:
            if os.path.isfile(self.ruta_contextos):
                with open(self.ruta_contextos, encoding="utf-8") as f:
                    cfg = json.load(f)
                agregar(cfg.get("mesa_directiva"))
                for com in (cfg.get("comisiones") or {}).values():
                    agregar(com)
        except Exception:
            pass

        sid = self._sid(q)
        if sid:
            con = self._db()
            try:
                for r in con.execute(
                    "SELECT DISTINCT orador FROM participaciones WHERE "
                    "sesion_id=? AND orador IS NOT NULL AND TRIM(orador)<>''",
                    (sid,)):
                    nombres.add((r["orador"] or "").strip())
            finally:
                con.close()
        # Fuera vacíos y el marcador de 'sin identificar'.
        for basura in ("", "DESCONOCIDO", "Desconocido", "desconocido"):
            nombres.discard(basura)
        self._responder({"nombres": sorted(nombres, key=lambda s: s.lower())})

    def _esteno_estado(self, q):
        sid = self._sid(q)
        if not sid:
            return self._responder({"error": "falta sesión"}, codigo=400)
        con = self._db()
        try:
            _esteno_extender(con, sid)   # refleja el crecimiento en vivo
            cfg = _esteno_config(con, sid)
            total = _esteno_dur_sesion(con, sid)
            ahora = time.time()
            bloques = []
            for b in con.execute("SELECT * FROM esteno_bloques WHERE "
                                 "sesion_id=? ORDER BY indice", (sid,)):
                cnt = con.execute(
                    "SELECT COUNT(*) c FROM participaciones WHERE sesion_id=? "
                    "AND inicio_seg>=? AND inicio_seg<?",
                    (sid, b["inicio_seg"], b["fin_seg"])).fetchone()["c"]
                vencido = False
                if b["estado"] == "editando" and b["tomado_en"]:
                    try:
                        t = datetime.fromisoformat(b["tomado_en"]).timestamp()
                        vencido = (ahora - t) > ESTENO_LOCK_STALE
                    except ValueError:
                        pass
                bloques.append({
                    "indice": b["indice"], "inicio_seg": b["inicio_seg"],
                    "fin_seg": b["fin_seg"], "slot": b["slot"],
                    "estado": b["estado"], "tomado_por": b["tomado_por"],
                    "vencido": vencido, "terminado_por": b["terminado_por"],
                    "segmentos": cnt,
                    # Un bloque está listo para corregir cuando la transcripción
                    # ya avanzó más allá de su fin: fin_seg < total de la sesión
                    # (con un margen de 2 s para segmentos que llegan con retraso).
                    "listo": float(b["fin_seg"]) <= total + 2.0
                             if b["estado"] != "terminado" else True})
            ruta_audio = _esteno_ruta_audio(con, sid)
            ses = con.execute("SELECT id,titulo,inicio,url FROM sesiones "
                              "WHERE id=?", (sid,)).fetchone()
            # ¿Hay una transcripción activa en este momento?
            proc = Manejador.proc_activo
            sesion_corriendo = bool(proc and proc.poll() is None)
        finally:
            con.close()
        conf = None
        if cfg:
            conf = {"num_correctores": cfg["num_correctores"],
                    "bloque_seg": cfg["bloque_seg"],
                    "nombres": (json.loads(cfg["nombres"])
                                if cfg["nombres"] else [])}

        # Recalcular 'listo' con conocimiento de si la sesión sigue viva.
        # Un bloque está listo cuando:
        #   a) la sesión YA terminó → todos los bloques están listos, o
        #   b) la sesión sigue viva → solo los bloques cuyo fin_seg está al
        #      menos dur/2 segundos antes del total actual (hay margen suficiente
        #      para que el último bloque no sea el que está grabándose ahora).
        dur_bloque = cfg["bloque_seg"] if cfg else 600
        margen = max(30, dur_bloque // 2)
        for b in bloques:
            if b["estado"] == "terminado":
                b["listo"] = True
            elif not sesion_corriendo:
                b["listo"] = True          # sesión terminada → todo listo
            else:
                # En vivo: listo solo si hay al menos 'margen' segundos
                # de transcripción más allá del fin del bloque.
                b["listo"] = float(b["fin_seg"]) <= total - margen

        terminados = sum(1 for b in bloques if b["estado"] == "terminado")
        self._responder({
            "sesion": (dict(ses) if ses else None),
            "config": conf, "total_seg": total, "bloques": bloques,
            "progreso": {"terminados": terminados, "total": len(bloques)},
            "audio": bool(ruta_audio)})

    def _esteno_leer_segmentos(self, q):
        sid = self._sid(q)
        try:
            indice = int(q.get("indice", ["-1"])[0])
        except (TypeError, ValueError):
            indice = -1
        if not sid or indice < 0:
            return self._responder({"error": "faltan datos"}, codigo=400)
        con = self._db()
        try:
            b = con.execute("SELECT * FROM esteno_bloques WHERE sesion_id=? "
                            "AND indice=?", (sid, indice)).fetchone()
            if not b:
                return self._responder({"error": "bloque inexistente"},
                                       codigo=404)
            segs = [dict(r) for r in _esteno_segmentos(
                con, sid, b["inicio_seg"], b["fin_seg"])]
        finally:
            con.close()
        self._responder({"segmentos": segs,
                         "bloque": {"indice": indice,
                                    "inicio_seg": b["inicio_seg"],
                                    "fin_seg": b["fin_seg"]}})

    def _esteno_servir_audio(self, q):
        sid = self._sid(q)
        con = self._db()
        try:
            ruta = _esteno_ruta_audio(con, sid)
        finally:
            con.close()
        if not ruta or not os.path.isfile(ruta):
            return self._responder(
                {"error": "No hay audio para esta sesión."}, codigo=404)
        tam = os.path.getsize(ruta)
        tipo = mimetypes.guess_type(ruta)[0] or "application/octet-stream"
        rango = self.headers.get("Range")
        inicio, fin, parcial = 0, tam - 1, False
        if rango:
            mm = re.match(r"bytes=(\d*)-(\d*)", rango.strip())
            if mm:
                g1, g2 = mm.group(1), mm.group(2)
                if g1 == "" and g2 != "":
                    inicio, fin = max(0, tam - int(g2)), tam - 1
                else:
                    inicio = int(g1) if g1 else 0
                    fin = int(g2) if g2 else tam - 1
                inicio, fin = max(0, inicio), min(fin, tam - 1)
                if inicio > fin:
                    inicio, fin = 0, tam - 1
                parcial = True
        longitud = fin - inicio + 1
        try:
            self.send_response(206 if parcial else 200)
            self.send_header("Content-Type", tipo)
            self.send_header("Accept-Ranges", "bytes")
            self.send_header("Content-Length", str(longitud))
            if parcial:
                self.send_header("Content-Range",
                                 f"bytes {inicio}-{fin}/{tam}")
            self.send_header("Cache-Control", "no-store")
            self.end_headers()
            with open(ruta, "rb") as f:
                f.seek(inicio)
                restante = longitud
                while restante > 0:
                    trozo = f.read(min(65536, restante))
                    if not trozo:
                        break
                    self.wfile.write(trozo)
                    restante -= len(trozo)
        except (BrokenPipeError, ConnectionResetError, ConnectionAbortedError):
            pass   # el navegador cambió de rango o cerró: normal con audio

    def _esteno_configurar(self, con, datos):
        if not self._usuario_actual_admin():
            return self._responder(
                {"error": "Requiere permisos de administrador"}, codigo=403)
        sid = int(datos.get("sesion_id", 0) or 0)
        if not sid:
            return self._responder({"error": "falta sesión"}, codigo=400)
        num = max(1, min(50, int(datos.get("num_correctores", 1) or 1)))
        if datos.get("bloque_seg"):
            dur = int(datos["bloque_seg"])
        else:
            dur = int(float(datos.get("bloque_min", 10) or 10) * 60)
        dur = max(30, dur)
        nombres = [str(x).strip() for x in (datos.get("nombres") or [])
                   if str(x).strip()]
        old = _esteno_config(con, sid)
        reset = (old is None) or (old["bloque_seg"] != dur)
        ahora = datetime.now().isoformat(timespec="seconds")
        njson = json.dumps(nombres, ensure_ascii=False)
        if old:
            con.execute("UPDATE esteno_config SET num_correctores=?, "
                        "bloque_seg=?, nombres=?, actualizado=? "
                        "WHERE sesion_id=?", (num, dur, njson, ahora, sid))
        else:
            con.execute("INSERT INTO esteno_config (sesion_id, "
                        "num_correctores, bloque_seg, nombres, creado, "
                        "actualizado) VALUES (?,?,?,?,?,?)",
                        (sid, num, dur, njson, ahora, ahora))
        n = _esteno_generar(con, sid, num, dur, reset=reset)
        con.commit()
        self._responder({"ok": True, "bloques": n, "reinicio": reset})

    def _esteno_reasignar(self, con, datos):
        if not self._usuario_actual_admin():
            return self._responder(
                {"error": "Requiere permisos de administrador"}, codigo=403)
        sid = int(datos.get("sesion_id", 0) or 0)
        indice = int(datos.get("indice", -1))
        slot = int(datos.get("slot", 0))
        if not sid or indice < 0:
            return self._responder({"error": "faltan datos"}, codigo=400)
        con.execute("UPDATE esteno_bloques SET slot=? WHERE sesion_id=? "
                    "AND indice=?", (slot, sid, indice))
        con.commit()
        self._responder({"ok": True})

    def _esteno_tomar(self, con, datos):
        sid = int(datos.get("sesion_id", 0) or 0)
        indice = int(datos.get("indice", -1))
        corrector = self._usuario_actual_email() or "desconocido"
        if not sid or indice < 0:
            return self._responder({"error": "faltan datos"}, codigo=400)
        b = con.execute("SELECT * FROM esteno_bloques WHERE sesion_id=? "
                        "AND indice=?", (sid, indice)).fetchone()
        if not b:
            return self._responder({"error": "bloque inexistente"},
                                   codigo=404)
        # Verificar que la transcripción ya cubrió todo el bloque.
        if b["estado"] != "terminado":
            proc = Manejador.proc_activo
            sesion_corriendo = bool(proc and proc.poll() is None)
            if sesion_corriendo:
                total = _esteno_dur_sesion(con, sid)
                cfg = _esteno_config(con, sid)
                dur_bloque = cfg["bloque_seg"] if cfg else 600
                margen = max(30, dur_bloque // 2)
                if float(b["fin_seg"]) > total - margen:
                    mins = int(b["fin_seg"] // 60)
                    return self._responder(
                        {"error": f"Este bloque aún está siendo grabado. "
                                  f"Espera a que la transcripción pase del "
                                  f"minuto {mins} para abrirlo."},
                        codigo=409)
        if (b["estado"] == "editando" and b["tomado_por"]
                and b["tomado_por"] != corrector):
            fresco = True
            if b["tomado_en"]:
                try:
                    t = datetime.fromisoformat(b["tomado_en"]).timestamp()
                    fresco = (time.time() - t) <= ESTENO_LOCK_STALE
                except ValueError:
                    fresco = False
            if fresco:
                return self._responder(
                    {"error": "Este bloque lo está editando "
                              + b["tomado_por"] + " en este momento."},
                    codigo=409)
        ahora = datetime.now().isoformat(timespec="seconds")
        con.execute("UPDATE esteno_bloques SET estado='editando', "
                    "tomado_por=?, tomado_en=? WHERE sesion_id=? AND indice=?",
                    (corrector, ahora, sid, indice))
        con.commit()
        segs = [dict(r) for r in _esteno_segmentos(
            con, sid, b["inicio_seg"], b["fin_seg"])]
        self._responder({"ok": True, "segmentos": segs,
                         "bloque": {"indice": indice,
                                    "inicio_seg": b["inicio_seg"],
                                    "fin_seg": b["fin_seg"],
                                    "slot": b["slot"]}})

    def _esteno_latido(self, con, datos):
        sid = int(datos.get("sesion_id", 0) or 0)
        indice = int(datos.get("indice", -1))
        corrector = self._usuario_actual_email() or "desconocido"
        ahora = datetime.now().isoformat(timespec="seconds")
        con.execute("UPDATE esteno_bloques SET tomado_en=? WHERE sesion_id=? "
                    "AND indice=? AND tomado_por=? AND estado='editando'",
                    (ahora, sid, indice, corrector))
        con.commit()
        self._responder({"ok": True})

    def _esteno_soltar(self, con, datos):
        sid = int(datos.get("sesion_id", 0) or 0)
        indice = int(datos.get("indice", -1))
        corrector = self._usuario_actual_email() or "desconocido"
        con.execute("UPDATE esteno_bloques SET estado='pendiente', "
                    "tomado_por=NULL, tomado_en=NULL WHERE sesion_id=? "
                    "AND indice=? AND tomado_por=? AND estado='editando'",
                    (sid, indice, corrector))
        con.commit()
        self._responder({"ok": True})

    def _esteno_guardar(self, con, datos):
        sid = int(datos.get("sesion_id", 0) or 0)
        indice = int(datos.get("indice", -1))
        corrector = self._usuario_actual_email() or "desconocido"
        terminar = bool(datos.get("terminar"))
        segmentos = datos.get("segmentos") or []
        if not sid or indice < 0:
            return self._responder({"error": "faltan datos"}, codigo=400)
        b = con.execute("SELECT * FROM esteno_bloques WHERE sesion_id=? "
                        "AND indice=?", (sid, indice)).fetchone()
        if not b:
            return self._responder({"error": "bloque inexistente"},
                                   codigo=404)
        # Verifica que nadie más tenga un bloqueo fresco sobre el bloque.
        if (b["estado"] == "editando" and b["tomado_por"]
                and b["tomado_por"] != corrector):
            fresco = True
            if b["tomado_en"]:
                try:
                    t = datetime.fromisoformat(b["tomado_en"]).timestamp()
                    fresco = (time.time() - t) <= ESTENO_LOCK_STALE
                except ValueError:
                    fresco = False
            if fresco:
                return self._responder(
                    {"error": "Otro corrector (" + b["tomado_por"] + ") tomó "
                              "el bloque mientras editabas. Recarga para no "
                              "pisar su trabajo."}, codigo=409)
        # Solo se permiten guardar segmentos que caen dentro del bloque.
        ids_validos = {r["id"] for r in _esteno_segmentos(
            con, sid, b["inicio_seg"], b["fin_seg"])}
        n = 0
        for s in segmentos:
            try:
                idseg = int(s.get("id"))
            except (TypeError, ValueError):
                continue
            if idseg not in ids_validos:
                continue
            orador = (s.get("orador") or "").strip()
            texto = (s.get("texto") or "")
            con.execute("UPDATE participaciones SET orador=?, texto=? "
                        "WHERE id=? AND sesion_id=?",
                        (orador, texto, idseg, sid))
            n += 1
        ahora = datetime.now().isoformat(timespec="seconds")
        if terminar:
            con.execute("UPDATE esteno_bloques SET estado='terminado', "
                        "tomado_por=NULL, tomado_en=NULL, terminado_por=?, "
                        "terminado_en=? WHERE sesion_id=? AND indice=?",
                        (corrector, ahora, sid, indice))
        else:
            con.execute("UPDATE esteno_bloques SET estado='editando', "
                        "tomado_por=?, tomado_en=? WHERE sesion_id=? "
                        "AND indice=?", (corrector, ahora, sid, indice))
        con.commit()
        self._responder({"ok": True, "guardados": n,
                         "terminado": bool(terminar)})

    def _esteno_audio_ruta(self, con, datos):
        if not self._usuario_actual_admin():
            return self._responder(
                {"error": "Requiere permisos de administrador"}, codigo=403)
        sid = int(datos.get("sesion_id", 0) or 0)
        ruta = (datos.get("ruta") or "").strip()
        if not sid:
            return self._responder({"error": "falta sesión"}, codigo=400)
        con.execute("INSERT OR REPLACE INTO esteno_audio (sesion_id, ruta) "
                    "VALUES (?, ?)", (sid, ruta))
        con.commit()
        self._responder({"ok": True, "existe": bool(ruta)
                         and os.path.isfile(ruta)})

    def _token_cookie(self):
        m = re.search(rf"{COOKIE_SESION}=([^;]+)",
                      self.headers.get("Cookie", ""))
        return m.group(1) if m else None

    def _usuario_actual_email(self):
        """Email de la cuenta autenticada (según la cookie de sesión), o
        None si no se pudo determinar. Es la fuente de verdad de 'quién es'
        para el módulo estenográfico — ya no se confía en nada que mande el
        cliente. En modo sin --requiere-login no hay cuentas reales; se usa
        un identificador fijo (no hay más de un operador local a la vez)."""
        if not self.requiere_login:
            return "operador-local"
        token = self._token_cookie()
        if not token:
            return None
        try:
            payload = _jwt.decode(token, _jwt_secret, algorithms=[_jwt_algoritmo])
        except Exception:
            return None
        return payload.get("sub")

    def _usuario_actual_admin(self):
        """True si la cuenta autenticada es administrador. En modo sin
        --requiere-login no hay control de acceso que aplicar."""
        if not self.requiere_login:
            return True
        email = self._usuario_actual_email()
        if not email:
            return False
        usuario = _obtener_usuario(email)
        return bool(usuario and usuario["es_admin"])

    def _llamar_api(self, metodo, ruta, cuerpo=None):
        """Reenvía una petición a la API (api/) usando el mismo token de
        sesión del operador (no requiere iniciar sesión otra vez).
        Devuelve (datos, error) — error es un dict {"error":...,"codigo":...}
        listo para responder si algo falló, o None si salió bien."""
        token = self._token_cookie()
        if not token:
            return None, {"error": "No autenticado", "codigo": 401}

        import urllib.error
        import urllib.request
        datos_bin = json.dumps(cuerpo).encode("utf-8") if cuerpo is not None else None
        peticion = urllib.request.Request(
            f"{self.api_interna}{ruta}", data=datos_bin,
            headers={"Content-Type": "application/json",
                    "Authorization": f"Bearer {token}"},
            method=metodo)
        try:
            with urllib.request.urlopen(peticion, timeout=20) as r:
                return json.loads(r.read().decode("utf-8")), None
        except urllib.error.HTTPError as e:
            crudo = e.read().decode("utf-8", errors="replace")
            try:
                detalle = json.loads(crudo).get("detail", crudo)
            except ValueError:
                detalle = crudo
            return None, {"error": detalle, "codigo": e.code}
        except Exception as e:
            return None, {
                "error": f"No pude contactar la API ({self.api_interna}): {e}",
                "codigo": 502}

    def _crear_transcripcion_evento(self, datos):
        """Crea un trabajo a partir de un evento real del Congreso."""
        cuerpo = {
            "evento_id": datos.get("evento_id"),
            "tipo": int(datos.get("tipo", 1)),
            "modelo": datos.get("modelo") or "small",
        }
        resultado, err = self._llamar_api("POST", "/transcripciones/desde-evento",
                                          cuerpo)
        if err:
            return self._responder({"error": err["error"]}, codigo=err["codigo"])
        return self._responder(resultado)

    def _listar_trabajos_api(self):
        resultado, err = self._llamar_api("GET", "/transcripciones")
        if err:
            return self._responder({"error": err["error"]}, codigo=err["codigo"])
        return self._responder(resultado)

    def _detener_trabajo_api(self, job_id):
        resultado, err = self._llamar_api(
            "POST", f"/transcripciones/{job_id}/detener")
        if err:
            return self._responder({"error": err["error"]}, codigo=err["codigo"])
        return self._responder(resultado)

    def _listar_usuarios(self):
        """Cuentas existentes (operador de audio, correctores, admins). Solo
        responde si quien llama es admin: lo decide la API (usuario_admin_actual),
        no aquí — así el panel se puede ocultar en el front sin duplicar lógica
        de permisos."""
        resultado, err = self._llamar_api("GET", "/usuarios")
        if err:
            return self._responder({"error": err["error"]}, codigo=err["codigo"])
        return self._responder(resultado)

    def _crear_usuario(self, datos):
        """Crea una cuenta (típicamente para el agente de captura .exe o un
        corrector del módulo estenográfico) con una contraseña generada aquí
        mismo, para poder mostrarla una sola vez al admin que la está creando
        (la API solo guarda el hash; la contraseña en claro no se puede
        recuperar después)."""
        email = (datos.get("email") or "").strip().lower()
        if not email:
            return self._responder({"error": "Falta el email"}, codigo=400)
        es_admin = bool(datos.get("es_admin"))
        password = _generar_password_aleatoria()
        resultado, err = self._llamar_api(
            "POST", "/usuarios",
            {"email": email, "password": password, "es_admin": es_admin})
        if err:
            return self._responder({"error": err["error"]}, codigo=err["codigo"])
        return self._responder({"ok": True, "email": resultado["email"],
                                "password": password,
                                "es_admin": resultado["es_admin"]})

    def do_GET(self):
        p = urlparse(self.path)
        if p.path == "/login":
            return self._pagina_login()
        if p.path == "/logout":
            return self._cerrar_sesion()
        if not self._usuario_autenticado():
            return self._redirigir_login()
        if p.path not in _RUTAS_NO_ADMIN and not self._usuario_actual_admin():
            if p.path in ("/", "/index.html"):
                # Una cuenta no-admin no tiene nada que hacer en la consola
                # completa (lanzar transcripciones, trabajos, usuarios) —
                # se le manda directo a lo suyo.
                self.send_response(302)
                self.send_header("Location", "/esteno")
                self.end_headers()
                return
            return self._responder(
                {"error": "Requiere permisos de administrador"}, codigo=403)
        if p.path == "/api/yo":
            # Identidad real de la sesión actual — el módulo estenográfico
            # la usa para saber quién eres sin pedírtelo por URL/formulario.
            return self._responder({
                "email": self._usuario_actual_email(),
                "es_admin": self._usuario_actual_admin()})
        if p.path in ("/", "/index.html"):
            pagina = PAGINA.replace(
                "__REQUIERE_LOGIN__", "true" if self.requiere_login else "false")
            self._responder(pagina.encode("utf-8"),
                            "text/html; charset=utf-8")
        elif p.path in ("/esteno", "/esteno/", "/esteno.html"):
            self._responder(PAGINA_ESTENO.encode("utf-8"),
                            "text/html; charset=utf-8")
        elif p.path == "/api/esteno/estado":
            self._esteno_estado(parse_qs(p.query))
        elif p.path == "/api/esteno/segmentos":
            self._esteno_leer_segmentos(parse_qs(p.query))
        elif p.path == "/api/esteno/nombres":
            self._esteno_nombres(parse_qs(p.query))
        elif p.path == "/api/esteno/audio":
            self._esteno_servir_audio(parse_qs(p.query))
        elif p.path == "/api/sesiones":
            con = self._db()
            filas = [dict(r) for r in con.execute(
                "SELECT s.*, (SELECT COUNT(*) FROM participaciones p "
                " WHERE p.sesion_id = s.id) AS segmentos "
                "FROM sesiones s ORDER BY s.id DESC")]
            con.close()
            self._responder(filas)
        elif p.path == "/api/participaciones":
            q = parse_qs(p.query)
            sid = int(q.get("sesion", ["0"])[0] or 0)
            con = self._db()
            ses = con.execute("SELECT * FROM sesiones WHERE id=?",
                              (sid,)).fetchone()
            cols = {r[1] for r in con.execute(
                "PRAGMA table_info(participaciones)")}
            opcionales = [c for c in ("voz_orador", "voz_similitud",
                                      "revisado_ia", "motivo_ia")
                          if c in cols]
            extra = ("," + ",".join(opcionales)) if opcionales else ""
            filas = [dict(r) for r in con.execute(
                f"SELECT id, orador, inicio_seg, fin_seg, texto{extra} "
                "FROM participaciones WHERE sesion_id=? "
                "ORDER BY inicio_seg, id", (sid,))]
            resumenes = {r["ancla_id"]: r["resumen"] for r in con.execute(
                "SELECT ancla_id, resumen FROM resumenes WHERE sesion_id=?",
                (sid,))}
            estructura = None
            fila_est = con.execute(
                "SELECT datos FROM estructura WHERE sesion_id=?",
                (sid,)).fetchone()
            if fila_est and fila_est["datos"]:
                try:
                    estructura = json.loads(fila_est["datos"])
                except (ValueError, TypeError):
                    estructura = None
            con.close()
            self._responder({"sesion": dict(ses) if ses else None,
                             "filas": filas, "resumenes": resumenes,
                             "estructura": estructura})
        elif p.path == "/api/catalogo":
            self._responder(cargar_catalogo())
        elif p.path == "/api/contextos":
            # Lista de comisiones del JSON, para llenar el select. Si no hay
            # archivo o está mal, el lanzador sigue sirviendo para el pleno.
            info = {"disponible": False, "comisiones": [],
                    "tiene_transcriptor": os.path.isfile(
                        self.ruta_transcriptor),
                    "archivo": self.ruta_contextos}
            try:
                if os.path.isfile(self.ruta_contextos):
                    with open(self.ruta_contextos, encoding="utf-8") as f:
                        cfg = json.load(f)
                    info["disponible"] = True
                    info["comisiones"] = sorted(
                        (cfg.get("comisiones") or {}).keys())
                    info["tiene_mesa"] = bool(cfg.get("mesa_directiva"))
            except Exception as e:
                info["error"] = str(e)
            self._responder(info)
        elif p.path == "/api/eventos_parlamentarios":
            # Eventos reales (sesiones/comisiones) del sistema de registro
            # parlamentario, para elegir uno en vez de escribir URL/tipo a
            # mano — participantes salen automáticos de ahí.
            q = parse_qs(p.query)
            tipo = q.get("tipo", ["1"])[0]
            base = os.environ.get(
                "PARLAMENTARIO_API_URL",
                "https://parlamentario.congresoedomex.gob.mx/backend/api/"
                "eventos/ultimoseventos")
            import urllib.request
            try:
                with urllib.request.urlopen(f"{base}/{tipo}", timeout=15) as r:
                    datos = json.loads(r.read().decode("utf-8"))
                eventos = [{"id": e.get("id"),
                           "descripcion": (e.get("descripcion") or "").strip(),
                           "fecha": e.get("fecha"),
                           "tipoevento": e.get("tipoevento"),
                           "liga": e.get("liga")}
                          for e in datos.get("data", [])]
                self._responder(eventos)
            except Exception as e:
                self._responder(
                    {"error": f"No pude consultar el Congreso: {e}"},
                    codigo=502)
        elif p.path == "/api/trabajos_api":
            # Trabajos creados vía la API (fuente youtube/srt) — para verlos
            # y poder detenerlos desde aquí mismo, sin usar /docs.
            self._listar_trabajos_api()
        elif p.path == "/api/usuarios":
            self._listar_usuarios()
        elif p.path == "/api/exportar_word":
            q = parse_qs(p.query)
            try:
                sid = int(q.get("sesion", ["0"])[0] or 0)
            except (TypeError, ValueError):
                sid = 0
            if not sid:
                return self._responder({"error": "Falta el ID de la sesión"},
                                       codigo=400)
            # Si se pide 'corrector', el Word sale acotado solo a los
            # tramos que esa persona tomó o terminó en el módulo
            # estenográfico — para que el admin pueda revisar puntualmente
            # qué escribió, sin tener que abrir el editor en vivo.
            corrector = (q.get("corrector", [""])[0] or "").strip()
            # 'marcar_autores': Word completo (todos los oradores), pero con
            # una nota antes de cada tramo diciendo quién lo corrigió — para
            # revisar de un vistazo sin tener que sacar un Word por persona.
            marcar_autores = (q.get("marcar_autores", ["0"])[0] or "") == "1"
            con = self._db()
            try:
                ses = con.execute("SELECT * FROM sesiones WHERE id=?",
                                  (sid,)).fetchone()
                filas = con.execute(
                    "SELECT id, orador, texto, inicio_seg FROM participaciones "
                    "WHERE sesion_id=? ORDER BY inicio_seg", (sid,)).fetchall()
                rangos_corrector = None
                if corrector:
                    rangos_corrector = con.execute(
                        "SELECT inicio_seg, fin_seg FROM esteno_bloques "
                        "WHERE sesion_id=? AND (tomado_por=? OR "
                        "terminado_por=?)", (sid, corrector, corrector)
                    ).fetchall()
                    filas = [f for f in filas if any(
                        r["inicio_seg"] <= f["inicio_seg"] < r["fin_seg"]
                        for r in rangos_corrector)]
                bloques_autor = (con.execute(
                    "SELECT inicio_seg, fin_seg, tomado_por, terminado_por "
                    "FROM esteno_bloques WHERE sesion_id=? "
                    "ORDER BY inicio_seg", (sid,)).fetchall()
                    if marcar_autores else [])
                # Resúmenes ejecutivos guardados (si la tabla existe) — no
                # aplican al recorte por corrector, son de la sesión completa.
                try:
                    resumenes = ([] if corrector else con.execute(
                        "SELECT orador, resumen FROM resumenes "
                        "WHERE sesion_id=? ORDER BY ancla_id", (sid,)).fetchall())
                except sqlite3.Error:
                    resumenes = []
                # Estructura (encabezados del orden del día), si existe
                estructura = {"secciones": []}
                try:
                    fe = con.execute(
                        "SELECT datos FROM estructura WHERE sesion_id=?",
                        (sid,)).fetchone()
                    if fe and fe["datos"]:
                        estructura = json.loads(fe["datos"])
                except (sqlite3.Error, ValueError, TypeError):
                    pass
            finally:
                con.close()
            # Mapa ancla_id -> título de sección para insertar encabezados.
            titulos_seccion = {s["ancla_id"]: s["titulo"]
                               for s in estructura.get("secciones", [])
                               if s.get("ancla_id") is not None}

            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = Document()

            # --- CONFIGURAR MÁRGENES ---
            for section in doc.sections:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.top_margin = Cm(3.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(3)
                section.right_margin = Cm(2.5)

                # --- INSERTAR IMAGEN (MEMBRETE) ---
                # Ruta configurable con la variable de entorno MEMBRETE_IMG;
                # si no, se buscan varios nombres junto al script.
                header = section.header
                header_para = (header.paragraphs[0] if header.paragraphs
                               else header.add_paragraph())
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ruta_membrete = _ruta_membrete()
                if ruta_membrete:
                    try:
                        header_para.add_run().add_picture(
                            ruta_membrete, width=Inches(7.5))
                    except Exception as e:
                        print("No se pudo insertar el membrete:", e)

            # --- CONFIGURAR FUENTE ARIAL 11, JUSTIFICADO Y ESPACIADO ---
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            pf = style.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing = 1.15
            pf.space_after = Pt(12)

            # --- AGREGAR TÍTULOS AL DOCUMENTO ---
            titulo_texto = ("VERSIÓN ESTENOGRÁFICA: "
                            + (ses['titulo'] if ses and ses['titulo']
                               else 'SESIÓN LEGISLATIVA'))
            p_titulo = doc.add_paragraph()
            p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_titulo = p_titulo.add_run(titulo_texto.upper())
            run_titulo.bold = True
            run_titulo.font.size = Pt(12)

            p_fecha = doc.add_paragraph(
                "Fecha: " + _fecha_legible(ses['inicio'] if ses else None))
            p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_fecha.runs[0].font.size = Pt(10)
            p_fecha.runs[0].font.italic = True

            if corrector:
                p_corr = doc.add_paragraph(
                    "Tramos corregidos por: " + corrector + "  ("
                    + ", ".join(hms(r["inicio_seg"]) + "–" + hms(r["fin_seg"])
                                for r in rangos_corrector) + ")")
                p_corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_corr.runs[0].font.size = Pt(10)
                p_corr.runs[0].font.italic = True

            doc.add_paragraph("-" * 65).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Helper: inserta un encabezado de sección (punto del orden del día).
            def agregar_encabezado(titulo):
                pe = doc.add_paragraph()
                pe.paragraph_format.space_before = Pt(14)
                pe.paragraph_format.space_after = Pt(6)
                pe.paragraph_format.keep_with_next = True
                pe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run_e = pe.add_run(titulo)
                run_e.bold = True
                run_e.font.size = Pt(11)

            # Helper: vuelca un texto en uno o varios párrafos, respetando
            # los saltos de línea (\n) que introdujo la corrección de la IA.
            # El primer párrafo puede empezar con el nombre del orador en negrita.
            def volcar_intervencion(orador, texto):
                html_a_word(doc, texto, orador=orador,
                            fuente_nombre="Arial", fuente_pt=11,
                            justify=True, first_indent_cm=1)

            # Agrupamos por orador consecutivo para que cada intervención
            # sea un bloque coherente (aunque en la BD estén partida en filas).
            orador_actual = None
            ancla_actual = None
            buffer_texto = []

            def vaciar_buffer():
                if orador_actual is not None and buffer_texto:
                    # Si esta intervención inicia un punto del orden del día,
                    # insertamos primero su encabezado de sección.
                    if ancla_actual in titulos_seccion:
                        agregar_encabezado(titulos_seccion[ancla_actual])
                    volcar_intervencion(orador_actual,
                                        "\n".join(buffer_texto).strip())

            def agregar_marcador_autor(autor):
                pm = doc.add_paragraph()
                pm.paragraph_format.space_before = Pt(10)
                pm.paragraph_format.space_after = Pt(4)
                pm.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_m = pm.add_run(
                    "— Tramo corregido por: " + autor + " —" if autor
                    else "— Tramo sin corrección registrada —")
                run_m.italic = True
                run_m.font.size = Pt(9)
                run_m.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            # Punteros para saber, sobre la marcha, en qué bloque estenográfico
            # cae cada fila (ambas listas vienen ordenadas por inicio_seg).
            ptr_bloque = 0
            idx_bloque_actual = "__ninguno__"  # distinto de None a propósito

            for f in filas:
                texto_segmento = (f['texto'] or "").strip()
                if not texto_segmento:
                    continue
                if marcar_autores and bloques_autor:
                    while (ptr_bloque < len(bloques_autor)
                           and f['inicio_seg'] >= bloques_autor[ptr_bloque]['fin_seg']):
                        ptr_bloque += 1
                    en_bloque = (ptr_bloque < len(bloques_autor)
                                and f['inicio_seg'] >= bloques_autor[ptr_bloque]['inicio_seg'])
                    nuevo_idx = ptr_bloque if en_bloque else None
                    if nuevo_idx != idx_bloque_actual:
                        vaciar_buffer()
                        orador_actual = None
                        buffer_texto = []
                        idx_bloque_actual = nuevo_idx
                        b = bloques_autor[nuevo_idx] if nuevo_idx is not None else None
                        autor = (b["terminado_por"] or b["tomado_por"]) if b else None
                        agregar_marcador_autor(autor)
                if f['orador'] != orador_actual:
                    vaciar_buffer()
                    orador_actual = f['orador']
                    ancla_actual = f['id']
                    buffer_texto = [texto_segmento]
                else:
                    buffer_texto.append(texto_segmento)
            vaciar_buffer()

            # --- ANEXO: RESÚMENES EJECUTIVOS (si existen) ---
            resumenes_utiles = [r for r in resumenes
                                if (r["resumen"] or "").strip()]
            if resumenes_utiles:
                doc.add_page_break()
                p_anexo = doc.add_paragraph()
                p_anexo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_anexo = p_anexo.add_run("ANEXO — RESÚMENES EJECUTIVOS")
                r_anexo.bold = True
                r_anexo.font.size = Pt(12)
                doc.add_paragraph("-" * 65).alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in resumenes_utiles:
                    if r["orador"]:
                        p_ro = doc.add_paragraph()
                        p_ro.add_run((r["orador"] or "").upper()).bold = True
                    for linea in (r["resumen"] or "").replace("\r", "").split("\n"):
                        linea = linea.strip()
                        if linea:
                            doc.add_paragraph(linea)

            archivo_memoria = BytesIO()
            doc.save(archivo_memoria)
            datos_docx = archivo_memoria.getvalue()

            self.send_response(200)
            self.send_header(
                "Content-Type",
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document")
            nombre_archivo = f"Sesion_{sid}.docx"
            if corrector:
                sufijo = re.sub(r"[^A-Za-z0-9]+", "_", corrector.split("@")[0])
                nombre_archivo = f"Sesion_{sid}_{sufijo}.docx"
            elif marcar_autores:
                nombre_archivo = f"Sesion_{sid}_con_autores.docx"
            self.send_header(
                "Content-Disposition",
                f"attachment; filename={nombre_archivo}")
            self.send_header("Content-Length", str(len(datos_docx)))
            self.end_headers()
            self.wfile.write(datos_docx)

        elif p.path == "/api/transcripcion_estado":
            proc = Manejador.proc_activo
            corriendo = bool(proc and proc.poll() is None)
            codigo = (None if corriendo
                      else (proc.returncode if proc else None))
            lineas = ""
            log = Manejador.log_transcripcion
            if log and os.path.isfile(log):
                try:
                    with open(log, "rb") as f:
                        # solo la cola del log, para no mandar todo el archivo
                        try:
                            f.seek(-8000, os.SEEK_END)
                        except OSError:
                            f.seek(0)
                        crudo = f.read()
                    lineas = crudo.decode("utf-8", errors="replace")
                    # descartar una posible primera línea partida por el seek
                    if len(lineas) >= 8000 and "\n" in lineas:
                        lineas = lineas.split("\n", 1)[1]
                except Exception:
                    lineas = ""
            self._responder({"corriendo": corriendo,
                             "pid": (proc.pid if proc else None),
                             "codigo": codigo, "lineas": lineas})
        else:
            self._responder({"error": "no encontrado"}, codigo=404)

    def do_POST(self):
        p = urlparse(self.path)
        n = int(self.headers.get("Content-Length", 0) or 0)
        try:
            datos = json.loads(self.rfile.read(n) or b"{}")
        except json.JSONDecodeError:
            return self._responder({"error": "JSON inválido"}, codigo=400)

        if p.path == "/login":
            return self._procesar_login(datos)
        if not self._usuario_autenticado():
            return self._responder({"error": "No autenticado"}, codigo=401)
        if p.path not in _RUTAS_NO_ADMIN and not self._usuario_actual_admin():
            return self._responder(
                {"error": "Requiere permisos de administrador"}, codigo=403)

        # --- Lanzador de transcripción (costura 2) ---------------------
        # Estas rutas no tocan la base de datos, así que se resuelven antes
        # de abrir la conexión.
        if p.path == "/api/transcribir":
            return self._lanzar_transcripcion(datos)
        if p.path == "/api/detener":
            return self._detener_transcripcion()
        if p.path == "/api/transcripciones_evento":
            return self._crear_transcripcion_evento(datos)
        if p.path == "/api/trabajos_api/detener":
            return self._detener_trabajo_api(datos.get("id"))
        if p.path == "/api/usuarios":
            return self._crear_usuario(datos)

        con = self._db()
        try:
            if p.path == "/api/esteno/configurar":
                return self._esteno_configurar(con, datos)
            elif p.path == "/api/esteno/reasignar":
                return self._esteno_reasignar(con, datos)
            elif p.path == "/api/esteno/tomar":
                return self._esteno_tomar(con, datos)
            elif p.path == "/api/esteno/latido":
                return self._esteno_latido(con, datos)
            elif p.path == "/api/esteno/soltar":
                return self._esteno_soltar(con, datos)
            elif p.path == "/api/esteno/guardar":
                return self._esteno_guardar(con, datos)
            elif p.path == "/api/esteno/audio_ruta":
                return self._esteno_audio_ruta(con, datos)
            elif p.path == "/api/actualizar":
                ids = [int(i) for i in datos.get("ids", [])]
                orador = (datos.get("orador") or "").strip()
                sid = int(datos.get("sesion_id", 0) or 0)
                if not ids or not orador:
                    return self._responder({"error": "faltan datos"},
                                           codigo=400)
                marcas = ",".join("?" * len(ids))
                # Una corrección manual manda sobre el veredicto del
                # corrector: se limpia la insignia de IA de esas filas (si
                # la base tiene esas columnas).
                cols = {r[1] for r in con.execute(
                    "PRAGMA table_info(participaciones)")}
                limpiar = (", revisado_ia=NULL, motivo_ia=NULL"
                           if "revisado_ia" in cols else "")
                cur = con.execute(
                    f"UPDATE participaciones SET orador=?{limpiar} "
                    f"WHERE id IN ({marcas})", [orador] + ids)
                con.commit()
                # Unir en la base los registros que quedaron consecutivos
                # con el mismo orador (los corregidos y sus vecinos)
                unidos = compactar(con, sid, solo_ids=ids) if sid else 0
                self._responder({"ok": True, "cambios": cur.rowcount,
                                 "unidos": unidos})
            elif p.path == "/api/dividir":
                # Separar dos oradores dentro de un mismo bloque. 'corte' es
                # el desplazamiento en caracteres sobre el texto del bloque
                # unido con espacios (igual que lo arma el front).
                sid = int(datos.get("sesion_id", 0) or 0)
                ids = [int(i) for i in datos.get("ids", [])]
                corte = int(datos.get("corte", 0) or 0)
                orador1 = (datos.get("orador1") or "").strip()
                orador2 = (datos.get("orador2") or "").strip()
                if not (ids and orador1 and orador2):
                    return self._responder({"error": "faltan datos"},
                                           codigo=400)
                if not self._usuario_actual_admin():
                    # Un corrector solo puede dividir dentro del bloque
                    # estenográfico que tiene abierto ahora mismo (no
                    # cualquier fila de la sesión).
                    email = self._usuario_actual_email()
                    filas_chk = con.execute(
                        f"SELECT inicio_seg FROM participaciones WHERE "
                        f"id IN ({','.join('?' * len(ids))}) AND "
                        f"sesion_id=?", ids + [sid]).fetchall()
                    bloques_corrector = con.execute(
                        "SELECT inicio_seg, fin_seg FROM esteno_bloques "
                        "WHERE sesion_id=? AND tomado_por=? AND "
                        "estado='editando'", (sid, email)).fetchall()
                    propios = len(filas_chk) == len(ids) and all(
                        any(b["inicio_seg"] <= f["inicio_seg"] < b["fin_seg"]
                            for b in bloques_corrector)
                        for f in filas_chk)
                    if not propios:
                        return self._responder(
                            {"error": "Solo puedes dividir segmentos del "
                                      "bloque que tienes abierto"},
                            codigo=403)
                cols = {r[1] for r in con.execute(
                    "PRAGMA table_info(participaciones)")}
                limpiar = (", revisado_ia=NULL, motivo_ia=NULL"
                           if "revisado_ia" in cols else "")
                filas = {r["id"]: r for r in con.execute(
                    f"SELECT * FROM participaciones WHERE id IN "
                    f"({','.join('?' * len(ids))})", ids)}
                orden = [filas[i] for i in ids if i in filas]
                if not orden:
                    return self._responder({"error": "filas no encontradas"},
                                           codigo=400)
                # Localizar en qué fila (y en qué punto de ella) cae el corte
                pos, fila_k, off_local = 0, None, 0
                for k, f in enumerate(orden):
                    txt = f["texto"] or ""
                    if corte <= pos:
                        fila_k, off_local = k, 0
                        break
                    if corte < pos + len(txt):
                        fila_k, off_local = k, corte - pos
                        break
                    pos += len(txt) + 1          # +1 por el espacio de unión
                if fila_k is None:
                    return self._responder(
                        {"error": "el corte deja la segunda parte vacía"},
                        codigo=400)
                # 1ª parte: filas estrictamente anteriores al corte -> orador1
                antes = [orden[j]["id"] for j in range(fila_k)]
                if antes:
                    m = ",".join("?" * len(antes))
                    con.execute(
                        f"UPDATE participaciones SET orador=?{limpiar} "
                        f"WHERE id IN ({m})", [orador1] + antes)
                nuevo_id = None
                if off_local > 0:
                    # El corte cae DENTRO de una fila: se parte en dos
                    f = orden[fila_k]
                    txt = f["texto"] or ""
                    t1, t2 = txt[:off_local].strip(), txt[off_local:].strip()
                    ini_f = f["inicio_seg"] or 0
                    fin_f = f["fin_seg"] or ini_f
                    corte_seg = ini_f + (fin_f - ini_f) * (off_local / len(txt))
                    con.execute(
                        f"UPDATE participaciones SET orador=?, texto=?, "
                        f"fin_seg=?, fin_hms=?{limpiar} WHERE id=?",
                        (orador1, t1, corte_seg, hms(corte_seg), f["id"]))
                    campos = ("sesion_id, orador, inicio_seg, fin_seg, "
                              "inicio_hms, fin_hms, texto")
                    valores = [sid or f["sesion_id"], orador2, corte_seg,
                               fin_f, hms(corte_seg), hms(fin_f), t2]
                    if "fuente" in cols:
                        campos += ", fuente"
                        valores.append("manual")
                    cur = con.execute(
                        f"INSERT INTO participaciones ({campos}) VALUES "
                        f"({','.join('?' * len(valores))})", valores)
                    nuevo_id = cur.lastrowid
                    despues = [orden[j]["id"]
                               for j in range(fila_k + 1, len(orden))]
                else:
                    # Corte limpio en frontera de fila
                    despues = [orden[j]["id"]
                               for j in range(fila_k, len(orden))]
                if despues:
                    m = ",".join("?" * len(despues))
                    con.execute(
                        f"UPDATE participaciones SET orador=?{limpiar} "
                        f"WHERE id IN ({m})", [orador2] + despues)
                con.commit()
                self._responder({"ok": True, "nuevo_id": nuevo_id})
            elif p.path == "/api/renombrar":
                sid = int(datos.get("sesion_id", 0) or 0)
                de = (datos.get("de") or "").strip()
                a = (datos.get("a") or "").strip()
                if not (sid and de and a):
                    return self._responder({"error": "faltan datos"},
                                           codigo=400)
                cur = con.execute(
                    "UPDATE participaciones SET orador=? "
                    "WHERE sesion_id=? AND orador=?", (a, sid, de))
                con.commit()
                unidos = compactar(con, sid)
                self._responder({"ok": True, "cambios": cur.rowcount,
                                 "unidos": unidos})
            elif p.path == "/api/compactar":
                sid = int(datos.get("sesion_id", 0) or 0)
                if not sid:
                    return self._responder({"error": "faltan datos"},
                                           codigo=400)
                unidos = compactar(con, sid)
                self._responder({"ok": True, "unidos": unidos})
            elif p.path == "/api/textos":
                cambios = datos.get("cambios") or []
                n = 0
                for c in cambios:
                    try:
                        rid = int(c.get("id"))
                    except (TypeError, ValueError):
                        continue
                    cur = con.execute(
                        "UPDATE participaciones SET texto=? WHERE id=?",
                        (c.get("texto", ""), rid))
                    n += cur.rowcount
                con.commit()
                self._responder({"ok": True, "cambios": n})
            elif p.path == "/api/corregir_estilo":
                sid = int(datos.get("sesion_id", 0) or 0)
                if not sid:
                    return self._responder(
                        {"error": "Falta el ID de la sesión"}, codigo=400)
                if not os.environ.get("ANTHROPIC_API_KEY", "").strip():
                    return self._responder(
                        {"error": "No hay clave API configurada "
                                  "(variable ANTHROPIC_API_KEY)."}, codigo=400)

                # 1) Unimos bloques consecutivos del mismo orador para corregir
                #    intervenciones completas (la IA reparte mejor los párrafos).
                compactar(con, sid)

                # 2) Nos aseguramos de tener columna de respaldo del original,
                #    para poder comparar o revertir la corrección.
                cols = {r[1] for r in con.execute(
                    "PRAGMA table_info(participaciones)")}
                if "texto_original" not in cols:
                    con.execute("ALTER TABLE participaciones "
                                "ADD COLUMN texto_original TEXT")
                    con.commit()

                catalogo = cargar_catalogo()
                filas = con.execute(
                    "SELECT id, texto, texto_original FROM participaciones "
                    "WHERE sesion_id=? AND texto IS NOT NULL", (sid,)).fetchall()
                modificados = 0
                errores = []
                for f in filas:
                    texto_actual = (f["texto"] or "").strip()
                    if len(texto_actual) < 10:
                        continue
                    res = corregir_texto_ia(texto_actual, catalogo)
                    if "texto_corregido" in res:
                        # Guardamos el original solo la primera vez.
                        if not (f["texto_original"] or "").strip():
                            con.execute(
                                "UPDATE participaciones SET texto_original=? "
                                "WHERE id=?", (texto_actual, f["id"]))
                        con.execute(
                            "UPDATE participaciones SET texto=? WHERE id=?",
                            (res["texto_corregido"], f["id"]))
                        modificados += 1
                    elif res.get("error"):
                        errores.append(res["error"])
                con.commit()
                resp = {"ok": True, "cambios": modificados}
                if errores:
                    # Reportamos solo el primer error para no saturar.
                    resp["aviso"] = (f"{len(errores)} segmento(s) no se "
                                     f"pudieron corregir. Ej.: {errores[0]}")
                self._responder(resp)
            elif p.path == "/api/estructurar":
                sid = int(datos.get("sesion_id", 0) or 0)
                if not sid:
                    return self._responder(
                        {"error": "Falta el ID de la sesión"}, codigo=400)
                if not os.environ.get("ANTHROPIC_API_KEY", "").strip():
                    return self._responder(
                        {"error": "No hay clave API configurada "
                                  "(variable ANTHROPIC_API_KEY)."}, codigo=400)
                turnos = turnos_de_sesion(con, sid)
                res = estructurar_ia(turnos)
                if res.get("error"):
                    return self._responder({"error": res["error"]}, codigo=502)
                con.execute(
                    "INSERT INTO estructura (sesion_id, datos, creado) "
                    "VALUES (?,?,?) "
                    "ON CONFLICT(sesion_id) DO UPDATE SET "
                    "datos=excluded.datos, creado=excluded.creado",
                    (sid, json.dumps(res, ensure_ascii=False),
                     datetime.now().isoformat(timespec="seconds")))
                con.commit()
                self._responder({
                    "ok": True,
                    "secciones": len(res["secciones"]),
                    "estructura": res})
            elif p.path == "/api/borrar_estructura":
                sid = int(datos.get("sesion_id", 0) or 0)
                if sid:
                    con.execute("DELETE FROM estructura WHERE sesion_id=?",
                                (sid,))
                    con.commit()
                self._responder({"ok": True})
            elif p.path == "/api/resumen":
                orador = (datos.get("orador") or "").strip()
                texto = (datos.get("texto") or "").strip()
                sid = int(datos.get("sesion_id", 0) or 0)
                ancla = int(datos.get("ancla_id", 0) or 0)
                if not texto:
                    return self._responder({"error": "sin texto"},
                                           codigo=400)
                res = generar_resumen(orador, texto)
                # Si se generó automáticamente, se guarda pegado al bloque
                if res.get("modo") == "auto" and res.get("resumen") and ancla:
                    con.execute(
                        "INSERT INTO resumenes "
                        "(sesion_id, ancla_id, orador, resumen, creado) "
                        "VALUES (?,?,?,?,?) "
                        "ON CONFLICT(ancla_id) DO UPDATE SET "
                        "resumen=excluded.resumen, orador=excluded.orador, "
                        "creado=excluded.creado",
                        (sid, ancla, orador, res["resumen"],
                         datetime.now().isoformat(timespec="seconds")))
                    con.commit()
                self._responder(res)
            elif p.path == "/api/guardar_resumen":
                # Guardar un resumen editado o pegado a mano
                sid = int(datos.get("sesion_id", 0) or 0)
                ancla = int(datos.get("ancla_id", 0) or 0)
                orador = (datos.get("orador") or "").strip()
                resumen = (datos.get("resumen") or "").strip()
                if not ancla:
                    return self._responder({"error": "falta ancla"},
                                           codigo=400)
                if resumen:
                    con.execute(
                        "INSERT INTO resumenes "
                        "(sesion_id, ancla_id, orador, resumen, creado) "
                        "VALUES (?,?,?,?,?) "
                        "ON CONFLICT(ancla_id) DO UPDATE SET "
                        "resumen=excluded.resumen, orador=excluded.orador, "
                        "creado=excluded.creado",
                        (sid, ancla, orador, resumen,
                         datetime.now().isoformat(timespec="seconds")))
                else:
                    con.execute("DELETE FROM resumenes WHERE ancla_id=?",
                                (ancla,))
                con.commit()
                self._responder({"ok": True})
            elif p.path == "/api/borrar_resumen":
                ancla = int(datos.get("ancla_id", 0) or 0)
                con.execute("DELETE FROM resumenes WHERE ancla_id=?",
                            (ancla,))
                con.commit()
                self._responder({"ok": True})
            else:
                self._responder({"error": "no encontrado"}, codigo=404)
        finally:
            con.close()


def main():
    ap = argparse.ArgumentParser(
        description="Interfaz web local para revisar y corregir oradores.")
    ap.add_argument("--db", default="sesiones.db",
                    help="Base de datos SQLite (default: sesiones.db)")
    ap.add_argument("--puerto", type=int, default=8756,
                    help="Puerto local (default: 8756)")
    ap.add_argument("--transcriptor", default="transcribir_en_vivo_c3.py",
                    help="script de transcripción que arranca el botón "
                         "(default: transcribir_en_vivo_c3.py)")
    ap.add_argument("--contextos", default="contextos.json",
                    help="archivo de contextos para llenar el select de "
                         "comisiones (default: contextos.json)")
    ap.add_argument("--api-interna", default="http://api:8000",
                    help="URL de la API (api/) para crear trabajos desde "
                        "un evento real del Congreso (default: "
                        "http://api:8000, el nombre del servicio en Docker)")
    ap.add_argument("--host", default="127.0.0.1",
                    help="Dirección donde escuchar (default: 127.0.0.1; "
                        "usa 0.0.0.0 para exponerlo fuera de esta máquina, "
                        "p. ej. dentro de un contenedor Docker)")
    ap.add_argument("--requiere-login", action="store_true",
                    help="exige iniciar sesión con un usuario de la API "
                        "(mismo MySQL/usuarios de api/; ver README) antes "
                        "de usar la interfaz. Recomendado si --host no es "
                        "127.0.0.1")
    args = ap.parse_args()

    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    if not os.path.isfile(args.db):
        print(f"No encuentro la base de datos: {os.path.abspath(args.db)}")
        print("Corre primero transcribir_en_vivo.py, o indica la ruta "
              "con --db.")
        sys.exit(1)

    if args.requiere_login:
        global _verificar_password, _obtener_usuario, _crear_token
        global _jwt, _jwt_secret, _jwt_algoritmo
        from jose import jwt as _jwt_modulo
        from api.config import settings
        from api.security import (crear_token, obtener_usuario_por_email,
                                  verificar_password)
        _verificar_password = verificar_password
        _obtener_usuario = obtener_usuario_por_email
        _crear_token = crear_token
        _jwt = _jwt_modulo
        _jwt_secret = settings.jwt_secret
        _jwt_algoritmo = settings.jwt_algoritmo
        Manejador.requiere_login = True
        print("Login requerido: usuarios de MySQL (api/), ver /login")

    Manejador.ruta_db = args.db
    Manejador.ruta_transcriptor = args.transcriptor
    Manejador.ruta_contextos = args.contextos
    Manejador.api_interna = args.api_interna
    direccion = f"http://{args.host}:{args.puerto}"
    servidor = ThreadingHTTPServer((args.host, args.puerto), Manejador)

    catalogo = cargar_catalogo()
    print(f"Base de datos : {os.path.abspath(args.db)}")
    print(f"Catálogo      : {len(catalogo)} diputados"
          if catalogo else
          "Catálogo      : sin diputados.txt (opcional)")
    print(f"Interfaz      : {direccion}  (Ctrl+C para cerrar)")

    if args.host in ("127.0.0.1", "localhost"):
        # Solo tiene sentido abrir un navegador local si el servidor
        # también escucha en local (no dentro de un contenedor Docker).
        def _abrir():
            try:
                webbrowser.open(direccion)
            except Exception:
                pass
        threading.Timer(1.0, _abrir).start()
    try:
        servidor.serve_forever()
    except KeyboardInterrupt:
        print("\nCerrado. Tus cambios ya quedaron guardados en la "
              "base de datos.")


if __name__ == "__main__":
    main()
